"""
模板编辑器服务
提供 docx 解析和导出功能
"""

import io
import re
from typing import Dict, Any, List, Set, Optional
from docx import Document
from loguru import logger
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from sqlalchemy.orm import selectinload

from .mappers import DocxToProseMirrorMapper, ProseMirrorToDocxMapper
from .logging import template_log_service
from .models import DocumentTemplate, TemplatePlaceholder


class TemplateEditorService:
    """模板编辑器服务"""

    def __init__(self):
        self.docx_to_pm_mapper = DocxToProseMirrorMapper()
        self.pm_to_docx_mapper = ProseMirrorToDocxMapper()
    
    def extract_placeholders(self, prosemirror_json: Dict[str, Any]) -> Dict[str, Any]:
        """
        从 ProseMirror JSON 中提取占位符
        支持两种格式：
        1. placeholder 节点（type="placeholder", attrs.fieldKey）
        2. 文本中的 {{name}} 格式
        
        Args:
            prosemirror_json: ProseMirror JSON 格式的文档
            
        Returns:
            占位符信息字典，格式：
            {
                'placeholders': ['name', 'date', 'address'],  # 去重后的占位符名称列表
                'metadata': {
                    'name': {'count': 3, 'positions': [...]},  # 每个占位符的元数据
                    ...
                }
            }
        """
        placeholders_set: Set[str] = set()
        placeholder_metadata: Dict[str, Dict[str, Any]] = {}
        
        def traverse_node(node: Dict[str, Any], path: List[str] = []):
            """递归遍历 ProseMirror 节点树"""
            node_type = node.get("type")
            
            # 如果是 placeholder 节点，直接提取 fieldKey
            if node_type == "placeholder":
                attrs = node.get("attrs", {})
                field_key = attrs.get("fieldKey") or attrs.get("field_key")
                if field_key:
                    placeholder_name = str(field_key).strip()
                    if placeholder_name:
                        placeholders_set.add(placeholder_name)
                        
                        # 记录占位符的元数据
                        if placeholder_name not in placeholder_metadata:
                            placeholder_metadata[placeholder_name] = {
                                'count': 0,
                                'positions': []
                            }
                        
                        placeholder_metadata[placeholder_name]['count'] += 1
                        placeholder_metadata[placeholder_name]['positions'].append({
                            'path': path.copy(),
                            'node_type': 'placeholder',
                            'fieldKey': placeholder_name
                        })
            
            # 如果是文本节点，检查文本内容中的 {{placeholder}} 格式
            elif node_type == "text":
                text = node.get("text", "")
                # 使用正则表达式提取 {{placeholder}} 格式的占位符
                pattern = r'\{\{([^}]+)\}\}'
                matches = re.finditer(pattern, text)
                
                for match in matches:
                    placeholder_name = match.group(1).strip()
                    if placeholder_name:
                        placeholders_set.add(placeholder_name)
                        
                        # 记录占位符的元数据
                        if placeholder_name not in placeholder_metadata:
                            placeholder_metadata[placeholder_name] = {
                                'count': 0,
                                'positions': []
                            }
                        
                        placeholder_metadata[placeholder_name]['count'] += 1
                        placeholder_metadata[placeholder_name]['positions'].append({
                            'path': path.copy(),
                            'text': text,
                            'start': match.start(),
                            'end': match.end()
                        })
            
            # 递归处理子节点
            content = node.get("content", [])
            if isinstance(content, list):
                for idx, child in enumerate(content):
                    child_path = path + [f"{node_type}[{idx}]"]
                    traverse_node(child, child_path)
        
        # 从根节点开始遍历
        traverse_node(prosemirror_json)
        
        # 返回结果
        return {
            'placeholders': sorted(list(placeholders_set)),
            'metadata': placeholder_metadata
        }

    def parse_docx_to_prosemirror(self, docx_bytes: bytes) -> Dict[str, Any]:
        """
        解析 docx 文件为 ProseMirror JSON 格式

        Args:
            docx_bytes: docx 文件的字节流

        Returns:
            ProseMirror JSON 格式的文档

        Raises:
            ValueError: 如果文件格式无效
        """
        try:
            # 从字节流创建 Document 对象
            doc = Document(io.BytesIO(docx_bytes))

            template_log_service.clear()
            template_log_service.add_record("parse", "read_docx", "start")
            # 使用映射器转换为 ProseMirror JSON
            prosemirror_json = self.docx_to_pm_mapper.map_document(doc)
            template_log_service.add_record(
                "parse",
                "read_docx",
                "success",
                {"elements": len(prosemirror_json.get("content", []))},
            )

            # 记录解析结果（用于调试）
            content_count = len(prosemirror_json.get("content", []))
            logger.info(f"解析 docx 成功，共 {content_count} 个元素")
            if content_count > 0:
                first_element = prosemirror_json["content"][0]
                logger.info(f"第一个元素类型: {first_element.get('type')}")

            return prosemirror_json

        except Exception as e:
            logger.error(f"解析 docx 文件失败: {str(e)}")
            raise ValueError(f"无法解析 docx 文件: {str(e)}") from e

    def export_prosemirror_to_docx(
        self, prosemirror_json: Dict[str, Any], is_narrative_style: bool = False
    ) -> Dict[str, Any]:
        """
        从 ProseMirror JSON 导出为 docx 文件

        Args:
            prosemirror_json: ProseMirror JSON 格式的文档
            is_narrative_style: 是否为陈述式模板（用于决定是否显示表格边框）

        Returns:
            包含 docx 字节流和警告信息的字典

        Raises:
            ValueError: 如果 JSON 格式无效
        """
        try:
            # 验证 JSON 格式
            if not isinstance(prosemirror_json, dict):
                raise ValueError("ProseMirror JSON 必须是字典类型")

            if prosemirror_json.get("type") != "doc":
                raise ValueError(
                    "ProseMirror JSON 的根节点类型必须是 'doc'"
                )

            # 使用映射器转换为 docx
            logger.info("开始转换 ProseMirror JSON 到 DOCX")
            template_log_service.clear()
            template_log_service.add_record("export", "map_document", "start")
            doc = self.pm_to_docx_mapper.map_document(prosemirror_json, is_narrative_style=is_narrative_style)
            template_log_service.add_record("export", "map_document", "success")
            logger.info("DOCX 文档创建成功")
            
            # 验证表格列宽（在保存前）
            if doc.tables:
                from docx.oxml.ns import qn
                ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
                for table_idx, table in enumerate(doc.tables):
                    tbl = table._tbl
                    tbl_grid = tbl.tblGrid
                    if tbl_grid is not None:
                        grid_cols = tbl_grid.findall(f'.//{ns}gridCol')
                        logger.info(f"[SERVICE] 保存前验证 - 表格{table_idx} tblGrid 列宽:")
                        for i, grid_col in enumerate(grid_cols):
                            width = grid_col.get(qn('w:w'))
                            width_type = grid_col.get(qn('w:type'))
                            logger.info(f"[SERVICE]   列 {i}: width={width}, type={width_type}")

            # 保存到字节流
            logger.info("开始保存 DOCX 到字节流")
            output = io.BytesIO()
            doc.save(output)
            logger.info("DOCX 保存到字节流成功")
            output.seek(0)
            docx_bytes = output.getvalue()
            logger.info(f"DOCX 字节流大小: {len(docx_bytes)} bytes")
            template_log_service.add_record(
                "export", "save_doc", "success", {"size": len(docx_bytes)}
            )
            
            # 验证保存后的列宽（重新读取）
            from docx import Document
            doc_verify = Document(io.BytesIO(docx_bytes))
            if doc_verify.tables:
                from docx.oxml.ns import qn
                ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
                for table_idx, table in enumerate(doc_verify.tables):
                    tbl = table._tbl
                    tbl_grid = tbl.tblGrid
                    if tbl_grid is not None:
                        grid_cols = tbl_grid.findall(f'.//{ns}gridCol')
                        logger.info(f"[SERVICE] 保存后验证 - 表格{table_idx} tblGrid 列宽:")
                        for i, grid_col in enumerate(grid_cols):
                            width = grid_col.get(qn('w:w'))
                            width_type = grid_col.get(qn('w:type'))
                            logger.info(f"[SERVICE]   列 {i}: width={width}, type={width_type}")
            
            warnings = getattr(self.pm_to_docx_mapper, "export_warnings", [])
            for warning in warnings:
                template_log_service.with_warning(warning)

            return {
                "docx": docx_bytes,
                "warnings": warnings,
                "logs": template_log_service.get_records(),
            }

        except UnicodeEncodeError as e:
            logger.error(f"编码错误 - 位置: {e.start}-{e.end}, 对象: {e.object}, 错误: {str(e)}")
            raise ValueError(f"无法导出 docx 文件: 编码错误 - {str(e)}") from e
        except Exception as e:
            logger.error(f"导出 docx 文件失败: {str(e)}", exc_info=True)
            raise ValueError(f"无法导出 docx 文件: {str(e)}") from e


class TemplateService:
    """模板管理服务"""
    
    async def create_template(
        self,
        db: AsyncSession,
        name: str,
        prosemirror_json: Dict[str, Any],
        docx_url: Optional[str] = None,
        description: Optional[str] = None,
        category: Optional[str] = None,
        status: str = "draft",
        created_by_id: Optional[int] = None,
    ) -> DocumentTemplate:
        """
        创建文书模板
        
        Args:
            db: 数据库会话
            name: 模板名称
            prosemirror_json: ProseMirror JSON 内容
            docx_url: DOCX 文件在 COS 中的 URL
            description: 模板描述
            category: 分类名称
            status: 状态（draft/published）
            created_by_id: 创建人ID
            
        Returns:
            创建的模板对象
        """
        # 提取占位符
        placeholder_info = template_editor_service.extract_placeholders(prosemirror_json)
        placeholder_names = placeholder_info.get('placeholders', [])
        
        # 创建模板对象
        template = DocumentTemplate(
            name=name,
            description=description,
            category=category,
            status=status,
            prosemirror_json=prosemirror_json,
            docx_url=docx_url,
            created_by_id=created_by_id,
            updated_by_id=created_by_id,
        )
        
        db.add(template)
        await db.flush()  # 先 flush 以获取 template.id
        
        # 判断是否为陈述式模板，确定适用的模板类型
        is_narrative_style = category and ("陈述" in category or category == "陈述式")
        applicable_category = "陈述式" if is_narrative_style else ("要素式" if category and ("要素" in category or category == "要素式") else None)
        
        # 为每个占位符名称创建或获取 TemplatePlaceholder 对象并关联到模板
        placeholder_objects = []
        for placeholder_name in placeholder_names:
            placeholder = None
            
            if is_narrative_style:
                # 陈述式：优先查询适用于"陈述式"的 type="text" 占位符
                result = await db.execute(
                    select(TemplatePlaceholder).where(
                        TemplatePlaceholder.name == placeholder_name,
                        TemplatePlaceholder.applicable_template_category == "陈述式",
                        TemplatePlaceholder.type == "text"
                    )
                )
                placeholder = result.scalar_one_or_none()
                
                # 如果没有找到 type="text" 的，检查是否存在其他类型的（避免唯一性约束冲突）
                if not placeholder:
                    result = await db.execute(
                        select(TemplatePlaceholder).where(
                            TemplatePlaceholder.name == placeholder_name,
                            TemplatePlaceholder.applicable_template_category == "陈述式"
                        )
                    )
                    existing_placeholder = result.scalar_one_or_none()
                    if existing_placeholder:
                        # 如果已存在其他类型的占位符，复用它（不强制改为 text）
                        placeholder = existing_placeholder
                        logger.info(f"复用已存在的占位符: {placeholder_name} (type={placeholder.type}, 适用于陈述式)")
                
                if not placeholder:
                    # 创建新的 type="text" 占位符，标记为适用于"陈述式"
                    placeholder = TemplatePlaceholder(
                        name=placeholder_name,
                        type="text",  # 陈述式统一使用 text
                        options=None,  # 无选项
                        applicable_template_category="陈述式",  # 标记为适用于陈述式模板
                        created_by_id=created_by_id,
                        updated_by_id=created_by_id,
                    )
                    db.add(placeholder)
                    await db.flush()  # flush 以获取 placeholder.id
                    logger.info(f"创建占位符: {placeholder_name} (type=text, 适用于陈述式)")
                else:
                    logger.info(f"使用已存在的占位符: {placeholder_name} (type={placeholder.type}, 适用于陈述式)")
            else:
                # 要素式：查询适用于"要素式"的占位符（任意类型）
                result = await db.execute(
                    select(TemplatePlaceholder).where(
                        TemplatePlaceholder.name == placeholder_name,
                        TemplatePlaceholder.applicable_template_category == "要素式"
                    )
                )
                placeholder = result.scalar_one_or_none()
                
                if not placeholder:
                    # 创建新占位符（默认 type="text"），标记为适用于"要素式"
                    placeholder = TemplatePlaceholder(
                        name=placeholder_name,
                        type="text",  # 默认类型
                        options=None,  # 默认值
                        applicable_template_category="要素式",  # 标记为适用于要素式模板
                        created_by_id=created_by_id,
                        updated_by_id=created_by_id,
                    )
                    db.add(placeholder)
                    await db.flush()  # flush 以获取 placeholder.id
                    logger.info(f"创建占位符: {placeholder_name} (type=text, 适用于要素式)")
                else:
                    logger.info(f"使用已存在的占位符: {placeholder_name} (type={placeholder.type}, 适用于要素式)")
            
            placeholder_objects.append(placeholder)
        
        # 关联占位符到模板（直接操作中间表，避免触发懒加载）
        if placeholder_objects:
            from .models import template_placeholder_association
            # 直接插入到中间表，避免访问关系字段触发懒加载
            for placeholder in placeholder_objects:
                await db.execute(
                    template_placeholder_association.insert().values(
                        template_id=template.id,
                        placeholder_id=placeholder.id
                    )
                )
        
        await db.commit()
        await db.refresh(template)
        
        logger.info(f"创建模板成功: {template.id}, 名称: {name}, 占位符数量: {len(placeholder_names)}")
        
        return template
    
    async def get_template(self, db: AsyncSession, template_id: int) -> Optional[DocumentTemplate]:
        """根据ID获取文书模板"""
        result = await db.execute(
            select(DocumentTemplate)
            .where(DocumentTemplate.id == template_id)
            .options(selectinload(DocumentTemplate.placeholders))
        )
        return result.scalar_one_or_none()
    
    async def list_templates(
        self,
        db: AsyncSession,
        status: Optional[str] = None,
        category: Optional[str] = None,
        skip: int = 0,
        limit: int = 100,
    ) -> List[DocumentTemplate]:
        """列出文书模板"""
        query = select(DocumentTemplate).options(selectinload(DocumentTemplate.placeholders))
        
        if status:
            query = query.where(DocumentTemplate.status == status)
        if category:
            query = query.where(DocumentTemplate.category == category)
        
        query = query.offset(skip).limit(limit).order_by(DocumentTemplate.created_at.desc())
        
        result = await db.execute(query)
        return list(result.scalars().unique().all())
    
    async def update_template(
        self,
        db: AsyncSession,
        template: DocumentTemplate,
        name: Optional[str] = None,
        prosemirror_json: Optional[Dict[str, Any]] = None,
        docx_url: Optional[str] = None,
        description: Optional[str] = None,
        category: Optional[str] = None,
        status: Optional[str] = None,
        updated_by_id: Optional[int] = None,
    ) -> DocumentTemplate:
        """更新文书模板"""
        if name is not None:
            template.name = name
        if description is not None:
            template.description = description
        if category is not None:
            template.category = category
        if status is not None:
            template.status = status
        if docx_url is not None:
            template.docx_url = docx_url
        if updated_by_id is not None:
            template.updated_by_id = updated_by_id
        
        # 如果更新了 ProseMirror JSON，重新提取占位符并更新关联关系
        if prosemirror_json is not None:
            template.prosemirror_json = prosemirror_json
            placeholder_info = template_editor_service.extract_placeholders(prosemirror_json)
            placeholder_names = placeholder_info.get('placeholders', [])
            
            from .models import template_placeholder_association
            
            # 获取当前模板的所有关联占位符ID
            existing_associations_result = await db.execute(
                select(template_placeholder_association.c.placeholder_id).where(
                    template_placeholder_association.c.template_id == template.id
                )
            )
            existing_placeholder_ids = {row[0] for row in existing_associations_result.all()}
            
            # 判断是否为陈述式模板，确定适用的模板类型
            is_narrative_style = template.category and ("陈述" in template.category or template.category == "陈述式")
            applicable_category = "陈述式" if is_narrative_style else ("要素式" if template.category and ("要素" in template.category or template.category == "要素式") else None)
            
            # 获取文档中占位符名称对应的占位符对象
            placeholder_objects = []
            placeholder_name_to_id: Dict[str, int] = {}
            
            for placeholder_name in placeholder_names:
                placeholder = None
                
                if is_narrative_style:
                    # 陈述式：优先查询适用于"陈述式"的 type="text" 占位符
                    result = await db.execute(
                        select(TemplatePlaceholder).where(
                            TemplatePlaceholder.name == placeholder_name,
                            TemplatePlaceholder.applicable_template_category == "陈述式",
                            TemplatePlaceholder.type == "text"
                        )
                    )
                    placeholder = result.scalar_one_or_none()
                    
                    # 如果没有找到 type="text" 的，检查是否存在其他类型的（避免唯一性约束冲突）
                    if not placeholder:
                        result = await db.execute(
                            select(TemplatePlaceholder).where(
                                TemplatePlaceholder.name == placeholder_name,
                                TemplatePlaceholder.applicable_template_category == "陈述式"
                            )
                        )
                        existing_placeholder = result.scalar_one_or_none()
                        if existing_placeholder:
                            # 如果已存在其他类型的占位符，复用它（不强制改为 text）
                            placeholder = existing_placeholder
                            logger.info(f"复用已存在的占位符: {placeholder_name} (type={placeholder.type}, 适用于陈述式)")
                    
                    if not placeholder:
                        # 创建新的 type="text" 占位符，标记为适用于"陈述式"
                        placeholder = TemplatePlaceholder(
                            name=placeholder_name,
                            type="text",  # 陈述式统一使用 text
                            options=None,  # 无选项
                            applicable_template_category="陈述式",  # 标记为适用于陈述式模板
                            created_by_id=updated_by_id,
                            updated_by_id=updated_by_id,
                        )
                        db.add(placeholder)
                        await db.flush()  # flush 以获取 placeholder.id
                        logger.info(f"创建占位符: {placeholder_name} (type=text, 适用于陈述式)")
                    else:
                        logger.info(f"使用已存在的占位符: {placeholder_name} (type={placeholder.type}, 适用于陈述式)")
                else:
                    # 要素式：查询适用于"要素式"的占位符（任意类型）
                    result = await db.execute(
                        select(TemplatePlaceholder).where(
                            TemplatePlaceholder.name == placeholder_name,
                            TemplatePlaceholder.applicable_template_category == "要素式"
                        )
                    )
                    placeholder = result.scalar_one_or_none()
                    
                    if not placeholder:
                        # 创建新占位符（默认 type="text"），标记为适用于"要素式"
                        placeholder = TemplatePlaceholder(
                            name=placeholder_name,
                            type="text",  # 默认类型
                            options=None,  # 默认值
                            applicable_template_category="要素式",  # 标记为适用于要素式模板
                            created_by_id=updated_by_id,
                            updated_by_id=updated_by_id,
                        )
                        db.add(placeholder)
                        await db.flush()  # flush 以获取 placeholder.id
                        logger.info(f"创建占位符: {placeholder_name} (type=text, 适用于要素式)")
                    else:
                        logger.info(f"使用已存在的占位符: {placeholder_name} (type={placeholder.type}, 适用于要素式)")
                
                placeholder_objects.append(placeholder)
                placeholder_name_to_id[placeholder_name] = placeholder.id
            
            # 计算需要添加和删除的关联
            new_placeholder_ids = {p.id for p in placeholder_objects}
            ids_to_add = new_placeholder_ids - existing_placeholder_ids
            # 注意：只删除文档中不再存在的占位符关联，保留手动关联的占位符
            # 这里我们只删除那些在文档中不再出现的占位符
            ids_to_remove = existing_placeholder_ids - new_placeholder_ids
            
            # 删除文档中不再存在的占位符关联
            if ids_to_remove:
                await db.execute(
                    template_placeholder_association.delete().where(
                        (template_placeholder_association.c.template_id == template.id) &
                        (template_placeholder_association.c.placeholder_id.in_(ids_to_remove))
                    )
                )
                logger.info(f"删除 {len(ids_to_remove)} 个不再使用的占位符关联")
            
            # 添加新的占位符关联（如果不存在）
            if ids_to_add:
                for placeholder in placeholder_objects:
                    if placeholder.id in ids_to_add:
                        await db.execute(
                            template_placeholder_association.insert().values(
                                template_id=template.id,
                                placeholder_id=placeholder.id
                            )
                        )
                logger.info(f"添加 {len(ids_to_add)} 个新的占位符关联")
        
        await db.commit()
        await db.refresh(template)
        
        logger.info(f"更新模板成功: {template.id}")
        
        return template
    
    async def delete_template(self, db: AsyncSession, template_id: int) -> bool:
        """删除文书模板"""
        template = await self.get_template(db, template_id)
        if template is None:
            return False
        
        await db.delete(template)
        await db.commit()
        
        logger.info(f"删除模板成功: {template_id}")
        
        return True


class PlaceholderService:
    """占位符服务"""
    
    async def create_placeholder(
        self,
        db: AsyncSession,
        name: str,
        type: str,
        options: Optional[List[Dict[str, Any]]] = None,
        applicable_template_category: Optional[str] = None,
        created_by_id: Optional[int] = None,
    ) -> TemplatePlaceholder:
        """
        创建占位符（以 name + applicable_template_category 为唯一标识）
        
        如果占位符已存在（相同的 name 和 applicable_template_category），抛出 ValueError 异常
        """
        from sqlalchemy import select
        
        result = await db.execute(
            select(TemplatePlaceholder).where(
                TemplatePlaceholder.name == name,
                TemplatePlaceholder.applicable_template_category == applicable_template_category
            )
        )
        placeholder = result.scalar_one_or_none()
        
        if placeholder:
            category_str = applicable_template_category or "通用"
            raise ValueError(f"占位符 '{name}' (适用于{category_str}) 已存在")
        
        placeholder = TemplatePlaceholder(
            name=name,
            type=type,
            options=options,
            applicable_template_category=applicable_template_category,
            created_by_id=created_by_id,
            updated_by_id=created_by_id,
        )
        db.add(placeholder)
        await db.commit()
        await db.refresh(placeholder)
        logger.info(f"创建占位符成功: {name}")
        
        return placeholder
    
    async def get_placeholder(
        self,
        db: AsyncSession,
        placeholder_name: str,
    ) -> Optional[TemplatePlaceholder]:
        """
        根据占位符名称获取占位符
        
        Args:
            db: 数据库会话
            placeholder_name: 占位符名称
            
        Returns:
            Optional[TemplatePlaceholder]: 占位符对象，不存在返回 None
        """
        from sqlalchemy import select
        
        result = await db.execute(
            select(TemplatePlaceholder).where(
                TemplatePlaceholder.name == placeholder_name
            )
        )
        return result.scalar_one_or_none()
    
    async def list_placeholders(
        self,
        db: AsyncSession,
        template_id: Optional[int] = None,
        template_category: Optional[str] = None,
        skip: int = 0,
        limit: int = 100,
    ) -> tuple[List[TemplatePlaceholder], int]:
        """
        列出占位符
        
        Args:
            db: 数据库会话
            template_id: 模板ID（可选，如果提供则只返回该模板关联的占位符）
            template_category: 模板类型（可选，如果提供则只返回该类型或通用的占位符）
            skip: 跳过记录数
            limit: 返回记录数限制
            
        Returns:
            tuple[List[TemplatePlaceholder], int]: (占位符列表, 总数)
        """
        from sqlalchemy import select, func, or_
        from .models import template_placeholder_association, DocumentTemplate
        
        if template_id:
            # 查询指定模板关联的占位符
            # 如果提供了 template_category，先获取模板的 category（如果模板存在）
            if template_category is None:
                template_result = await db.execute(
                    select(DocumentTemplate.category).where(DocumentTemplate.id == template_id)
                )
                template_category = template_result.scalar_one_or_none()
            
            query = select(TemplatePlaceholder).join(
                template_placeholder_association
            ).where(
                template_placeholder_association.c.template_id == template_id
            )
            count_query = select(func.count()).select_from(
                TemplatePlaceholder
            ).join(
                template_placeholder_association
            ).where(
                template_placeholder_association.c.template_id == template_id
            )
        else:
            # 查询所有占位符
            query = select(TemplatePlaceholder)
            count_query = select(func.count(TemplatePlaceholder.id))
        
        # 如果指定了模板类型，过滤占位符（只返回该类型或通用的占位符）
        if template_category:
            # 判断是要素式还是陈述式
            is_narrative_style = "陈述" in template_category or template_category == "陈述式"
            category_filter = "陈述式" if is_narrative_style else ("要素式" if ("要素" in template_category or template_category == "要素式") else None)
            
            if category_filter:
                # 只返回匹配的 applicable_template_category 或 None（通用）的占位符
                category_condition = or_(
                    TemplatePlaceholder.applicable_template_category == category_filter,
                    TemplatePlaceholder.applicable_template_category.is_(None)
                )
                query = query.where(category_condition)
                count_query = count_query.where(category_condition)
        
        # 获取总数
        total_result = await db.execute(count_query)
        total = total_result.scalar() or 0
        
        # 获取列表
        query = query.order_by(TemplatePlaceholder.created_at.desc()).offset(skip).limit(limit)
        result = await db.execute(query)
        placeholders = list(result.scalars().unique().all())
        
        return placeholders, total
    
    async def update_placeholder(
        self,
        db: AsyncSession,
        placeholder_name: str,
        new_name: Optional[str] = None,
        type: Optional[str] = None,
        options: Optional[List[Dict[str, Any]]] = None,
        applicable_template_category: Optional[str] = None,
        updated_by_id: Optional[int] = None,
    ) -> Optional[TemplatePlaceholder]:
        """
        更新占位符
        
        注意：由于占位符的唯一性约束是 (name, applicable_template_category) 组合，更新时需要检查冲突
        """
        from sqlalchemy import select
        
        # 先根据 name 查询（可能有多个，取第一个）
        result = await db.execute(
            select(TemplatePlaceholder).where(
                TemplatePlaceholder.name == placeholder_name
            )
        )
        placeholder = result.scalar_one_or_none()
        if not placeholder:
            return None
        
        # 确定最终的名称和适用的模板类型
        final_name = new_name if new_name else placeholder_name
        
        # 检查 applicable_template_category 是否应该更新
        # 由于前端总是会发送这个字段（即使是 null），我们可以通过比较来判断是否需要更新
        # 如果传入的值与当前值不同，则需要更新
        should_update_category = False
        if applicable_template_category != placeholder.applicable_template_category:
            # 值发生了变化，需要更新
            should_update_category = True
            final_category = applicable_template_category
        else:
            # 值没有变化，不更新
            final_category = placeholder.applicable_template_category
        
        # 如果名称或适用的模板类型发生变化，检查 (final_name, final_category) 组合是否已存在
        if (new_name and new_name != placeholder_name) or should_update_category:
            existing_result = await db.execute(
                select(TemplatePlaceholder).where(
                    TemplatePlaceholder.name == final_name,
                    TemplatePlaceholder.applicable_template_category == final_category
                )
            )
            existing = existing_result.scalar_one_or_none()
            # 如果找到的占位符不是当前要更新的占位符，说明冲突
            if existing and existing.id != placeholder.id:
                category_str = final_category or "通用"
                raise ValueError(f"占位符 '{final_name}' (适用于{category_str}) 已存在")
        
        # 更新字段
        if new_name and new_name != placeholder_name:
            placeholder.name = new_name
        if type is not None:
            placeholder.type = type
        if should_update_category:
            placeholder.applicable_template_category = final_category
        if options is not None:
            placeholder.options = options
        if updated_by_id is not None:
            placeholder.updated_by_id = updated_by_id
        
        await db.commit()
        await db.refresh(placeholder)
        
        logger.info(f"更新占位符成功: {placeholder_name}" + (f" -> {new_name}" if new_name else ""))
        return placeholder
    
    async def delete_placeholder(
        self,
        db: AsyncSession,
        placeholder_name: str,
    ) -> bool:
        """
        删除占位符
        
        注意：由于唯一性约束改为 (name, applicable_template_category)，
        同一个 name 可能有多个占位符。此方法会删除所有匹配 name 的占位符。
        
        Args:
            db: 数据库会话
            placeholder_name: 占位符名称
            
        Returns:
            bool: 是否删除成功
        """
        from sqlalchemy import select
        from .models import template_placeholder_association
        
        # 查询所有匹配 name 的占位符（可能有多个，因为唯一性约束是 (name, applicable_template_category)）
        result = await db.execute(
            select(TemplatePlaceholder).where(
                TemplatePlaceholder.name == placeholder_name
            )
        )
        placeholders = result.scalars().all()
        
        if not placeholders:
            return False
        
        # 先删除所有关联记录（中间表）
        placeholder_ids = [p.id for p in placeholders]
        if placeholder_ids:
            await db.execute(
                template_placeholder_association.delete().where(
                    template_placeholder_association.c.placeholder_id.in_(placeholder_ids)
                )
            )
            logger.info(f"删除 {len(placeholder_ids)} 个占位符的关联记录")
        
        # 然后删除占位符本身
        for placeholder in placeholders:
            await db.delete(placeholder)
        
        await db.commit()
        
        logger.info(f"删除占位符成功: {placeholder_name} (共 {len(placeholders)} 个)")
        return True
    
    async def associate_placeholder_to_template(
        self,
        db: AsyncSession,
        template_id: int,
        placeholder_name: str,
    ) -> bool:
        """
        将占位符关联到模板
        
        注意：由于唯一性约束改为 (name, applicable_template_category)，
        同一个 name 可能有多个占位符。此方法会根据模板类型选择合适的占位符。
        优先级：匹配模板类型的占位符 > 通用占位符（applicable_template_category 为 None）
        
        Args:
            db: 数据库会话
            template_id: 模板ID
            placeholder_name: 占位符名称
            
        Returns:
            bool: 是否关联成功
        """
        from sqlalchemy import select, or_
        from .models import template_placeholder_association
        
        # 获取模板信息（包括 category）
        template_result = await db.execute(
            select(DocumentTemplate).where(DocumentTemplate.id == template_id)
        )
        template = template_result.scalar_one_or_none()
        if not template:
            return False
        
        # 根据模板类型选择合适的占位符
        # 优先级：1. 匹配模板类型的占位符 2. 通用占位符（applicable_template_category 为 None）
        template_category = template.category or ""
        
        # 查询匹配模板类型的占位符
        result = await db.execute(
            select(TemplatePlaceholder).where(
                TemplatePlaceholder.name == placeholder_name,
                TemplatePlaceholder.applicable_template_category == template_category
            )
        )
        placeholder = result.scalar_one_or_none()
        
        # 如果没有找到匹配模板类型的，查找通用占位符
        if not placeholder:
            result = await db.execute(
                select(TemplatePlaceholder).where(
                    TemplatePlaceholder.name == placeholder_name,
                    TemplatePlaceholder.applicable_template_category.is_(None)
                )
            )
            placeholder = result.scalar_one_or_none()
        
        if not placeholder:
            return False
        
        # 检查是否已存在关联
        existing = await db.execute(
            select(template_placeholder_association.c.template_id).where(
                (template_placeholder_association.c.template_id == template_id)
                & (template_placeholder_association.c.placeholder_id == placeholder.id)
            )
        )
        if existing.scalar_one_or_none():
            return True
        
        # 直接操作中间表插入关联，避免懒加载
        await db.execute(
            template_placeholder_association.insert().values(
                template_id=template_id,
                placeholder_id=placeholder.id,
            )
        )
        await db.commit()
        
        logger.info(f"关联占位符到模板成功: template_id={template_id}, placeholder_name={placeholder_name}")
        return True
    
    async def disassociate_placeholder_from_template(
        self,
        db: AsyncSession,
        template_id: int,
        placeholder_name: str,
    ) -> bool:
        """
        从模板中移除占位符关联
        
        注意：由于唯一性约束改为 (name, applicable_template_category)，
        同一个 name 可能有多个占位符。此方法会删除所有匹配 name 的占位符与模板的关联。
        
        Args:
            db: 数据库会话
            template_id: 模板ID
            placeholder_name: 占位符名称
            
        Returns:
            bool: 是否移除成功
        """
        from sqlalchemy import select
        from .models import template_placeholder_association
        
        # 校验模板存在
        template_result = await db.execute(
            select(DocumentTemplate.id).where(DocumentTemplate.id == template_id)
        )
        if not template_result.scalar_one_or_none():
            return False
        
        # 查找所有匹配 name 的占位符（可能有多个，因为唯一性约束是 (name, applicable_template_category)）
        placeholders_result = await db.execute(
            select(TemplatePlaceholder.id).where(
                TemplatePlaceholder.name == placeholder_name
            )
        )
        placeholder_ids = [row[0] for row in placeholders_result.all()]
        
        if not placeholder_ids:
            return False
        
        # 删除所有匹配的关联记录（通过中间表）
        delete_result = await db.execute(
            template_placeholder_association.delete().where(
                (template_placeholder_association.c.template_id == template_id)
                & (template_placeholder_association.c.placeholder_id.in_(placeholder_ids))
            )
        )
        
        if delete_result.rowcount > 0:
            await db.commit()
            logger.info(f"移除占位符关联成功: template_id={template_id}, placeholder_name={placeholder_name} (删除了 {delete_result.rowcount} 条关联)")
            return True
        
        return False


# 创建服务实例
template_editor_service = TemplateEditorService()
template_service = TemplateService()
placeholder_service = PlaceholderService()

