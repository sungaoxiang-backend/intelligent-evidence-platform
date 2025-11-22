"""
映射器
实现 docx 和 ProseMirror JSON 之间的双向转换
"""

import io
from loguru import logger
from typing import Dict, Any, List, Optional
import re
from docx import Document
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.table import Table, _Cell
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from .mapping_config import (
    ALIGNMENT_MAPPING,
    ALIGNMENT_REVERSE_MAPPING,
    VERTICAL_ALIGNMENT_MAPPING,
    VERTICAL_ALIGNMENT_REVERSE_MAPPING,
    rgb_to_hex,
    hex_to_rgb,
)

W_NAMESPACE = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W_NS = f"{{{W_NAMESPACE}}}"
W_NS_MAP = {"w": W_NAMESPACE}

HIGHLIGHT_COLOR_MAP = {
    "yellow": "#ffff00",
    "pink": "#ff00ff",
    "red": "#ff0000",
    "blue": "#0000ff",
    "darkblue": "#00008b",
    "darkred": "#8b0000",
    "darkyellow": "#b8860b",
    "darkgreen": "#006400",
    "darkcyan": "#008b8b",
    "darkmagenta": "#8b008b",
    "green": "#00ff00",
    "cyan": "#00ffff",
    "magenta": "#ff00ff",
    "black": "#000000",
    "white": "#ffffff",
    "lightgray": "#d3d3d3",
    "darkgray": "#a9a9a9",
}


class DocxToProseMirrorMapper:
    """docx → ProseMirror JSON 映射器"""

    def map_document(self, doc: Document) -> Dict[str, Any]:
        """
        映射整个文档到 ProseMirror JSON 格式

        Args:
            doc: python-docx Document 对象

        Returns:
            ProseMirror JSON 格式的文档
        """
        return {
            "type": "doc",
            "content": self._map_body_elements(doc),
        }

    def _map_body_elements(self, doc: Document) -> List[Dict[str, Any]]:
        """映射文档主体元素"""
        elements = []

        # 遍历文档的所有段落和表格
        # python-docx 的文档结构：paragraphs 和 tables 是分开的
        # 我们需要按照它们在文档中的顺序合并

        # 获取所有元素（段落和表格）的 XML 元素
        body = doc.element.body
        para_idx = 0
        table_idx = 0

        for element in body:
            if element.tag.endswith("p"):  # 段落
                if para_idx < len(doc.paragraphs):
                    para = doc.paragraphs[para_idx]
                    elements.append(self.map_paragraph(para))
                    para_idx += 1
            elif element.tag.endswith("tbl"):  # 表格
                if table_idx < len(doc.tables):
                    table = doc.tables[table_idx]
                    elements.append(self.map_table(table))
                    table_idx += 1

        return elements

    def map_paragraph(self, para: Paragraph) -> Dict[str, Any]:
        """
        映射段落

        Args:
            para: Paragraph 对象

        Returns:
            ProseMirror 段落或标题节点
        """
        attrs: Dict[str, Any] = {}

        # 对齐方式
        alignment = ALIGNMENT_MAPPING.get(para.alignment, "left")
        if alignment != "left":  # 默认左对齐，可以不设置
            attrs["textAlign"] = alignment

        # 缩进 - 分别处理左缩进和首行缩进
        if para.paragraph_format.left_indent:
            attrs["indent"] = para.paragraph_format.left_indent.pt
        if para.paragraph_format.first_line_indent:
            attrs["firstLineIndent"] = para.paragraph_format.first_line_indent.pt

        # 间距（可选）
        spacing = {}
        if para.paragraph_format.space_before:
            spacing["before"] = para.paragraph_format.space_before.pt
        if para.paragraph_format.space_after:
            spacing["after"] = para.paragraph_format.space_after.pt
        if spacing:
            attrs["spacing"] = spacing

        line_height = self._get_line_height(para)
        if line_height is not None:
            attrs["lineHeight"] = line_height

        # 文本内容
        content = self._map_runs(para.runs)

        # 检查是否是标题样式
        node_type = "paragraph"
        if para.style and para.style.name:
            style_name = para.style.name
            attrs["style"] = style_name
            
            # 将 Word 的标题样式映射到 ProseMirror 的 heading 节点
            if style_name.startswith("Heading") or style_name.startswith("标题"):
                # 提取标题级别
                level = self._extract_heading_level(style_name)
                if level:
                    node_type = "heading"
                    attrs["level"] = level
                    # 移除 style 属性，因为已经转换为 heading
                    attrs.pop("style", None)

        list_info = self._extract_list_info(para)
        if list_info:
            attrs["list"] = list_info

        # 如果段落为空，返回空段落（不包含文本节点）
        # ProseMirror 允许空段落，但不允许空文本节点
        return {
            "type": node_type,
            "attrs": attrs if attrs else {},
            "content": content,  # 可能为空数组，这是允许的
        }

    def _extract_heading_level(self, style_name: str) -> Optional[int]:
        """从样式名称中提取标题级别"""
        # 处理 "Heading 1", "Heading 2" 等
        if "Heading" in style_name or "标题" in style_name:
            # 尝试提取数字
            match = re.search(r'(\d+)', style_name)
            if match:
                level = int(match.group(1))
                # 限制在 1-6 之间
                return min(max(level, 1), 6)
        return None

    def _get_line_height(self, para: Paragraph) -> Optional[float]:
        """提取行距并统一为 float"""
        line_spacing = para.paragraph_format.line_spacing
        if not line_spacing:
            return None

        try:
            if hasattr(line_spacing, "pt"):
                return float(line_spacing.pt)
            if isinstance(line_spacing, (int, float)):
                return float(line_spacing)
        except (TypeError, ValueError):
            logger.debug("无法解析行距，使用默认值")
            return None

        return None

    def _extract_list_info(self, para: Paragraph) -> Optional[Dict[str, Any]]:
        """提取段落的列表信息（若存在）"""
        try:
            p_pr = para._p.pPr  # type: ignore[attr-defined]
            num_pr = getattr(p_pr, "numPr", None) if p_pr is not None else None

            style = getattr(para, "style", None)
            num_id = getattr(num_pr.numId, "val", None) if num_pr is not None else None
            ilvl = (
                getattr(num_pr.ilvl, "val", "0") if num_pr is not None else "0"
            )

            if num_id is None and style is not None:
                style_num_ids = style.element.xpath("./w:pPr/w:numPr/w:numId/@w:val")
                if style_num_ids:
                    num_id = style_num_ids[0]
                style_ilvls = style.element.xpath("./w:pPr/w:numPr/w:ilvl/@w:val")
                if style_ilvls:
                    ilvl = style_ilvls[0]

            if num_id is None:
                return None

            try:
                num_id_int = int(num_id)
            except (TypeError, ValueError):
                num_id_int = 0

            try:
                level = int(ilvl)
            except (TypeError, ValueError):
                level = 0

            list_info: Dict[str, Any] = {
                "numId": num_id_int,
                "level": level,
            }

            numbering_part = getattr(para.part, "numbering_part", None)
            list_type = "ordered"
            fmt_value = None
            marker_value = None

            if numbering_part is not None:
                numbering = numbering_part.element
                abstract_id = None
                for num in numbering.findall(f".//{W_NS}num"):
                    if num.get(qn("w:numId")) == str(num_id):
                        abstract_num = num.find(f"./{W_NS}abstractNumId")
                        if abstract_num is not None:
                            abstract_id = abstract_num.get(qn("w:val"))
                        break

                if abstract_id is not None:
                    target_level = str(ilvl or "0")
                    lvl_node = None
                    for abstract in numbering.findall(f".//{W_NS}abstractNum"):
                        if abstract.get(qn("w:abstractNumId")) == abstract_id:
                            for lvl in abstract.findall(f"./{W_NS}lvl"):
                                if lvl.get(qn("w:ilvl")) == target_level:
                                    lvl_node = lvl
                                    break
                            if lvl_node is None:
                                lvl_node = abstract.find(f"./{W_NS}lvl")
                            break

                    if lvl_node is not None:
                        num_fmt = lvl_node.find(f"./{W_NS}numFmt")
                        if num_fmt is not None:
                            fmt_value = num_fmt.get(qn("w:val"))
                        lvl_text = lvl_node.find(f"./{W_NS}lvlText")
                        if lvl_text is not None:
                            marker_value = lvl_text.get(qn("w:val"))

            if fmt_value:
                list_info["format"] = fmt_value
                if fmt_value == "bullet":
                    list_type = "unordered"
            list_info["type"] = list_type

            if marker_value:
                list_info["marker"] = marker_value

            return list_info
        except Exception as exc:
            logger.warning(f"解析列表信息失败: {exc}")
            return {"type": "ordered", "__fallback": True}

    def _get_highlight_color(self, run: Run) -> Optional[str]:
        """提取高亮颜色"""
        try:
            highlight = run.font.highlight_color
            if not highlight or highlight == WD_COLOR_INDEX.AUTO:
                return None

            name = getattr(highlight, "name", None)
            if not name:
                name = str(highlight).split()[0]
            if not name:
                return None
            return HIGHLIGHT_COLOR_MAP.get(name.lower())
        except Exception:
            return None

    def _map_runs(self, runs: List[Run]) -> List[Dict[str, Any]]:
        """映射文本运行"""
        content = []
        for run in runs:
            if run.text is None:
                continue

            text_value = run.text
            if text_value == "":
                continue

            parts = text_value.split("\n")
            for idx, part in enumerate(parts):
                if part:
                    text_node: Dict[str, Any] = {
                        "type": "text",
                        "text": part,
                    }

                    marks = self._map_marks(run)
                    if marks:
                        text_node["marks"] = marks

                    content.append(text_node)

                if idx < len(parts) - 1:
                    content.append({"type": "hardBreak"})

        return content

    def _map_marks(self, run: Run) -> List[Dict[str, Any]]:
        """映射文本标记（粗体、斜体、颜色等）"""
        marks = []

        if run.bold:
            marks.append({"type": "bold"})
        if run.italic:
            marks.append({"type": "italic"})
        if run.underline:
            marks.append({"type": "underline"})
        if run.font.strike:
            marks.append({"type": "strike"})

        # 收集样式属性到textStyle中
        style_attrs = {}

        # 字体颜色
        color_hex = rgb_to_hex(run.font.color) if run.font.color else None
        if color_hex:
            style_attrs["color"] = color_hex

        # 字体大小
        if run.font.size:
            font_size = f"{run.font.size.pt}pt"
            style_attrs["fontSize"] = font_size
        else:
            # 如果没有直接设置字体大小，尝试从样式中获取
            try:
                if run.style and hasattr(run.style, 'font') and run.style.font.size:
                    font_size = f"{run.style.font.size.pt}pt"
                    style_attrs["fontSize"] = font_size
            except (AttributeError, TypeError):
                pass

        # 字体名称
        try:
            if run.font.name:
                style_attrs["fontFamily"] = run.font.name
            elif run.style and hasattr(run.style, 'font') and run.style.font.name:
                style_attrs["fontFamily"] = run.style.font.name
        except (AttributeError, TypeError):
            pass

        highlight_color = self._get_highlight_color(run)
        if highlight_color:
            style_attrs["backgroundColor"] = highlight_color

        if getattr(run.font, "small_caps", False):
            style_attrs["smallCaps"] = True
        if getattr(run.font, "all_caps", False):
            style_attrs["allCaps"] = True

        # 如果有样式属性，添加到marks中
        if style_attrs:
            marks.append({
                "type": "textStyle",
                "attrs": style_attrs,
            })

        return marks

    def map_table(self, table: Table) -> Dict[str, Any]:
        """
        映射表格

        Args:
            table: Table 对象

        Returns:
            ProseMirror 表格节点
        """
        rows = []
        
        # 直接从 XML 读取 tc 元素，避免 python-docx 展开合并单元格的问题
        for row in table.rows:
            cells = []
            # 获取行的 XML 元素
            tr = row._tr
            # 查找所有 tc 元素（使用命名空间）
            ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
            tc_elements = tr.findall(f'.//{ns}tc')
            
            # 使用一个集合来跟踪已处理的单元格，避免重复处理
            processed_cells = set()
            
            for tc in tc_elements:
                # 检查这个 tc 是否已经被处理过（通过 id 或位置）
                tc_id = id(tc)
                if tc_id in processed_cells:
                    continue
                processed_cells.add(tc_id)
                
                # 从 tc 元素创建 _Cell 对象
                # 注意：我们需要找到对应的 cell 对象，或者直接从 XML 解析
                # 由于 python-docx 的限制，我们需要通过 row.cells 来找到对应的 cell
                # 但我们需要检查 tc 是否对应某个 cell
                cell = None
                for c in row.cells:
                    if c._tc is tc:
                        cell = c
                        break
                
                if cell is None:
                    # 如果找不到对应的 cell，可能是被合并的单元格，跳过
                    continue
                
                # 映射单元格（包含合并信息）
                cell_node = self.map_table_cell(cell)
                
                # 如果单元格为 None（被合并的单元格），跳过
                if cell_node is None:
                    continue
                
                cells.append(cell_node)
            
            rows.append({"type": "tableRow", "content": cells})

        attrs: Dict[str, Any] = {"border": True}
        if table.style and table.style.name:
            attrs["style"] = table.style.name
        
        # 提取表格列宽信息（从 tblGrid）
        # DOCX 中的列宽单位是 twips (dxa)，1 twip = 1/20 point = 1/1440 inch
        try:
            tbl = table._tbl
            tbl_pr = tbl.tblPr
            tbl_grid = tbl.tblGrid
            logger.info(f"[PARSE] 开始提取列宽，tbl_grid 是否存在: {tbl_grid is not None}")
            if tbl_grid is not None:
                col_widths = []
                grid_cols = tbl_grid.findall(f'.//{W_NS}gridCol')
                logger.info(f"[PARSE] 找到 {len(grid_cols)} 个 gridCol 元素")
                for i, grid_col in enumerate(grid_cols):
                    width = grid_col.get(qn('w:w'))
                    width_type = grid_col.get(qn('w:type'))
                    logger.info(f"[PARSE] gridCol[{i}]: width={width}, type={width_type}")
                    if width:
                        try:
                            # 列宽单位是 twips (dxa)，直接保存
                            col_widths.append(int(width))
                            logger.info(f"[PARSE] ✅ 成功提取 gridCol[{i}] 宽度: {width}")
                        except (ValueError, TypeError) as e:
                            logger.warning(f"[PARSE] ❌ 无法转换列宽 {width}: {e}")
                if col_widths:
                    attrs["colWidths"] = col_widths
                    logger.info(f"[PARSE] ✅✅✅ 最终提取到列宽并保存到 attrs: {col_widths}")
                else:
                    logger.warning("[PARSE] ⚠️ 未提取到任何列宽")
            else:
                logger.warning("[PARSE] ⚠️ tblGrid 不存在，无法提取列宽")

            if tbl_pr is not None:
                tbl_w = tbl_pr.find(f"./{W_NS}tblW")
                if tbl_w is not None:
                    width_val = tbl_w.get(qn("w:w"))
                    width_type = tbl_w.get(qn("w:type"))
                    if width_val:
                        attrs["tableWidth"] = {
                            "width": int(width_val),
                            "type": width_type or "auto",
                        }
                tbl_layout = tbl_pr.find(f"./{W_NS}tblLayout")
                if tbl_layout is not None and tbl_layout.get(qn("w:type")):
                    attrs["tableLayout"] = tbl_layout.get(qn("w:type"))
        except (AttributeError, Exception) as e:
            logger.error(f"[PARSE] ❌ 提取列宽时出错: {str(e)}", exc_info=True)

        return {
            "type": "table",
            "attrs": attrs,
            "content": rows,
        }

    def map_table_cell(self, cell: _Cell, row_idx: int = -1, col_idx: int = -1) -> Optional[Dict[str, Any]]:
        """
        映射表格单元格

        Args:
            cell: _Cell 对象

        Returns:
            ProseMirror 表格单元格节点
        """
        attrs: Dict[str, Any] = {
            "colspan": 1,
            "rowspan": 1,
        }

        # 检测合并单元格（colspan 和 rowspan）
        try:
            tc = cell._tc  # table cell element
            tc_pr = tc.tcPr  # table cell properties
            
            if tc_pr is not None:
                # 检测横向合并（colspan）
                grid_span = tc_pr.gridSpan
                if grid_span is not None:
                    # gridSpan 的值存储在 XML 属性中，需要通过 qn 获取
                    colspan_val = grid_span.get(qn('w:val'))
                    if colspan_val is not None:
                        try:
                            colspan_int = int(colspan_val)
                            if colspan_int > 1:
                                attrs["colspan"] = colspan_int
                                logger.debug(f"检测到 colspan: {colspan_int}")
                        except (ValueError, TypeError):
                            pass
                
                # 检测纵向合并（rowspan）
                v_merge = tc_pr.vMerge
                if v_merge is not None:
                    # vMerge 存在说明这个单元格参与了纵向合并
                    # 获取 vMerge 的 val 属性
                    v_merge_val = v_merge.get(qn('w:val'))
                    
                    # 如果 vMerge.val 不存在或为 "restart"，说明这是合并的起始单元格
                    if v_merge_val is None or v_merge_val == "restart":
                        # 这是合并的起始单元格，需要计算 rowspan
                        rowspan = self._calculate_rowspan(cell, tc)
                        if rowspan > 1:
                            attrs["rowspan"] = rowspan
                            logger.debug(f"检测到 rowspan: {rowspan}")
                    # 注意：被合并的单元格（vMerge.val 不存在但不是起始单元格）不会出现在 row.cells 中
        except (AttributeError, Exception) as e:
            logger.debug(f"无法检测合并单元格: {str(e)}")
            # 如果无法检测，使用默认值

        # 背景色（暂时跳过，python-docx 的 _Cell 对象没有直接的 shading 属性）
        fill = None
        if tc_pr is not None:
            shading = tc_pr.find(f"./{W_NS}shd")
            if shading is not None:
                fill = shading.get(qn("w:fill"))

        if fill and fill.lower() != "auto":
            fill_str = fill.lstrip("#").upper()
            if len(fill_str) == 6:
                attrs["backgroundColor"] = f"#{fill_str}"

        # 提取单元格宽度（tcW）- 完全基于原始数据，不做任何计算
        if tc_pr is not None:
            tc_w = tc_pr.tcW
            if tc_w is not None:
                width = tc_w.get(qn('w:w'))
                width_type = tc_w.get(qn('w:type'))
                if width:
                    try:
                        attrs["cellWidth"] = {
                            "width": int(width),
                            "type": width_type if width_type else "dxa"
                        }
                        logger.debug(f"[PARSE] 提取单元格宽度: {width} {width_type}")
                    except (ValueError, TypeError):
                        pass

        # 垂直对齐
        if cell.vertical_alignment:
            vertical_align = VERTICAL_ALIGNMENT_MAPPING.get(
                cell.vertical_alignment, "top"
            )
            if vertical_align != "top":  # 默认顶部对齐
                attrs["verticalAlign"] = vertical_align

        # 单元格内容（段落列表）
        content = []
        for para in cell.paragraphs:
            para_node = self.map_paragraph(para)
            content.append(para_node)

        # 如果单元格为空，至少添加一个空段落（不包含空文本节点）
        if not content:
            content = [
                {
                    "type": "paragraph",
                    "attrs": {},
                    "content": [],  # 空段落，不包含空文本节点
                }
            ]

        # 如果 rowspan 为 0，说明这是被合并的单元格，不应该输出
        if attrs.get("rowspan") == 0:
            return None

        return {
            "type": "tableCell",
            "attrs": attrs,
            "content": content,
        }

    def _calculate_rowspan(self, cell: _Cell, tc_element) -> int:
        """
        计算单元格的 rowspan（向下合并的行数）
        
        Args:
            cell: _Cell 对象
            tc_element: 单元格的 XML 元素
            
        Returns:
            rowspan 值
        """
        try:
            # 通过 table 对象找到单元格位置
            # 找到单元格所在的表格和行
            table = None
            row = None
            row_idx = None
            col_idx = None
            
            # 遍历所有表格（从文档中）
            # 但我们需要知道 cell 属于哪个 table
            # 通过 XML 元素向上查找表格
            parent = tc_element.getparent()
            while parent is not None:
                if parent.tag.endswith('tr'):
                    # 找到行元素，现在找表格
                    tbl_parent = parent.getparent()
                    while tbl_parent is not None:
                        if tbl_parent.tag.endswith('tbl'):
                            # 找到了表格元素
                            # 现在需要找到对应的 Table 对象
                            # 通过遍历文档的表格来匹配
                            break
                        tbl_parent = tbl_parent.getparent()
                    break
                parent = parent.getparent()
            
            # 简化方法：直接通过 XML 查找
            # 找到当前单元格所在的行
            row_elem = tc_element.getparent()
            while row_elem is not None and not row_elem.tag.endswith('tr'):
                row_elem = row_elem.getparent()
            
            if row_elem is None:
                return 1
            
            # 找到表格元素
            tbl_elem = row_elem.getparent()
            while tbl_elem is not None and not tbl_elem.tag.endswith('tbl'):
                tbl_elem = tbl_elem.getparent()
            
            if tbl_elem is None:
                return 1
            
            # 找到当前单元格在行中的位置
            cells_in_row = row_elem.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tc')
            for j, tc in enumerate(cells_in_row):
                if tc == tc_element:
                    col_idx = j
                    break
            
            if col_idx is None:
                return 1
            
            # 找到当前行在表格中的位置
            rows_in_table = tbl_elem.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tr')
            for i, r in enumerate(rows_in_table):
                if r == row_elem:
                    row_idx = i
                    break
            
            if row_idx is None:
                return 1
            
            # 向下查找被合并的单元格
            rowspan = 1
            for i in range(row_idx + 1, len(rows_in_table)):
                next_row = rows_in_table[i]
                next_cells = next_row.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tc')
                
                if col_idx >= len(next_cells):
                    break
                
                next_tc = next_cells[col_idx]
                next_tc_pr = next_tc.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcPr')
                
                if next_tc_pr is not None:
                    next_v_merge = next_tc_pr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}vMerge')
                    # 如果下一个单元格的 vMerge 没有 val 属性，说明它被合并了
                    if next_v_merge is not None:
                        # 检查命名空间
                        ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
                        val_attr = next_v_merge.get(ns + 'val')
                        if val_attr is None:
                            rowspan += 1
                        else:
                            break
                    else:
                        break
                else:
                    break
            
            return rowspan
        except (AttributeError, Exception) as e:
            logger.debug(f"计算 rowspan 失败: {str(e)}")
            return 1


class ProseMirrorToDocxMapper:
    """ProseMirror JSON → docx 映射器"""

    def __init__(self) -> None:
        self.export_warnings: List[Dict[str, Any]] = []

    def map_document(self, prosemirror_json: Dict[str, Any]) -> Document:
        """
        从 ProseMirror JSON 映射到 docx 文档

        Args:
            prosemirror_json: ProseMirror JSON 格式的文档

        Returns:
            Document 对象
        """
        self.export_warnings = []
        doc = Document()

        # 遍历内容节点
        content = prosemirror_json.get("content", [])
        for idx, node in enumerate(content):
            node_type = node.get("type")
            try:
                if node_type == "paragraph" or node_type == "heading":
                    self.map_paragraph_node(node, doc)
                elif node_type == "table":
                    self.map_table_node(node, doc)
                else:
                    self._log_warning(
                        "unsupported_node",
                        f"暂不支持的节点类型: {node_type}",
                        {"index": idx, "type": node_type},
                    )
            except Exception as exc:
                self._log_warning(
                    "node_export_failed",
                    f"节点导出失败: {node_type}",
                    {"index": idx, "type": node_type, "error": str(exc)},
                )

        return doc

    def _log_warning(self, code: str, message: str, meta: Optional[Dict[str, Any]] = None) -> None:
        entry = {
            "code": code,
            "message": message,
            "meta": meta or {},
        }
        self.export_warnings.append(entry)
        logger.warning(f"[EXPORT] {code}: {message} - {meta}")


    def map_paragraph_node(
        self, node: Dict[str, Any], container: Any
    ) -> Paragraph:
        """
        映射段落节点

        Args:
            node: ProseMirror 段落节点
            container: Document 或 _Cell 对象

        Returns:
            Paragraph 对象
        """
        # 根据容器类型添加段落
        # 检查是否是 Document 对象（通过检查是否有 add_paragraph 方法且不是 _Cell）
        # Document 对象有 element 属性，_Cell 对象有 _tc 属性
        if hasattr(container, 'element') and not hasattr(container, '_tc'):
            para = container.add_paragraph()
        else:  # _Cell
            para = container.add_paragraph()

        # 应用格式
        attrs = node.get("attrs", {})

        # 对齐方式
        if "textAlign" in attrs:
            alignment = ALIGNMENT_REVERSE_MAPPING.get(
                attrs["textAlign"], WD_ALIGN_PARAGRAPH.LEFT
            )
            para.alignment = alignment

        # 缩进 - 分别处理左缩进和首行缩进
        indent_value = attrs.get("indent")
        if indent_value is not None:
            try:
                para.paragraph_format.left_indent = Pt(float(indent_value))
            except (TypeError, ValueError):
                logger.warning(f"无法设置左缩进: {indent_value}")

        first_line_indent = attrs.get("firstLineIndent")
        if first_line_indent is not None:
            try:
                para.paragraph_format.first_line_indent = Pt(float(first_line_indent))
            except (TypeError, ValueError):
                logger.warning(f"无法设置首行缩进: {first_line_indent}")

        line_height = attrs.get("lineHeight")
        if line_height is not None:
            try:
                numeric_line_height = float(line_height)
                if numeric_line_height < 5:
                    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                    para.paragraph_format.line_spacing = numeric_line_height
                else:
                    para.paragraph_format.line_spacing = Pt(numeric_line_height)
            except (TypeError, ValueError):
                logger.warning(f"无法设置行距: {line_height}")

        # 样式（如果是 heading，映射为 Word 的标题样式）
        if node.get("type") == "heading":
            level = attrs.get("level", 1)
            try:
                para.style = f"Heading {level}"
            except Exception:
                # 如果 Heading 样式不存在，尝试使用中文样式名
                try:
                    para.style = f"标题 {level}"
                except Exception:
                    logger.warning(f"标题样式 'Heading {level}' 不存在，使用默认样式")
        elif "style" in attrs:
            style_name = attrs["style"]
            # 跳过不存在的样式（如 'Table Text'），使用默认样式
            if style_name and style_name not in ["Table Text"]:
                try:
                    para.style = style_name
                except Exception:
                    logger.warning(f"样式 '{style_name}' 不存在，使用默认样式")
            # 对于 'Table Text' 等特殊样式，直接使用默认样式，不报错

        list_attrs = attrs.get("list")
        if list_attrs:
            self._apply_list_format(para, list_attrs)

        # 间距 - 完全按照ProseMirror JSON中的值设置
        spacing = attrs.get("spacing")
        if isinstance(spacing, dict):
            before = spacing.get("before")
            after = spacing.get("after")
            if before is not None:
                try:
                    para.paragraph_format.space_before = Pt(float(before))
                except (TypeError, ValueError):
                    logger.warning(f"无法设置段前间距: {before}")
            if after is not None:
                try:
                    para.paragraph_format.space_after = Pt(float(after))
                except (TypeError, ValueError):
                    logger.warning(f"无法设置段后间距: {after}")

        # 添加文本内容
        content = node.get("content", [])
        for text_node in content:
            node_type = text_node.get("type")
            if node_type == "text":
                run = para.add_run(text_node.get("text", ""))
                self._apply_marks(run, text_node.get("marks", []))
                if not run.font.color or not run.font.color.rgb:
                    run.font.color.rgb = RGBColor(0, 0, 0)
            elif node_type == "placeholder":
                # 处理 placeholder 节点（要素式模板保留的占位符节点）
                attrs = text_node.get("attrs", {})
                value = attrs.get("value")
                placeholder_type = attrs.get("placeholderType", "text")
                options = attrs.get("options", [])
                
                # 格式化占位符值
                formatted_text = self._format_placeholder_value_for_export(value, placeholder_type, options)
                
                if formatted_text:
                    run = para.add_run(formatted_text)
                    if not run.font.color or not run.font.color.rgb:
                        run.font.color.rgb = RGBColor(0, 0, 0)
            elif node_type == "hardBreak":
                para.add_run().add_break()

        return para

    def _format_placeholder_value_for_export(
        self, 
        value: Any, 
        placeholder_type: str, 
        options: Optional[List[Dict[str, Any]]] = None
    ) -> str:
        """格式化占位符值用于导出到 DOCX"""
        if value is None:
            return ""
        
        # 对于 radio/checkbox 类型，显示所有选项和选中状态
        if placeholder_type in ["radio", "checkbox"] and options:
            # 处理 radio：值是单个字符串
            # 处理 checkbox：值是数组
            if placeholder_type == "radio":
                selected_values = [value] if value else []
            else:
                selected_values = value if isinstance(value, list) else ([value] if value else [])
            
            # 格式化选项：☑ 已选 ☐ 未选
            formatted_options = []
            for opt in options:
                opt_value = opt.get("value", "")
                opt_label = opt.get("label", opt_value)
                
                if opt_value in selected_values:
                    formatted_options.append(f"☑ {opt_label}")
                else:
                    formatted_options.append(f"☐ {opt_label}")
            
            return "  ".join(formatted_options)
        
        # 其他类型正常处理
        elif isinstance(value, list):
            # 数组转换为顿号分隔的字符串
            if len(value) == 0:
                return ""
            return "、".join(str(v) for v in value if v)
        elif isinstance(value, (int, float)):
            # 数字转换为字符串
            return str(value)
        elif isinstance(value, str):
            if value == "":
                return ""
            return value
        elif isinstance(value, bool):
            # 布尔值转换为字符串
            return "是" if value else "否"
        else:
            # 其他类型转换为字符串
            return str(value) if value else ""

    def _apply_list_format(self, para: Paragraph, list_info: Dict[str, Any]) -> None:
        list_type = list_info.get("type")
        level = int(list_info.get("level", 0))

        preferred_style = list_info.get("style")
        fallback_style = "List Number" if list_type == "ordered" else "List Bullet"
        style_to_apply = preferred_style or fallback_style

        if style_to_apply:
            try:
                para.style = style_to_apply
            except Exception:
                logger.warning(f"列表样式 '{style_to_apply}' 不存在，使用默认样式")

        try:
            indent_base = 18  # approx 0.25 inch
            para.paragraph_format.left_indent = Pt(indent_base * (level + 1))
            para.paragraph_format.first_line_indent = Pt(-indent_base / 2)
        except Exception:
            pass

    def map_table_node(self, node: Dict[str, Any], doc: Document) -> Table:
        """
        映射表格节点

        Args:
            node: ProseMirror 表格节点
            doc: Document 对象

        Returns:
            Table 对象
        """
        rows = node.get("content", [])
        if not rows:
            return None

        table_attrs = node.get("attrs", {})

        # 计算实际需要的列数（考虑 colspan）
        # 这是表格的最大列数，用于创建基础表格结构
        max_cols = 1
        for row_node in rows:
            cells = row_node.get("content", [])
            current_cols = 0
            for cell_node in cells:
                attrs = cell_node.get("attrs", {})
                colspan = attrs.get("colspan", 1)
                current_cols += colspan
            max_cols = max(max_cols, current_cols)

        # 创建表格：先创建基础结构（所有行都有 max_cols 列）
        # 然后按照实际结构填充，合并单元格通过 gridSpan 和 vMerge 实现
        table = doc.add_table(rows=len(rows), cols=max_cols)

        table_style_name = table_attrs.get("style")
        if table_style_name:
            try:
                table.style = table_style_name
            except Exception:
                logger.warning(f"[EXPORT] 表格样式 '{table_style_name}' 不存在，使用默认样式")

        ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        
        # 应用列宽（如果存在，从原始表格提取的）
        # 注意：必须在创建表格后立即设置，在填充内容之前
        logger.info(f"[EXPORT] 表格节点 attrs: {table_attrs}")
        col_widths = table_attrs.get("colWidths")
        logger.info(f"[EXPORT] 从 attrs 获取的 colWidths: {col_widths}, max_cols: {max_cols}")
        
        if col_widths and len(col_widths) == max_cols:
            logger.info(f"[EXPORT] ✅ 开始应用列宽: {col_widths}")
            # 设置表格的列定义（tblGrid）
            tbl = table._tbl
            tbl_grid = tbl.tblGrid
            logger.info(f"[EXPORT] tbl_grid 是否存在: {tbl_grid is not None}")
            if tbl_grid is not None:
                # 更新现有的 gridCol 元素
                grid_cols = tbl_grid.findall(f'.//{ns}gridCol')
                logger.info(f"[EXPORT] 找到 {len(grid_cols)} 个 gridCol 元素")
                for i, (grid_col, width) in enumerate(zip(grid_cols, col_widths)):
                    if i < len(grid_cols):
                        # 设置列宽（单位：twips/dxa）
                        old_width = grid_col.get(qn('w:w'))
                        grid_col.set(qn('w:w'), str(width))
                        grid_col.set(qn('w:type'), 'dxa')
                        new_width = grid_col.get(qn('w:w'))
                        logger.info(f"[EXPORT] ✅ 设置 gridCol[{i}]: {old_width} -> {new_width} (type=dxa)")
                # 验证设置结果
                grid_cols_verify = tbl_grid.findall(f'.//{ns}gridCol')
                logger.info(f"[EXPORT] 验证：重新读取 gridCol")
                for i, grid_col in enumerate(grid_cols_verify):
                    width = grid_col.get(qn('w:w'))
                    width_type = grid_col.get(qn('w:type'))
                    logger.info(f"[EXPORT] ✅✅✅ gridCol[{i}]: width={width}, type={width_type}")
            else:
                logger.error("[EXPORT] ❌ tblGrid 不存在，无法应用列宽")
        elif col_widths:
            logger.warning(f"[EXPORT] ⚠️ 列宽数量不匹配: colWidths={col_widths} (长度={len(col_widths)}), max_cols={max_cols}")
        else:
            logger.warning("[EXPORT] ⚠️ 未找到 colWidths 属性，使用默认列宽")
        
        tbl = table._tbl
        tbl_pr = tbl.tblPr
        if tbl_pr is None:
            tbl_pr = OxmlElement('w:tblPr')
            tbl.insert(0, tbl_pr)
        
        table_width_info = table_attrs.get("tableWidth")
        table_width_value = None
        table_width_type = "auto"
        if table_width_info:
            table_width_value = table_width_info.get("width")
            table_width_type = table_width_info.get("type", "auto")
        if table_width_value is None and col_widths:
            table_width_value = sum(col_widths)
            table_width_type = "dxa"

        if table_width_value is not None:
            existing_tbl_w = tbl_pr.find(f'.//{ns}tblW')
            if existing_tbl_w is None:
                existing_tbl_w = OxmlElement('w:tblW')
                tbl_pr.append(existing_tbl_w)
            existing_tbl_w.set(qn('w:w'), str(table_width_value))
            existing_tbl_w.set(qn('w:type'), table_width_type)
            logger.info(f"[EXPORT] 设置表格总宽度: {table_width_value} ({table_width_type})")

        layout_type = table_attrs.get("tableLayout")
        if not layout_type and col_widths:
            layout_type = "fixed"

        if layout_type:
            existing_tbl_layout = tbl_pr.find(f'.//{ns}tblLayout')
            if existing_tbl_layout is None:
                existing_tbl_layout = OxmlElement('w:tblLayout')
                tbl_pr.append(existing_tbl_layout)
            existing_tbl_layout.set(qn('w:type'), layout_type)
            logger.info(f"[EXPORT] 设置表格布局: {layout_type}")
        
        # 设置边框
        tbl_borders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tbl_borders.append(border)
        tbl_pr.append(tbl_borders)

        # 填充表格（考虑合并单元格）
        # 关键：按照实际结构填充，合并单元格通过 gridSpan 和 vMerge 实现
        # 对于横向合并（colspan），设置 gridSpan 后，被合并的单元格需要从 XML 中删除
        # 跟踪每行中哪些列位置已被占用（由于 rowspan）
        occupied_positions = {}  # {row_idx: {col_pos: True}}
        
        for row_idx, row_node in enumerate(rows):
            cells = row_node.get("content", [])
            col_pos = 0  # 当前行的逻辑列位置（考虑合并）
            row_xml = table.rows[row_idx]._tr  # 行的 XML 元素
            # 获取行的所有 tc 元素（用于删除被合并的单元格）
            # 注意：每次循环开始时重新获取，因为删除单元格后列表会变化
            tc_elements_in_row = row_xml.findall(f'.//{ns}tc')
            # 跟踪已处理的单元格索引（在 XML 中的实际位置）
            processed_tc_indices = set()
            # 记录需要删除的单元格（等所有单元格处理完后再删除）
            cells_to_remove_in_row = []
            logger.debug(f"[EXPORT] 处理第 {row_idx} 行，共 {len(cells)} 个单元格，XML 中有 {len(tc_elements_in_row)} 个 tc 元素")
            
            for cell_idx, cell_node in enumerate(cells):
                attrs = cell_node.get("attrs", {})
                colspan = attrs.get("colspan", 1)
                rowspan = attrs.get("rowspan", 1)
                
                logger.debug(f"[EXPORT] 处理单元格 {cell_idx}/{len(cells)}: colspan={colspan}, rowspan={rowspan}, 当前 col_pos={col_pos}")
                
                # 跳过已经被 rowspan 占用的列位置
                while col_pos in occupied_positions.get(row_idx, {}):
                    col_pos += 1
                    logger.debug(f"[EXPORT] 跳过被 rowspan 占用的列位置，col_pos 递增为 {col_pos}")
                
                # 找到对应的实际单元格（在 table.rows[row_idx].cells 中的位置）
                # 关键：逻辑位置（col_pos）对应的是 ProseMirror JSON 中的单元格顺序
                # 实际位置（actual_cell_index）对应的是 Word XML 中的 tc 元素索引
                # 对于合并单元格，逻辑位置会跳过被合并的列，但实际位置需要找到对应的 tc 元素
                actual_cell_index = None
                # 获取将被删除的单元格索引集合
                cells_to_remove_indices = {idx for idx, _ in cells_to_remove_in_row}
                
                logger.info(f"[EXPORT] 查找单元格: 行{row_idx}, 单元格{cell_idx}/{len(cells)}, 逻辑列{col_pos}, 总tc数={len(tc_elements_in_row)}, 已处理={processed_tc_indices}, 将删除={cells_to_remove_indices}")
                
                # 重新设计查找逻辑：由于前面行的合并单元格删除操作会影响整个表格的XML结构
                # 我们需要从Python的table.rows[row_idx].cells中获取正确的单元格
                # 而不是从可能不完整的XML中查找

                # 直接通过逻辑位置从Python对象获取单元格
                # Word表格的Python对象应该保持了正确的表格结构
                logger.debug(f"[EXPORT] 尝试从table.rows[{row_idx}].cells获取逻辑列{col_pos}的单元格")

                actual_cell_index = None
                tc = None
                cell = None

                # 方法1：直接从Python对象获取（推荐）
                try:
                    # 获取当前行的所有单元格
                    row_cells = table.rows[row_idx].cells
                    logger.debug(f"[EXPORT] 当前行Python对象中有 {len(row_cells)} 个单元格")

                    # 检查逻辑位置是否有效
                    if col_pos < len(row_cells):
                        cell = row_cells[col_pos]
                        tc = cell._tc
                        logger.info(f"[EXPORT]   ✅ 通过Python对象找到单元格: 逻辑列{col_pos} -> cell={cell}, tc={tc}")

                        # 在XML中找到对应的索引
                        for idx, tc_elem in enumerate(tc_elements_in_row):
                            if tc_elem == tc:
                                actual_cell_index = idx
                                logger.info(f"[EXPORT]   ✅ 在XML中找到对应索引: {actual_cell_index}")
                                break
                    else:
                        logger.warning(f"[EXPORT]   ⚠️ 逻辑位置{col_pos}超出范围，当前行只有{len(row_cells)}个单元格")

                except (IndexError, AttributeError) as e:
                    logger.warning(f"[EXPORT]   ⚠️ 从Python对象获取单元格失败: {e}")

                # 方法2：如果方法1失败，尝试从XML中查找（备用方案）
                if actual_cell_index is None:
                    logger.debug(f"[EXPORT]   方法1失败，尝试从XML查找逻辑列{col_pos}")
                    logical_col_position = 0

                    for idx, tc_elem in enumerate(tc_elements_in_row):
                        # 跳过已处理的单元格
                        if idx in processed_tc_indices:
                            logger.debug(f"[EXPORT]   跳过已处理的单元格: 索引{idx}, 逻辑位置={logical_col_position}")
                            continue
                        # 跳过将被删除的单元格（被当前行的合并单元格占用）
                        if idx in cells_to_remove_indices:
                            logger.debug(f"[EXPORT]   跳过将被删除的单元格: 索引{idx}, 逻辑位置={logical_col_position}")
                            continue

                        # 检查这个单元格是否是合并单元格
                        tc_pr = tc_elem.find(f'.//{ns}tcPr')
                        current_grid_span = None
                        if tc_pr is not None:
                            grid_span_elem = tc_pr.find(f'.//{ns}gridSpan')
                            if grid_span_elem is not None:
                                current_grid_span = int(grid_span_elem.get(qn('w:val'), 1))

                        logger.debug(f"[EXPORT]   检查单元格: 索引{idx}, 逻辑位置={logical_col_position}, 目标col_pos={col_pos}, gridSpan={current_grid_span}")

                        # 检查是否到达目标逻辑位置
                        if logical_col_position == col_pos:
                            actual_cell_index = idx
                            tc = tc_elem
                            logger.info(f"[EXPORT]   ✅ 通过XML查找找到单元格: 逻辑列{col_pos} -> 实际索引{idx}, 逻辑位置={logical_col_position}")
                            break

                        # 递增逻辑位置：如果当前单元格有gridSpan，则递增相应的值
                        if current_grid_span and current_grid_span > 1:
                            logical_col_position += current_grid_span
                            logger.debug(f"[EXPORT]   逻辑位置递增 {current_grid_span} 到 {logical_col_position}（因为gridSpan）")
                        else:
                            logical_col_position += 1
                            logger.debug(f"[EXPORT]   逻辑位置递增 1 到 {logical_col_position}")
                
                if actual_cell_index is None:
                    logger.warning(f"[EXPORT] ⚠️ 无法找到单元格 (行{row_idx}, 单元格{cell_idx}/{len(cells)}, 逻辑列{col_pos}, 总tc数={len(tc_elements_in_row)}, 已处理={processed_tc_indices}, 将删除={cells_to_remove_indices})")
                    # 即使找不到单元格，也要递增 col_pos，以便继续处理下一个单元格
                    col_pos += colspan
                    continue
                
                if row_idx >= len(table.rows):
                    logger.warning(f"[EXPORT] ⚠️ 行索引超出范围 (行{row_idx}, 总行数={len(table.rows)})")
                    col_pos += colspan
                    continue
                
                # 如果cell为None，尝试通过XML索引获取
                if cell is None and actual_cell_index is not None and actual_cell_index < len(tc_elements_in_row):
                    try:
                        tc = tc_elements_in_row[actual_cell_index]
                        # 从 table.rows 获取对应的 _Cell 对象（用于添加段落）
                        # 尝试通过 row.cells 查找对应的 cell 对象
                        for c in table.rows[row_idx].cells:
                            if c._tc == tc:
                                cell = c
                                break

                        # 如果找不到，尝试直接创建 _Cell 对象
                        if cell is None:
                            try:
                                cell = _Cell(tc, table.rows[row_idx])
                            except (TypeError, AttributeError) as e:
                                logger.warning(f"[EXPORT] ⚠️ 无法创建单元格对象: {e} (行{row_idx}, 逻辑列{col_pos})")
                                cell = None
                    except (IndexError, AttributeError):
                        cell = None

                if cell is None:
                    logger.warning(f"[EXPORT] ⚠️ 无法获取单元格对象 (行{row_idx}, 逻辑列{col_pos}, 实际索引{actual_cell_index})")
                    col_pos += colspan
                    continue
                else:
                    logger.info(f"[EXPORT] ✅ 成功获取单元格对象 (行{row_idx}, 逻辑列{col_pos})")
                
                tc_pr = tc.find(f'.//{ns}tcPr')
                if tc_pr is None:
                    tc_pr = OxmlElement('w:tcPr')
                    tc.insert(0, tc_pr)

                background_color = attrs.get("backgroundColor")
                if background_color:
                    self._set_cell_background_color(cell, background_color)
                
                # 标记这个单元格已处理
                processed_tc_indices.add(actual_cell_index)
                
                # 处理横向合并（colspan）
                # 注意：先设置 gridSpan，但暂时不删除被合并的单元格
                # 等所有单元格处理完后再删除，避免影响后续单元格的定位
                if colspan > 1:
                    # 设置 gridSpan 属性
                    grid_span = OxmlElement('w:gridSpan')
                    grid_span.set(qn('w:val'), str(colspan))
                    tc_pr.append(grid_span)
                    
                    # 记录需要删除的单元格索引（等所有单元格处理完后再删除）
                    for i in range(1, colspan):
                        if actual_cell_index + i < len(tc_elements_in_row):
                            tc_to_remove = tc_elements_in_row[actual_cell_index + i]
                            cells_to_remove_in_row.append((actual_cell_index + i, tc_to_remove))
                        
                        # 更新 occupied_positions，标记被删除的列位置
                        if row_idx not in occupied_positions:
                            occupied_positions[row_idx] = {}
                        occupied_positions[row_idx][col_pos + i] = True
                
                # 处理纵向合并（rowspan）
                if rowspan > 1:
                    # 标记起始单元格（vMerge 没有 val 或 val="restart"）
                    v_merge = OxmlElement('w:vMerge')
                    # 不设置 val 属性，表示这是合并的起始单元格
                    tc_pr.append(v_merge)
                    
                    # 标记后续行中对应的列位置为已占用
                    for r in range(1, rowspan):
                        if row_idx + r < len(table.rows):
                            if row_idx + r not in occupied_positions:
                                occupied_positions[row_idx + r] = {}
                            # 标记所有被 rowspan 覆盖的列
                            for c in range(colspan):
                                occupied_positions[row_idx + r][col_pos + c] = True
                            
                            # 为后续行的对应单元格设置 vMerge="continue"
                            # 需要找到后续行中对应位置的单元格
                            next_row_xml = table.rows[row_idx + r]._tr
                            next_tc_elements = next_row_xml.findall(f'.//{ns}tc')
                            # 找到对应逻辑列位置的单元格
                            next_tc_count = 0
                            for next_tc in next_tc_elements:
                                if next_tc_count == col_pos:
                                    next_tc_pr = next_tc.find(f'.//{ns}tcPr')
                                    if next_tc_pr is None:
                                        next_tc_pr = OxmlElement('w:tcPr')
                                        next_tc.insert(0, next_tc_pr)
                                    next_v_merge = OxmlElement('w:vMerge')
                                    next_v_merge.set(qn('w:val'), 'continue')
                                    next_tc_pr.append(next_v_merge)
                                    break
                                # 检查这个单元格是否有 gridSpan
                                next_tc_pr_temp = next_tc.find(f'.//{ns}tcPr')
                                if next_tc_pr_temp is not None:
                                    next_grid_span = next_tc_pr_temp.find(f'.//{ns}gridSpan')
                                    if next_grid_span is not None:
                                        next_colspan = int(next_grid_span.get(qn('w:val')) or '1')
                                        next_tc_count += next_colspan
                                    else:
                                        next_tc_count += 1
                                else:
                                    next_tc_count += 1

                # 设置单元格宽度（tcW）
                # 重要：在 Word 中，合并单元格的宽度应该由 gridSpan 和 tblGrid 自动计算
                # 因此，合并单元格不应该设置 tcW，而普通单元格应该设置 tcW 为对应列的宽度
                if colspan == 1:
                    # 普通单元格：设置 tcW 为对应列的宽度
                    cell_width_info = attrs.get("cellWidth")
                    if cell_width_info:
                        # 优先使用解析时提取的宽度数据
                        cell_width = cell_width_info.get("width")
                        width_type = cell_width_info.get("type", "dxa")
                    else:
                        # 如果没有提取到宽度信息，从 tblGrid 获取
                        table_attrs = node.get("attrs", {})
                        col_widths = table_attrs.get("colWidths", [])
                        if col_widths and col_pos < len(col_widths):
                            cell_width = col_widths[col_pos]
                            width_type = "dxa"
                        else:
                            cell_width = None
                    
                    if cell_width is not None:
                        # 检查是否已存在 tcW，如果存在则删除
                        existing_tc_w = tc_pr.tcW
                        if existing_tc_w is not None:
                            old_width = existing_tc_w.get(qn('w:w'))
                            tc_pr.remove(existing_tc_w)
                            logger.debug(f"[EXPORT] 删除已存在的 tcW: {old_width}")
                        
                        # 设置单元格宽度
                        tc_w = OxmlElement('w:tcW')
                        tc_w.set(qn('w:w'), str(cell_width))
                        tc_w.set(qn('w:type'), width_type)
                        tc_pr.append(tc_w)
                        logger.info(f"[EXPORT] ✅✅✅ 设置普通单元格宽度: {cell_width} {width_type} (行{row_idx}, 列{col_pos})")
                    else:
                        # 如果没有宽度信息，删除可能存在的 tcW
                        existing_tc_w = tc_pr.tcW
                        if existing_tc_w is not None:
                            old_width = existing_tc_w.get(qn('w:w'))
                            tc_pr.remove(existing_tc_w)
                            logger.warning(f"[EXPORT] ⚠️ 未找到普通单元格宽度信息，删除 tcW: {old_width} (行{row_idx}, 列{col_pos})")
                else:
                    # 合并单元格：需要设置 tcW 为被合并列的总宽度
                    # 虽然 Word 理论上可以根据 gridSpan 和 tblGrid 计算，但实际中需要显式设置 tcW
                    cell_width_info = attrs.get("cellWidth")
                    table_attrs = node.get("attrs", {})
                    col_widths = table_attrs.get("colWidths", [])
                    
                    if cell_width_info:
                        # 优先使用解析时提取的宽度数据（这是原始文档中的值）
                        merged_width = cell_width_info.get("width")
                        width_type = cell_width_info.get("type", "dxa")
                    elif col_widths and col_pos + colspan <= len(col_widths):
                        # 如果没有提取到宽度，从 tblGrid 计算被合并列的总宽度
                        merged_width = sum(col_widths[col_pos:col_pos + colspan])
                        width_type = "dxa"
                    else:
                        merged_width = None
                    
                    if merged_width is not None:
                        # 删除已存在的 tcW
                        existing_tc_w = tc_pr.tcW
                        if existing_tc_w is not None:
                            old_width = existing_tc_w.get(qn('w:w'))
                            tc_pr.remove(existing_tc_w)
                            logger.debug(f"[EXPORT] 删除已存在的 tcW: {old_width}")
                        
                        # 设置合并单元格宽度为被合并列的总宽度
                        tc_w = OxmlElement('w:tcW')
                        tc_w.set(qn('w:w'), str(merged_width))
                        tc_w.set(qn('w:type'), width_type)
                        tc_pr.append(tc_w)
                        logger.info(f"[EXPORT] ✅✅✅ 设置合并单元格宽度: {merged_width} {width_type} (行{row_idx}, 列{col_pos}, colspan={colspan})")
                    else:
                        # 如果无法确定宽度，删除 tcW，让 Word 自动计算
                        existing_tc_w = tc_pr.tcW
                        if existing_tc_w is not None:
                            old_width = existing_tc_w.get(qn('w:w'))
                            tc_pr.remove(existing_tc_w)
                            logger.warning(f"[EXPORT] ⚠️ 无法确定合并单元格宽度，删除 tcW: {old_width} (行{row_idx}, 列{col_pos})")

                # 应用单元格格式（垂直对齐）
                if "verticalAlign" in attrs:
                    alignment = VERTICAL_ALIGNMENT_REVERSE_MAPPING.get(
                        attrs["verticalAlign"], WD_ALIGN_VERTICAL.TOP
                    )
                    # 通过 XML 设置垂直对齐
                    # Word 中垂直对齐通过 tcPr 的 vAlign 属性设置
                    v_align = OxmlElement('w:vAlign')
                    if alignment == WD_ALIGN_VERTICAL.CENTER:
                        v_align.set(qn('w:val'), 'center')
                    elif alignment == WD_ALIGN_VERTICAL.BOTTOM:
                        v_align.set(qn('w:val'), 'bottom')
                    else:
                        v_align.set(qn('w:val'), 'top')
                    tc_pr.append(v_align)

                # 添加单元格内容
                # 使用之前获取的 _Cell 对象来添加段落
                cell_content = cell_node.get("content", [])
                # 优化空段落处理：避免连续的空段落导致额外的空行
                # 策略：保留所有有内容的段落，对于连续的空段落，只保留最后一个
                filtered_paras = []
                last_was_empty = False
                
                for para_node in cell_content:
                    if para_node.get("type") == "paragraph" or para_node.get("type") == "heading":
                        # 检查段落是否有文本内容（包括 text 和 placeholder 节点）
                        para_content = para_node.get("content", [])
                        has_text = any(
                            (item.get("type") == "text" and item.get("text", "").strip()) or
                            (item.get("type") == "placeholder" and item.get("attrs", {}).get("value") is not None)
                            for item in para_content
                        )
                        # 检查是否有特殊格式（对齐、缩进、间距等）
                        attrs = para_node.get("attrs", {})
                        has_formatting = any(
                            key in attrs
                            for key in ["textAlign", "indent", "spacing", "style"]
                        ) or para_node.get("type") == "heading"
                        
                        is_empty = not has_text and not has_formatting
                        
                        # 如果有内容，直接添加
                        if has_text or has_formatting:
                            filtered_paras.append(para_node)
                            last_was_empty = False
                        # 如果是空段落
                        elif is_empty:
                            # 如果上一个也是空段落，跳过这个（避免连续空行）
                            if not last_was_empty:
                                filtered_paras.append(para_node)
                                last_was_empty = True
                            # 否则跳过这个空段落
                        else:
                            filtered_paras.append(para_node)
                            last_was_empty = False
                
                # 如果过滤后没有段落，至少添加一个空段落（Word 要求）
                if not filtered_paras:
                    # 创建一个最小的空段落
                    filtered_paras = [{"type": "paragraph", "attrs": {}, "content": []}]
                
                # 添加过滤后的段落
                # 只有在成功获取到 cell 对象时才添加内容
                logger.info(f"[EXPORT] 检查单元格内容: cell={cell is not None}, filtered_paras={len(filtered_paras)} (行{row_idx}, 列{col_pos})")
                if cell is not None:
                    logger.info(f"[EXPORT] 准备添加 {len(filtered_paras)} 个段落到单元格 (行{row_idx}, 列{col_pos})")
                    for para_node in filtered_paras:
                        logger.info(f"[EXPORT] 添加段落: type={para_node.get('type')}, content_len={len(para_node.get('content', []))}")
                        self.map_paragraph_node(para_node, cell)
                    logger.info(f"[EXPORT] ✅ 成功添加单元格内容 (行{row_idx}, 列{col_pos})")
                else:
                    logger.warning(f"[EXPORT] ⚠️ 无法获取单元格对象，跳过内容添加 (行{row_idx}, 列{col_pos})")
                
                # 递增逻辑列位置
                old_col_pos = col_pos
                col_pos += colspan
                logger.debug(f"[EXPORT] 处理完单元格 {cell_idx}，col_pos: {old_col_pos} -> {col_pos} (colspan={colspan})")
            
            # 关键修复：在处理完所有单元格后，再删除被 gridSpan 占用的单元格
            # 这样可以确保所有单元格都能正确定位和添加内容
            if cells_to_remove_in_row:
                # 从后往前删除，避免索引变化影响
                cells_to_remove_in_row.sort(reverse=True, key=lambda x: x[0])
                for remove_idx, tc_to_remove in cells_to_remove_in_row:
                    parent = tc_to_remove.getparent()
                    if parent is not None:
                        parent.remove(tc_to_remove)
                        processed_tc_indices.add(remove_idx)
                        logger.info(f"[EXPORT] ✅ 删除被 gridSpan 占用的单元格 (行{row_idx}, 原索引{remove_idx})")
                logger.info(f"[EXPORT] 行{row_idx} 处理完成，删除了 {len(cells_to_remove_in_row)} 个被合并的单元格")

        return table

    def _apply_marks(self, run: Run, marks: List[Dict[str, Any]]) -> None:
        """应用文本标记到 Run 对象"""
        for mark in marks:
            mark_type = mark.get("type")
            if mark_type == "bold":
                run.bold = True
            elif mark_type == "italic":
                run.italic = True
            elif mark_type == "underline":
                run.underline = True
            elif mark_type == "strike":
                run.font.strike = True
            elif mark_type == "textStyle":
                attrs = mark.get("attrs", {})

                if "color" in attrs:
                    rgb_color = hex_to_rgb(attrs["color"])
                    if rgb_color:
                        run.font.color.rgb = rgb_color

                if "fontSize" in attrs:
                    # 解析字体大小（如 "14pt"）
                    font_size_str = attrs["fontSize"]
                    try:
                        if font_size_str.endswith("pt"):
                            size_pt = float(font_size_str[:-2])
                            run.font.size = Pt(size_pt)
                    except (ValueError, AttributeError):
                        logger.warning(f"无法解析字体大小: {font_size_str}")

                if "fontFamily" in attrs:
                    try:
                        run.font.name = attrs["fontFamily"]
                    except (AttributeError, TypeError):
                        logger.warning(f"无法设置字体名称: {attrs['fontFamily']}")

                if attrs.get("smallCaps"):
                    run.font.small_caps = True
                if attrs.get("allCaps"):
                    run.font.all_caps = True

    def _set_cell_background_color(self, cell: _Cell, color_hex: str) -> None:
        """设置单元格背景色"""
        if not color_hex:
            return

        color = color_hex.lstrip("#").upper()
        if len(color) != 6:
            return

        tc = cell._tc
        tc_pr = tc.tcPr
        if tc_pr is None:
            tc_pr = OxmlElement("w:tcPr")
            tc.insert(0, tc_pr)

        shading = tc_pr.find(f"./{W_NS}shd")
        if shading is None:
            shading = OxmlElement("w:shd")
            tc_pr.append(shading)

        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), color)

