from typing import List, Dict, Any, Optional
import json
import docx
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import Table
from docx.text.paragraph import Paragraph
from agno.tools import Toolkit

class SmartDocGenToolkit(Toolkit):
    def __init__(self):
        super().__init__(name="smart_doc_gen_toolkit")
        self.register(self.analyze_template)
        self.register(self.extract_structure)
        self.register(self.fill_template)

    def analyze_template(self, template_path: str) -> str:
        """
        Analyzes the DOCX template to determine its type (element-based vs narrative).
        Returns a JSON string with type and stats.
        
        Args:
            template_path: Absolute path to the .docx file.
        """
        try:
            doc = docx.Document(template_path)
            
            p_chars = 0
            t_chars = 0
            
            for p in doc.paragraphs:
                p_chars += len(p.text.strip())
                
            for t in doc.tables:
                for row in t.rows:
                    for cell in row.cells:
                        t_chars += len(cell.text.strip())
                        
            doc_type = "unknown"
            if p_chars == 0 and t_chars == 0:
                doc_type = "unknown"
            elif t_chars > p_chars:
                doc_type = "element"  # Mostly tables
            else:
                doc_type = "narrative" # Mostly text
                
            return json.dumps({
                "type": doc_type,
                "stats": {"p_chars": p_chars, "t_chars": t_chars}
            }, ensure_ascii=False)
        except Exception as e:
            return json.dumps({"error": str(e)})

    def extract_structure(self, template_path: str) -> str:
        """
        Extracts the structure (paragraphs and tables) from the DOCX template.
        Returns a JSON string list of blocks with IDs.
        
        Args:
            template_path: Absolute path to the .docx file.
        """
        try:
            doc = docx.Document(template_path)
            blocks = []
            
            p_counter = 0
            t_counter = 0
            
            p_iter = iter(doc.paragraphs)
            t_iter = iter(doc.tables)
            
            # Iterate over element.body.iterchildren to preserve order
            for child in doc.element.body.iterchildren():
                if isinstance(child, CT_P):
                    try:
                        p = next(p_iter)
                        blocks.append({
                            "id": f"p_{p_counter}",
                            "type": "paragraph",
                            "text": p.text
                        })
                        p_counter += 1
                    except StopIteration:
                        pass
                        
                elif isinstance(child, CT_Tbl):
                    try:
                        t = next(t_iter)
                        table_data = []
                        for r_idx, row in enumerate(t.rows):
                            row_data = []
                            for c_idx, cell in enumerate(row.cells):
                                row_data.append({
                                    "id": f"t_{t_counter}_r{r_idx}_c{c_idx}",
                                    "text": cell.text
                                })
                            table_data.append(row_data)
                        
                        blocks.append({
                            "id": f"t_{t_counter}",
                            "type": "table",
                            "rows": table_data
                        })
                        t_counter += 1
                    except StopIteration:
                        pass
                        
            return json.dumps(blocks, ensure_ascii=False, indent=2)
        except Exception as e:
            return json.dumps({"error": str(e)})

    def fill_template(self, template_path: str, output_path: str, fillings_json: str) -> str:
        """
        Fills the DOCX template with the provided content and saves to output_path.
        
        Args:
            template_path: Absolute path to the source .docx file.
            output_path: Absolute path where the filed document should be saved.
            fillings_json: JSON string mapping block IDs to new content. 
                           {"p_0": "New Text", "t_0_r0_c0": "Cell Content"}
        """
        try:
            fillings = json.loads(fillings_json)
            doc = docx.Document(template_path)
            
            # Fill Paragraphs
            paragraphs_to_process = [(i, p) for i, p in enumerate(doc.paragraphs)]
            for i, p in paragraphs_to_process:
                pid = f"p_{i}"
                if pid in fillings:
                    val = fillings[pid]
                    if val.startswith(p.text) and len(p.text) > 0:
                        # Smart Append: Keep original runs (preserving format/bolding) and append only the new content
                        diff = val[len(p.text):]
                        if diff:
                            p.add_run(diff)
                    else:
                        # Fallback: Overwrite if content doesn't match prefix
                        p.text = val
                    
            # Fill Tables
            for i, t in enumerate(doc.tables):
                for r_idx, row in enumerate(t.rows):
                    for c_idx, cell in enumerate(row.cells):
                        cid = f"t_{i}_r{r_idx}_c{c_idx}"
                        if cid in fillings:
                            val = fillings[cid]
                            current_text = cell.text
                            if val.startswith(current_text) and len(current_text) > 0:
                                # Smart Append for cells
                                diff = val[len(current_text):]
                                if diff:
                                    # Append to the last paragraph of the cell 
                                    # (Cells always have at least one paragraph)
                                    if len(cell.paragraphs) > 0:
                                        cell.paragraphs[-1].add_run(diff)
                                    else:
                                        cell.add_paragraph(diff)
                            else:
                                cell.text = val
                            
            doc.save(output_path)
            return json.dumps({"status": "success", "output": output_path})
        except Exception as e:
            return json.dumps({"error": str(e)})
