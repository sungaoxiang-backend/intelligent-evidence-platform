"""
模板解析和文档转换工具函数
"""
import io
import logging
import re
from pathlib import Path
from typing import Any, Dict, List, Optional, Set

import mammoth
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from app.lex_docx.schemas import PlaceholderMetadata, validate_placeholder_name

logger = logging.getLogger(__name__)


def _get_alignment_style(alignment) -> str:
    """
    获取段落对齐方式对应的 CSS 样式
    
    注意：此函数已被 FormatMapper.docx_alignment_to_css 替代
    保留此函数是为了向后兼容，但应该逐步迁移到 FormatMapper
    
    Args:
        alignment: WD_ALIGN_PARAGRAPH 枚举值或 None
        
    Returns:
        CSS text-align 值
    """
    from app.lex_docx.format_mapper import FormatMapper
    return FormatMapper.docx_alignment_to_css(alignment)


def _get_vertical_alignment_style(vertical_alignment) -> str:
    """
    获取单元格垂直对齐方式对应的 CSS 样式
    
    Args:
        vertical_alignment: 单元格垂直对齐枚举值
        
    Returns:
        CSS vertical-align 值
    """
    if vertical_alignment is None:
        return "top"
    
    # python-docx 中垂直对齐：0=TOP, 1=CENTER, 2=BOTTOM
    if vertical_alignment == 0:  # TOP
        return "top"
    elif vertical_alignment == 1:  # CENTER
        return "middle"
    elif vertical_alignment == 2:  # BOTTOM
        return "bottom"
    else:
        return "top"


def _get_indent_style_from_xml(para_xml) -> str:
    """
    直接从 XML 获取段落缩进样式（更可靠的方法）
    
    Args:
        para_xml: 段落的 XML 元素
        
    Returns:
        CSS 缩进样式字符串
    """
    from docx.oxml.ns import qn
    
    styles = []
    
    try:
        # 获取段落属性
        pPr = para_xml.find(qn('w:pPr'))
        if pPr is not None:
            # 获取缩进属性
            ind = pPr.find(qn('w:ind'))
            if ind is not None:
                # 首行缩进 (firstLine) - 这是首行相对于左缩进的额外缩进
                first_line = ind.get(qn('w:firstLine'))
                if first_line is not None:
                    # firstLine 单位是 twips (1/20 pt)
                    indent_pt = float(first_line) / 20.0
                    if abs(indent_pt) > 0.01:
                        styles.append(f"text-indent: {indent_pt:.2f}pt")
                        logger.info(f"从XML提取首行缩进: {indent_pt:.2f}pt (firstLine={first_line})")
                
                # 左缩进 (left) - 这是整个段落的左缩进
                left = ind.get(qn('w:left'))
                if left is not None:
                    indent_pt = float(left) / 20.0
                    if abs(indent_pt) > 0.01:
                        styles.append(f"padding-left: {indent_pt:.2f}pt")
                        logger.info(f"从XML提取左缩进: {indent_pt:.2f}pt (left={left})")
    except Exception as e:
        logger.debug(f"从XML提取缩进失败: {e}")
    
    return "; ".join(styles)


def _get_indent_style(paragraph_format, para_element=None) -> str:
    """
    获取段落缩进对应的 CSS 样式
    
    Args:
        paragraph_format: python-docx ParagraphFormat 对象
        para_element: 段落的 XML 元素（可选，如果提供则优先使用）
        
    Returns:
        CSS 缩进样式字符串
    """
    from docx.shared import Pt, Inches, Length, Twips
    
    styles = []
    
    # 首先尝试从 XML 直接读取（最可靠的方法）
    try:
        if para_element is not None:
            xml_style = _get_indent_style_from_xml(para_element)
            if xml_style:
                logger.info(f"从XML成功提取缩进样式: {xml_style}")
                return xml_style
        elif hasattr(paragraph_format, '_element'):
            xml_style = _get_indent_style_from_xml(paragraph_format._element)
            if xml_style:
                logger.info(f"从paragraph_format._element成功提取缩进样式: {xml_style}")
                return xml_style
    except Exception as e:
        logger.debug(f"尝试从XML读取缩进失败: {e}")
    
    # 回退到使用 paragraph_format API
    # 左缩进（left_indent）- 优先处理，因为这是最常见的缩进方式
    try:
        left_indent = paragraph_format.left_indent
        if left_indent is not None:
            indent_pt = None
            
            # Twips 对象可以直接使用 .pt 属性
            if isinstance(left_indent, Twips):
                indent_pt = left_indent.pt
            elif isinstance(left_indent, Length):
                indent_pt = left_indent.pt
            elif hasattr(left_indent, 'pt'):
                indent_pt = left_indent.pt
            elif isinstance(left_indent, (int, float)):
                # 如果是整数，可能是 twips (1/20 pt)
                indent_pt = float(left_indent) / 20.0
            
            if indent_pt is not None and abs(indent_pt) > 0.01:
                styles.append(f"padding-left: {indent_pt:.2f}pt")
                logger.info(f"从paragraph_format提取左缩进: {indent_pt:.2f}pt (类型: {type(left_indent)})")
    except Exception as e:
        logger.debug(f"获取左缩进失败: {e}")
    
    # 首行缩进（first_line_indent）
    try:
        first_line_indent = paragraph_format.first_line_indent
        if first_line_indent is not None:
            indent_pt = None
            
            if isinstance(first_line_indent, Twips):
                indent_pt = first_line_indent.pt
            elif isinstance(first_line_indent, Length):
                indent_pt = first_line_indent.pt
            elif hasattr(first_line_indent, 'pt'):
                indent_pt = first_line_indent.pt
            elif isinstance(first_line_indent, (int, float)):
                indent_pt = float(first_line_indent) / 20.0
            
            if indent_pt is not None and abs(indent_pt) > 0.01:
                styles.append(f"text-indent: {indent_pt:.2f}pt")
                logger.info(f"从paragraph_format提取首行缩进: {indent_pt:.2f}pt (类型: {type(first_line_indent)})")
    except Exception as e:
        logger.debug(f"获取首行缩进失败: {e}")
    
    result = "; ".join(styles)
    if result:
        logger.info(f"最终缩进样式: {result}")
    return result


def _inject_alignment_info(html: str, doc) -> str:
    """
    将 DOCX 文档中的对齐信息和缩进信息注入到 HTML 中
    
    Args:
        html: mammoth 生成的 HTML
        doc: python-docx Document 对象
        
    Returns:
        注入对齐和缩进信息后的 HTML
    """
    import re
    from html import escape
    
    # 1. 为没有 border 属性的 table 添加默认样式，并应用列宽度
    # 先提取列宽度（需要在处理表格之前）
    table_column_widths = {}  # {table_idx: [width1, width2, ...]}
    for table_idx, table in enumerate(doc.tables):
        column_widths = []
        try:
            # 从表格的 XML 中提取列宽度
            tbl = table._element
            from docx.oxml.ns import qn
            tblGrid = tbl.find(qn('w:tblGrid'))
            if tblGrid is not None:
                gridCols = tblGrid.findall(qn('w:gridCol'))
                for gridCol in gridCols:
                    w = gridCol.get(qn('w:w'))
                    if w:
                        # w 单位是 twips (1/20 pt)，转换为 pt
                        width_pt = float(w) / 20.0
                        column_widths.append(width_pt)
                    else:
                        column_widths.append(None)
        except Exception as e:
            logger.debug(f"提取表格 {table_idx} 列宽度失败: {e}")
        
        if column_widths:
            table_column_widths[table_idx] = column_widths
            logger.info(f"表格 {table_idx} 列宽度: {column_widths}")
    
    # 为表格添加样式和列宽度
    table_pattern = r'<table([^>]*)>'
    tables = list(re.finditer(table_pattern, html))
    offset = 0
    for table_idx, table_match in enumerate(tables):
        table_attrs = table_match.group(1)
        table_start = table_match.start() + offset
        table_end = table_match.end() + offset
        
        # 构建样式
        styles = ["border-collapse: collapse", "border: 1px solid #000"]
        if table_idx in table_column_widths:
            styles.append("table-layout: fixed")
        
        # 检查是否已有 style 属性
        if 'style=' in table_attrs:
            # 合并现有样式
            style_match = re.search(r'style="([^"]*)"', table_attrs)
            if style_match:
                existing_style = style_match.group(1)
                new_style = existing_style + "; " + "; ".join(styles)
                table_attrs = table_attrs.replace(f'style="{existing_style}"', f'style="{new_style}"')
            else:
                table_attrs += f' style="{"; ".join(styles)}"'
        else:
            table_attrs += f' style="{"; ".join(styles)}"'
        
        # 替换表格标签
        new_table_tag = f'<table{table_attrs}>'
        html = html[:table_start] + new_table_tag + html[table_end:]
        offset += len(new_table_tag) - (table_end - table_start)
        
        # 如果有列宽度信息，添加 <colgroup>
        if table_idx in table_column_widths:
            column_widths = table_column_widths[table_idx]
            if column_widths:
                colgroup_html = '<colgroup>'
                total_width = sum(w for w in column_widths if w is not None)
                if total_width > 0:
                    for width in column_widths:
                        if width is not None:
                            # 计算百分比宽度
                            percentage = (width / total_width) * 100
                            colgroup_html += f'<col style="width: {percentage:.2f}%;">'
                        else:
                            colgroup_html += '<col>'
                else:
                    # 如果没有宽度信息，平均分配
                    for _ in column_widths:
                        colgroup_html += '<col>'
                colgroup_html += '</colgroup>'
                # 在 <table> 标签后插入 <colgroup>
                html = html[:table_start + len(new_table_tag)] + colgroup_html + html[table_start + len(new_table_tag):]
                offset += len(colgroup_html)
    
    # 2. 处理文档级别段落对齐和缩进
    # 收集所有段落及其格式信息
    paragraph_info = []
    for para in doc.paragraphs:
        if para.text.strip():
            text = para.text.strip()
            alignment = _get_alignment_style(para.alignment)
            indent_style = _get_indent_style(para.paragraph_format)
            # 使用文本的前50个字符作为键（避免过长）
            key = text[:50] if len(text) > 50 else text
            paragraph_info.append({
                'text': key,
                'alignment': alignment,
                'indent_style': indent_style,
            })
    
    # 为文档级别段落注入对齐和缩进样式
    for para_info in paragraph_info:
        if para_info['alignment'] != "left" or para_info['indent_style']:
            escaped_text = re.escape(para_info['text'])
            # 匹配包含该文本的 p 或 h 标签（不在表格单元格内）
            pattern = rf'<(p|h[1-6])([^>]*)>(.*?{escaped_text}.*?)</\1>'
            def add_para_format(match):
                tag = match.group(1)
                attrs = match.group(2)
                content = match.group(3)
                
                # 构建样式字符串
                styles = []
                if 'style=' in attrs:
                    style_match = re.search(r'style="([^"]*)"', attrs)
                    if style_match:
                        existing_style = style_match.group(1)
                        styles.append(existing_style)
                        attrs = re.sub(r'style="[^"]*"', '', attrs)
                
                # 添加对齐样式
                if para_info and para_info['alignment'] != "left":
                    if 'text-align' not in '; '.join(styles):
                        styles.append(f"text-align: {para_info['alignment']}")
                
                # 添加缩进样式
                if para_info and para_info['indent_style']:
                    styles.append(para_info['indent_style'])
                
                if styles:
                    style_str = "; ".join(styles)
                    return f'<{tag}{attrs} style="{style_str}">{content}</{tag}>'
                return match.group(0)
            
            html = re.sub(pattern, add_para_format, html, flags=re.DOTALL)
    
    # 3. 处理表格单元格内的段落对齐和缩进
    # 新方法：直接遍历表格，为每个单元格建立精确的段落映射
    cell_info_map = {}  # {(table_idx, row_idx, col_idx): cell_info}
    table_column_widths = {}  # {table_idx: [width1, width2, ...]}
    
    # 提取表格列宽度
    for table_idx, table in enumerate(doc.tables):
        column_widths = []
        try:
            # 从表格的 XML 中提取列宽度
            tbl = table._element
            from docx.oxml.ns import qn
            tblGrid = tbl.find(qn('w:tblGrid'))
            if tblGrid is not None:
                gridCols = tblGrid.findall(qn('w:gridCol'))
                for gridCol in gridCols:
                    w = gridCol.get(qn('w:w'))
                    if w:
                        # w 单位是 twips (1/20 pt)，转换为 pt
                        width_pt = float(w) / 20.0
                        column_widths.append(width_pt)
                else:
                        column_widths.append(None)
        except Exception as e:
            logger.debug(f"提取表格 {table_idx} 列宽度失败: {e}")
        
        if column_widths:
            table_column_widths[table_idx] = column_widths
            logger.info(f"表格 {table_idx} 列宽度: {column_widths}")
    
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                cell_key = (table_idx, row_idx, col_idx)
                cell_vertical_alignment = _get_vertical_alignment_style(cell.vertical_alignment)
                
                # 收集单元格内所有段落的格式信息（包括空段落，保持顺序）
                paragraphs_info = []
                for para_idx, para in enumerate(cell.paragraphs):
                    text = para.text.strip()
                    alignment = _get_alignment_style(para.alignment)
                    # 使用改进的缩进提取方法，传递段落元素以便从XML读取
                    para_element = para._element if hasattr(para, '_element') else None
                    indent_style = _get_indent_style(para.paragraph_format, para_element)
                    
                    # 记录详细信息用于调试
                    logger.info(
                        f"单元格[{table_idx},{row_idx},{col_idx}]段落[{para_idx}]: "
                        f"文本='{text[:30]}...', 对齐={alignment}, 缩进={indent_style}"
                    )
                    
                    # 如果从XML提取失败，尝试直接从paragraph_format获取（使用Twips对象的.pt属性）
                    if not indent_style:
                        try:
                            from docx.shared import Twips
                            pf = para.paragraph_format
                            
                            # 优先处理左缩进（更常见）
                            if pf.left_indent is not None:
                                if isinstance(pf.left_indent, Twips):
                                    indent_pt = pf.left_indent.pt
                                    if abs(indent_pt) > 0.01:
                                        indent_style = f"padding-left: {indent_pt:.2f}pt"
                                        logger.info(f"从Twips对象提取左缩进: {indent_style}")
                            
                            # 然后处理首行缩进
                            if not indent_style and pf.first_line_indent is not None:
                                if isinstance(pf.first_line_indent, Twips):
                                    indent_pt = pf.first_line_indent.pt
                                    if abs(indent_pt) > 0.01:
                                        indent_style = f"text-indent: {indent_pt:.2f}pt"
                                        logger.info(f"从Twips对象提取首行缩进: {indent_style}")
                        except Exception as e:
                            logger.debug(f"尝试从Twips对象提取失败: {e}")
                    
                    # 即使文本为空，也记录格式信息（mammoth 可能合并了段落）
                    paragraphs_info.append({
                        'index': para_idx,  # 段落索引，用于精确匹配
                        'text': text[:50] if len(text) > 50 else text,
                        'full_text': text,
                        'alignment': alignment,
                        'indent_style': indent_style,  # 使用更新后的缩进样式
                        'has_text': bool(text),  # 标记是否有文本
                    })
                
                if paragraphs_info or cell_vertical_alignment != "top":
                    cell_info_map[cell_key] = {
                        'paragraphs': paragraphs_info,
                        'vertical_align': cell_vertical_alignment,
                    }
                    logger.debug(
                        f"单元格[{table_idx},{row_idx},{col_idx}]信息: "
                        f"{len(paragraphs_info)}个段落, 垂直对齐={cell_vertical_alignment}"
                    )
    
    # 为每个单元格注入格式样式
    # 新方法：使用表格索引来精确定位单元格，避免文本匹配的不确定性
    # 注意：对于合并单元格，只处理第一个单元格（col_idx=0），其他列会被合并单元格覆盖
    processed_cells = set()  # 记录已处理的HTML单元格位置，避免重复处理合并单元格
    
    for cell_key, cell_info in cell_info_map.items():
        table_idx, row_idx, col_idx = cell_key
        # 注意：mammoth可能不会将第一行转换为th，所以同时尝试td和th
        # 先尝试td（更常见）
        tag = 'td'
        
        # 方法：通过表格结构定位，而不是文本匹配
        # 首先找到对应的表格
        table_pattern = rf'<table[^>]*>.*?</table>'
        tables = list(re.finditer(table_pattern, html, flags=re.DOTALL))
        
        if table_idx >= len(tables):
            logger.warning(f"表格索引 {table_idx} 超出范围，共有 {len(tables)} 个表格")
            continue
        
        table_match = tables[table_idx]
        table_html = table_match.group(0)
    
        # 在表格内找到对应的行
        row_pattern = rf'<tr[^>]*>.*?</tr>'
        rows = list(re.finditer(row_pattern, table_html, flags=re.DOTALL))
        
        if row_idx >= len(rows):
            logger.warning(f"行索引 {row_idx} 超出范围，共有 {len(rows)} 行")
            continue
        
        row_match = rows[row_idx]
        row_html = row_match.group(0)
        
        # 在行内找到对应的单元格
        # 注意：需要处理合并单元格（colspan）的情况
        # 同时尝试td和th标签（mammoth可能使用不同的标签）
        cell_pattern_td = r'<td[^>]*>.*?</td>'
        cell_pattern_th = r'<th[^>]*>.*?</th>'
        cells_td = list(re.finditer(cell_pattern_td, row_html, flags=re.DOTALL))
        cells_th = list(re.finditer(cell_pattern_th, row_html, flags=re.DOTALL))
        
        # 优先使用td，如果没有则使用th
        cells = cells_td if cells_td else cells_th
        tag = 'td' if cells_td else 'th'
        
        if not cells:
            logger.warning(f"行 {row_idx} 中没有找到td或th标签")
            continue
        
        # 计算实际列位置（考虑colspan）
        # 对于合并单元格，我们需要找到包含目标列索引的HTML单元格
        actual_col_idx = None
        current_col = 0
        for idx, cell_match in enumerate(cells):
            cell_attrs = cell_match.group(0)
            # 检查是否有colspan属性
            colspan_match = re.search(r'colspan=["\']?(\d+)["\']?', cell_attrs)
            colspan = int(colspan_match.group(1)) if colspan_match else 1
            
            # 检查目标列是否在这个单元格的范围内
            if current_col <= col_idx < current_col + colspan:
                actual_col_idx = idx
                break
            current_col += colspan
        
        if actual_col_idx is None or actual_col_idx >= len(cells):
            logger.warning(
                f"列索引 {col_idx} 超出范围（考虑colspan后），"
                f"共有 {len(cells)} 个单元格（tag={tag}），当前列位置: {current_col}"
            )
            continue
        
        # 检查这个HTML单元格是否已经被处理过（合并单元格的情况）
        html_cell_key = (table_idx, row_idx, actual_col_idx)
        if html_cell_key in processed_cells:
            logger.debug(f"单元格[{table_idx},{row_idx},{col_idx}]对应的HTML单元格已处理，跳过")
            continue
        
        # 对于合并单元格，合并同一行所有单元格的段落格式信息
        # 找到映射到同一个HTML单元格的所有DOCX单元格
        merged_paragraphs = []
        for key, info in cell_info_map.items():
            if (key[0] == table_idx and key[1] == row_idx):
                # 检查这个DOCX单元格是否映射到同一个HTML单元格
                # 计算这个单元格对应的HTML列索引
                other_col_idx = key[2]
                other_current_col = 0
                other_actual_col_idx = None
                for idx, other_cell_match in enumerate(cells):
                    other_cell_attrs = other_cell_match.group(0)
                    other_colspan_match = re.search(r'colspan=["\']?(\d+)["\']?', other_cell_attrs)
                    other_colspan = int(other_colspan_match.group(1)) if other_colspan_match else 1
                    if other_current_col <= other_col_idx < other_current_col + other_colspan:
                        other_actual_col_idx = idx
                        break
                    other_current_col += other_colspan
                
                # 如果映射到同一个HTML单元格，合并段落格式
                if other_actual_col_idx == actual_col_idx:
                    merged_paragraphs.extend(info['paragraphs'])
        
        # 如果找到了合并的段落，使用合并后的段落信息
        if merged_paragraphs and len(merged_paragraphs) > len(cell_info['paragraphs']):
            cell_info = {
                'paragraphs': merged_paragraphs,
                'vertical_align': cell_info['vertical_align'],
            }
            logger.info(
                f"合并单元格[{table_idx},{row_idx}]: "
                f"合并了 {len(merged_paragraphs)} 个段落的格式信息（来自多个DOCX单元格）"
            )
        
        processed_cells.add(html_cell_key)
        
        cell_match = cells[actual_col_idx]
        cell_full_match = cell_match.group(0)
        cell_attrs_match = re.search(rf'<{tag}([^>]*)>', cell_full_match)
        cell_content_match = re.search(rf'<{tag}[^>]*>(.*?)</{tag}>', cell_full_match, flags=re.DOTALL)
        
        if not cell_attrs_match or not cell_content_match:
            logger.warning(f"无法解析单元格 [{table_idx},{row_idx},{col_idx}] 的结构")
            continue
        
        cell_attrs = cell_attrs_match.group(1)
        cell_content = cell_content_match.group(1)
                
        # 处理单元格样式
        cell_styles = []
        if 'style=' in cell_attrs:
            style_match = re.search(r'style="([^"]*)"', cell_attrs)
            if style_match:
                existing_style = style_match.group(1)
                cell_styles.append(existing_style)
                cell_attrs = re.sub(r'style="[^"]*"', '', cell_attrs)
            else:
                cell_styles.append("border: 1px solid #000; padding: 4pt 8pt")
                
        # 添加单元格垂直对齐（重要：必须应用，特别是对于拆分/合并的单元格）
        has_vertical_align = any('vertical-align' in s for s in cell_styles)
        if not has_vertical_align or cell_info['vertical_align'] != "top":
            cell_styles = [s for s in cell_styles if 'vertical-align' not in s]
            cell_styles.append(f"vertical-align: {cell_info['vertical_align']}")
        
        # 不再添加表头样式（背景色和粗体），只保留必要的格式
        
        # 处理单元格内容中的段落
        # 先记录原始HTML结构用于调试
        logger.info(f"单元格[{table_idx},{row_idx},{col_idx}]原始内容: {cell_content[:200]}...")
        logger.info(f"单元格[{table_idx},{row_idx},{col_idx}]DOCX段落数: {len(cell_info['paragraphs'])}")
        
        # 查找所有段落元素
        para_elements = list(re.finditer(r'<(p)([^>]*)>(.*?)</p>', cell_content, flags=re.DOTALL))
        logger.info(f"单元格[{table_idx},{row_idx},{col_idx}]HTML段落数: {len(para_elements)}")
        
        # 如果段落数不匹配，记录警告
        if len(para_elements) != len(cell_info['paragraphs']):
            logger.warning(
                f"单元格[{table_idx},{row_idx},{col_idx}]段落数不匹配: "
                f"HTML有{len(para_elements)}个段落，DOCX有{len(cell_info['paragraphs'])}个段落"
            )
        
        updated_paragraphs = []
        
        # 按顺序处理每个HTML段落
        for para_idx, para_match in enumerate(para_elements):
            p_tag = para_match.group(1)
            p_attrs = para_match.group(2)
            p_content = para_match.group(3)
            clean_text = re.sub(r'<[^>]+>', '', p_content).strip()[:30]
            
            # 直接使用索引匹配段落格式
            para_info = None
            if para_idx < len(cell_info['paragraphs']):
                para_info = cell_info['paragraphs'][para_idx]
            elif len(cell_info['paragraphs']) > 0:
                # 如果 HTML 段落数多于 DOCX 段落数，使用最后一个段落的格式
                para_info = cell_info['paragraphs'][-1]
            
            if para_info:
                p_styles = []
                if 'style=' in p_attrs:
                    style_match = re.search(r'style="([^"]*)"', p_attrs)
                    if style_match:
                        existing_style = style_match.group(1)
                        p_styles.append(existing_style)
                        p_attrs = re.sub(r'style="[^"]*"', '', p_attrs)
                
                # 添加段落对齐样式
                if para_info['alignment']:
                    p_styles = [s for s in p_styles if 'text-align' not in s]
                    p_styles.append(f"text-align: {para_info['alignment']}")
                
                # 添加段落缩进样式（关键：必须应用）
                if para_info['indent_style']:
                    p_styles = [s for s in p_styles if 'text-indent' not in s and 'padding-left' not in s]
                    p_styles.append(para_info['indent_style'])
                    logger.info(
                        f"✓ 应用缩进样式到单元格[{table_idx},{row_idx},{col_idx}]段落[{para_idx}]: "
                        f"{para_info['indent_style']}, 文本: '{clean_text}...'"
                    )
                else:
                    logger.warning(
                        f"✗ 单元格[{table_idx},{row_idx},{col_idx}]段落[{para_idx}]无缩进样式, "
                        f"文本: '{clean_text}...'"
                    )
                
                if p_styles:
                    p_style_str = "; ".join(p_styles)
                    updated_paragraphs.append(f'<{p_tag}{p_attrs} style="{p_style_str}">{p_content}</{p_tag}>')
                else:
                    updated_paragraphs.append(para_match.group(0))
            else:
                logger.warning(f"单元格[{table_idx},{row_idx},{col_idx}]段落[{para_idx}]未找到匹配的格式信息")
                updated_paragraphs.append(para_match.group(0))
        
        # 如果单元格内容中没有 p 标签，但需要应用格式，则添加 p 标签
        if not para_elements and cell_info['paragraphs']:
            logger.info(f"单元格[{table_idx},{row_idx},{col_idx}]内容中没有p标签，添加p标签")
            first_para = cell_info['paragraphs'][0]
            p_styles = []
            
            if first_para['alignment']:
                p_styles.append(f"text-align: {first_para['alignment']}")
            
            if first_para['indent_style']:
                p_styles.append(first_para['indent_style'])
                logger.info(f"✓ 为无p标签的单元格添加缩进样式: {first_para['indent_style']}")
            
            if p_styles:
                p_style_str = "; ".join(p_styles)
                cell_content = f'<p style="{p_style_str}">{cell_content}</p>'
        elif para_elements:
            # 替换所有段落（从后往前替换，避免索引问题）
            for para_idx in range(len(para_elements) - 1, -1, -1):
                para_match = para_elements[para_idx]
                if para_idx < len(updated_paragraphs):
                    new_para = updated_paragraphs[para_idx]
                    cell_content = cell_content[:para_match.start()] + new_para + cell_content[para_match.end():]
        
        # 构建新的单元格 HTML
        cell_style_str = "; ".join(cell_styles)
        new_cell_html = f'<{tag}{cell_attrs} style="{cell_style_str}">{cell_content}</{tag}>'
        
        # 计算在原始 HTML 中的绝对位置
        # row_match 和 cell_match 是相对于 table_html 的，需要加上 table_match.start()
        cell_absolute_start = table_match.start() + (row_match.start() - 0) + (cell_match.start() - 0)
        cell_absolute_end = table_match.start() + (row_match.start() - 0) + (cell_match.end() - 0)
        
        # 替换原始 HTML 中的单元格
        html = html[:cell_absolute_start] + new_cell_html + html[cell_absolute_end:]
    
    # 4. 为没有 style 的 td/th 添加默认边框（如果还没有被处理）
    html = re.sub(
        r'<td(?!\s+style)([^>]*)>',
        r'<td\1 style="border: 1px solid #000; padding: 4pt 8pt; vertical-align: top; text-align: left;">',
        html
    )
    html = re.sub(
        r'<th(?!\s+style)([^>]*)>',
        r'<th\1 style="border: 1px solid #000; padding: 4pt 8pt; vertical-align: top; text-align: center;">',
        html
    )
    
    return html


def extract_placeholders(html_content: str) -> Set[str]:
    """
    从 HTML 内容中提取所有 {{field_name}} 格式的占位符
    
    Args:
        html_content: HTML 格式的模板内容
        
    Returns:
        占位符名称集合（不包含 {{ 和 }}）
        
    Example:
        >>> html = "Hello {{name}}, your age is {{age}}"
        >>> extract_placeholders(html)
        {'name', 'age'}
    """
    if not html_content:
        return set()
    
    # 匹配 {{field_name}} 格式的占位符
    # 允许占位符名称包含字母、数字、下划线
    pattern = r'\{\{([a-zA-Z_][a-zA-Z0-9_]*)\}\}'
    matches = re.findall(pattern, html_content)
    
    # 返回去重后的占位符名称集合
    return set(matches)


def docx_to_html(docx_path: str | Path) -> str:
    """
    使用 mammoth-python 将 DOCX 文件转换为 HTML，保留样式信息
    
    Args:
        docx_path: DOCX 文件路径
        
    Returns:
        HTML 格式的字符串
        
    Raises:
        FileNotFoundError: 如果文件不存在
        Exception: 如果转换失败
    """
    docx_path = Path(docx_path)
    
    if not docx_path.exists():
        raise FileNotFoundError(f"DOCX 文件不存在: {docx_path}")
    
    try:
        with open(docx_path, "rb") as docx_file:
            # 使用样式映射来保留更多样式信息
            style_map = """
            p[style-name='Title'] => h1.title:fresh
            p[style-name='Heading 1'] => h1.heading1:fresh
            p[style-name='Heading 2'] => h2.heading2:fresh
            p[style-name='Heading 3'] => h3.heading3:fresh
            p[style-name='Heading 4'] => h4.heading4:fresh
            p[style-name='Heading 5'] => h5.heading5:fresh
            p[style-name='Heading 6'] => h6.heading6:fresh
            p[style-name='标题 1'] => h1.heading1:fresh
            p[style-name='标题 2'] => h2.heading2:fresh
            p[style-name='标题 3'] => h3.heading3:fresh
            p[style-name='标题'] => h1.title:fresh
            r[style-name='Strong'] => strong
            p[style-name='Normal'] => p.paragraph:fresh
            """
            
            result = mammoth.convert_to_html(
                docx_file,
                style_map=style_map,
                include_default_style_map=True  # 包含默认样式映射
            )
            html = result.value
            
            # 读取 DOCX 文件以获取对齐信息
            doc = Document(str(docx_path))
            
            # 后处理 HTML，注入对齐信息
            html = _inject_alignment_info(html, doc)
            
            # 记录警告（如果有）
            if result.messages:
                for message in result.messages:
                    logger.warning(f"DOCX 转换警告: {message}")
            
            return html
    except Exception as e:
        logger.error(f"DOCX 转 HTML 失败: {e}")
        raise


def docx_bytes_to_html(docx_bytes: bytes) -> str:
    """
    将 DOCX 字节数据转换为 HTML，保留样式信息
    
    Args:
        docx_bytes: DOCX 文件的字节数据
        
    Returns:
        HTML 格式的字符串
        
    Raises:
        Exception: 如果转换失败
    """
    try:
        docx_file = io.BytesIO(docx_bytes)
        # 使用样式映射来保留更多样式信息
        style_map = """
        p[style-name='Title'] => h1.title:fresh
        p[style-name='Heading 1'] => h1.heading1:fresh
        p[style-name='Heading 2'] => h2.heading2:fresh
        p[style-name='Heading 3'] => h3.heading3:fresh
        p[style-name='Heading 4'] => h4.heading4:fresh
        p[style-name='Heading 5'] => h5.heading5:fresh
        p[style-name='Heading 6'] => h6.heading6:fresh
        p[style-name='标题 1'] => h1.heading1:fresh
        p[style-name='标题 2'] => h2.heading2:fresh
        p[style-name='标题 3'] => h3.heading3:fresh
        p[style-name='标题'] => h1.title:fresh
        r[style-name='Strong'] => strong
        p[style-name='Normal'] => p.paragraph:fresh
        """
        
        result = mammoth.convert_to_html(
            docx_file,
            style_map=style_map,
            include_default_style_map=True  # 包含默认样式映射
        )
        html = result.value
        
        # 读取 DOCX 文件以获取对齐信息
        docx_file.seek(0)  # 重置文件指针
        doc = Document(docx_file)
        
        # 后处理 HTML，注入对齐信息
        html = _inject_alignment_info(html, doc)
        
        # 记录警告（如果有）
        if result.messages:
            for message in result.messages:
                logger.warning(f"DOCX 转换警告: {message}")
        
        return html
    except Exception as e:
        logger.error(f"DOCX 字节转 HTML 失败: {e}")
        raise


def update_docx_content_from_html(docx_bytes: bytes, html_content: str) -> bytes:
    """
    更新现有DOCX文件的内容，同时保留格式结构（表格、对齐、缩进等）
    
    这是保存编辑时的关键函数，确保格式不会丢失。
    
    Args:
        docx_bytes: 现有DOCX文件的字节数据
        html_content: 编辑后的HTML内容
        
    Returns:
        更新后的DOCX文件的字节数据
    """
    try:
        from docx import Document
        from docx.oxml.ns import qn
        import re
        
        # 读取现有DOCX文件
        docx_file = io.BytesIO(docx_bytes)
        doc = Document(docx_file)
        
        # 从HTML中提取文本内容（按段落和表格单元格）
        # 解析HTML，提取所有段落和表格单元格的文本
        html_paragraphs = []
        html_tables = []
        
        # 提取文档级别的段落和标题（不在表格内的）
        # 先提取表格，然后提取不在表格内的段落和标题
        table_pattern = r'<table[^>]*>.*?</table>'
        tables_html = list(re.finditer(table_pattern, html_content, flags=re.DOTALL))
        
        # 提取不在表格内的段落和标题（按顺序，包含格式信息）
        from app.lex_docx.format_mapper import FormatMapper
        
        last_pos = 0
        for table_match in tables_html:
            # 提取表格前的内容（包括段落和标题）
            before_table = html_content[last_pos:table_match.start()]
            # 提取标题（h1-h6）
            heading_pattern = r'<(h[1-6])([^>]*)>(.*?)</\1>'
            for match in re.finditer(heading_pattern, before_table, flags=re.DOTALL):
                heading_tag = match.group(1)
                heading_attrs = match.group(2)
                heading_text = re.sub(r'<[^>]+>', '', match.group(3))
                heading_text = heading_text.replace('&nbsp;', ' ').strip()
                if heading_text:
                    format_info = FormatMapper.extract_paragraph_format_from_html(f'<{heading_tag} {heading_attrs}>')
                    html_paragraphs.append({
                        'text': heading_text,
                        'format': format_info
                    })
            # 提取段落
            para_pattern = r'<p([^>]*)>(.*?)</p>'
            for match in re.finditer(para_pattern, before_table, flags=re.DOTALL):
                para_attrs = match.group(1)
                para_text = re.sub(r'<[^>]+>', '', match.group(2))
                para_text = para_text.replace('&nbsp;', ' ').strip()
                if para_text:
                    format_info = FormatMapper.extract_paragraph_format_from_html(f'<p {para_attrs}>')
                    html_paragraphs.append({
                        'text': para_text,
                        'format': format_info
                    })
            last_pos = table_match.end()
        
        # 提取最后一个表格后的段落和标题
        if last_pos < len(html_content):
            after_tables = html_content[last_pos:]
            # 提取标题（h1-h6）
            heading_pattern = r'<(h[1-6])([^>]*)>(.*?)</\1>'
            for match in re.finditer(heading_pattern, after_tables, flags=re.DOTALL):
                heading_tag = match.group(1)
                heading_attrs = match.group(2)
                heading_text = re.sub(r'<[^>]+>', '', match.group(3))
                heading_text = heading_text.replace('&nbsp;', ' ').strip()
                if heading_text:
                    format_info = FormatMapper.extract_paragraph_format_from_html(f'<{heading_tag} {heading_attrs}>')
                    html_paragraphs.append({
                        'text': heading_text,
                        'format': format_info
                    })
            # 提取段落
            para_pattern = r'<p([^>]*)>(.*?)</p>'
            for match in re.finditer(para_pattern, after_tables, flags=re.DOTALL):
                para_attrs = match.group(1)
                para_text = re.sub(r'<[^>]+>', '', match.group(2))
                para_text = para_text.replace('&nbsp;', ' ').strip()
                if para_text:
                    format_info = FormatMapper.extract_paragraph_format_from_html(f'<p {para_attrs}>')
                    html_paragraphs.append({
                        'text': para_text,
                        'format': format_info
                    })
        
        # 提取表格内容（保留单元格内的段落结构）
        table_pattern = r'<table[^>]*>(.*?)</table>'
        for table_match in re.finditer(table_pattern, html_content, flags=re.DOTALL):
            table_html = table_match.group(1)
            table_rows = []
            row_pattern = r'<tr[^>]*>(.*?)</tr>'
            for row_match in re.finditer(row_pattern, table_html, flags=re.DOTALL):
                row_html = row_match.group(1)
                row_cells = []
                cell_pattern = r'<(td|th)[^>]*>(.*?)</\1>'
                for cell_match in re.finditer(cell_pattern, row_html, flags=re.DOTALL):
                    cell_content = cell_match.group(2)
                    # 提取单元格内的所有段落（包含格式信息）
                    cell_paragraphs = []
                    para_pattern = r'<p([^>]*)>(.*?)</p>'
                    for para_match in re.finditer(para_pattern, cell_content, flags=re.DOTALL):
                        para_attrs = para_match.group(1)
                        para_html_content = para_match.group(2)
                        
                        # 提取格式信息
                        format_info = {}
                        style_match = re.search(r'style="([^"]*)"', para_attrs)
                        if style_match:
                            from app.lex_docx.format_mapper import FormatMapper
                            format_info = FormatMapper.extract_paragraph_format_from_html(f'<p {para_attrs}>')
                        
                        # 提取文本内容（处理换行）
                        # 先处理 <br> 标签，转换为换行标记
                        para_text = para_html_content.replace('<br>', '\n').replace('<br/>', '\n').replace('<br />', '\n')
                        # 移除其他HTML标签，但保留换行
                        para_text = re.sub(r'<[^>]+>', '', para_text)
                        para_text = para_text.replace('&nbsp;', ' ')
                        
                        # 如果包含换行，分割为多个段落
                        if '\n' in para_text:
                            lines = para_text.split('\n')
                            for line in lines:
                                line = line.strip()
                                if line:
                                    cell_paragraphs.append({
                                        'text': line,
                                        'format': format_info.copy()  # 每个段落都保留格式信息
                                    })
                        else:
                            para_text = para_text.strip()
                            if para_text:
                                cell_paragraphs.append({
                                    'text': para_text,
                                    'format': format_info
                                })
                    
                    # 如果没有段落，提取纯文本
                    if not cell_paragraphs:
                        cell_text = re.sub(r'<[^>]+>', '', cell_content)
                        cell_text = cell_text.replace('&nbsp;', ' ').strip()
                        if cell_text:
                            cell_paragraphs.append({
                                'text': cell_text,
                                'format': {}
                            })
                    row_cells.append(cell_paragraphs)
                if row_cells:
                    table_rows.append(row_cells)
            if table_rows:
                html_tables.append(table_rows)
        
        # 更新文档级别的段落
        # 如果HTML段落数少于DOCX段落数，只删除空段落，保留有内容的段落
        # 如果HTML段落数多于DOCX段落数，添加新段落（使用第一个段落的格式）
        para_idx = 0
        doc_paragraphs = list(doc.paragraphs)  # 转换为列表，避免迭代时修改
        
        # 使用统一的格式映射器
        from app.lex_docx.format_mapper import FormatMapper
        
        # 更新现有段落
        for para in doc_paragraphs:
            if para_idx < len(html_paragraphs):
                # 获取HTML段落信息（可能是字符串或字典）
                html_para_info = html_paragraphs[para_idx]
                if isinstance(html_para_info, dict):
                    text = html_para_info['text']
                    format_info = html_para_info.get('format', {})
                else:
                    # 兼容旧格式（字符串）
                    text = html_para_info
                    format_info = {}
                
                # 保留格式，只更新文本
                # 清除现有runs（但保留段落样式，如Heading 1, Heading 2等）
                original_style = para.style
                original_style_name = original_style.name if original_style else None
                para.clear()
                # 恢复段落样式（如果是标题样式，需要保留）
                if original_style and original_style_name and original_style_name.startswith('Heading'):
                    para.style = original_style
                
                # 应用格式信息（对齐、缩进等）
                FormatMapper.apply_format_to_docx_paragraph(para, format_info)
                
                # 处理占位符和文本（保留原始格式，不修改颜色）
                parts = re.split(r'(\{\{[a-zA-Z_][a-zA-Z0-9_]*\}\})', text)
                for part in parts:
                    if not part:
                        continue
                    # 占位符和普通文本都直接添加，保留原始格式
                    para.add_run(part)
                para_idx += 1
            else:
                # HTML中没有对应段落
                # 只删除空段落，保留有内容的段落（可能是标题或其他重要内容）
                if not para.text.strip():
                    # 空段落，可以删除
                    para_element = para._element
                    para_element.getparent().remove(para_element)
                else:
                    # 有内容的段落，保留（可能是标题或其他重要内容）
                    # 但清空内容，避免显示旧内容
                    para.clear()
        
        # 如果HTML段落数多于DOCX段落数，添加新段落
        if len(html_paragraphs) > len(doc_paragraphs):
            # 使用第一个段落的格式作为新段落的格式
            base_para = doc_paragraphs[0] if doc_paragraphs else None
            for extra_idx in range(len(doc_paragraphs), len(html_paragraphs)):
                # 获取HTML段落信息
                html_para_info = html_paragraphs[extra_idx]
                if isinstance(html_para_info, dict):
                    text = html_para_info['text']
                    format_info = html_para_info.get('format', {})
                else:
                    text = html_para_info
                    format_info = {}
                
                # 创建新段落
                if base_para:
                    # 复制段落格式
                    new_para = doc.add_paragraph()
                    # 先复制基础格式
                    if base_para.alignment is not None:
                        new_para.alignment = base_para.alignment
                    if base_para.paragraph_format.left_indent is not None:
                        new_para.paragraph_format.left_indent = base_para.paragraph_format.left_indent
                    if base_para.paragraph_format.first_line_indent is not None:
                        new_para.paragraph_format.first_line_indent = base_para.paragraph_format.first_line_indent
                else:
                    new_para = doc.add_paragraph()
                
                # 应用HTML中的格式信息（覆盖基础格式）
                FormatMapper.apply_format_to_docx_paragraph(new_para, format_info)
                
                # 处理占位符和文本（保留原始格式，不修改颜色）
                parts = re.split(r'(\{\{[a-zA-Z_][a-zA-Z0-9_]*\}\})', text)
                for part in parts:
                    if not part:
                        continue
                    # 占位符和普通文本都直接添加，保留原始格式
                    new_para.add_run(part)
        
        # 更新表格内容（保留格式结构）
        table_idx = 0
        for table in doc.tables:
            if table_idx < len(html_tables):
                html_table = html_tables[table_idx]
                row_idx = 0
                for row in table.rows:
                    if row_idx < len(html_table):
                        html_row = html_table[row_idx]
                        col_idx = 0
                        for cell in row.cells:
                            if col_idx < len(html_row):
                                # 更新单元格内容，保留格式
                                html_cell_paragraphs = html_row[col_idx]  # 这是一个段落列表
                                
                                # 保留现有段落格式，更新内容
                                # 如果HTML段落数少于DOCX段落数，保留多余的段落（可能包含格式信息）
                                # 如果HTML段落数多于DOCX段落数，添加新段落（使用第一个段落的格式）
                                
                                # 使用统一的格式映射器
                                from app.lex_docx.format_mapper import FormatMapper
                                
                                # 先更新现有段落
                                # 注意：需要先收集要删除的段落，避免在迭代时修改列表
                                para_idx = 0
                                paragraphs_to_remove = []
                                cell_paragraphs_list = list(cell.paragraphs)  # 转换为列表，避免迭代时修改
                                for para in cell_paragraphs_list:
                                    if para_idx < len(html_cell_paragraphs):
                                        # 获取HTML段落信息（可能是字符串或字典）
                                        html_para_info = html_cell_paragraphs[para_idx]
                                        if isinstance(html_para_info, dict):
                                            text = html_para_info['text']
                                            format_info = html_para_info.get('format', {})
                                        else:
                                            # 兼容旧格式（字符串）
                                            text = html_para_info
                                            format_info = {}
                                        
                                        # 保留格式，只更新文本
                                        para.clear()
                                        
                                        # 应用格式信息（对齐、缩进等）
                                        FormatMapper.apply_format_to_docx_paragraph(para, format_info)
                                        
                                        # 处理占位符和文本（保留原始格式，不修改颜色）
                                        parts = re.split(r'(\{\{[a-zA-Z_][a-zA-Z0-9_]*\}\})', text)
                                        for part in parts:
                                            if not part:
                                                continue
                                            # 占位符和普通文本都直接添加，保留原始格式
                                            para.add_run(part)
                                        para_idx += 1
                                    else:
                                        # HTML中没有对应段落，标记为删除（用户已删除）
                                        paragraphs_to_remove.append(para)
                                
                                # 删除标记的段落
                                for para in paragraphs_to_remove:
                                    para_element = para._element
                                    para_element.getparent().remove(para_element)
                                
                                # 如果HTML段落数多于DOCX段落数，添加新段落
                                if len(html_cell_paragraphs) > len(cell.paragraphs):
                                    # 使用第一个段落的格式作为新段落的格式
                                    base_para = cell.paragraphs[0] if cell.paragraphs else None
                                    for extra_idx in range(len(cell.paragraphs), len(html_cell_paragraphs)):
                                        # 获取HTML段落信息
                                        html_para_info = html_cell_paragraphs[extra_idx]
                                        if isinstance(html_para_info, dict):
                                            text = html_para_info['text']
                                            format_info = html_para_info.get('format', {})
                                        else:
                                            text = html_para_info
                                            format_info = {}
                                        
                                        # 创建新段落
                                        if base_para:
                                            # 复制段落格式
                                            new_para = cell.add_paragraph()
                                            # 先复制基础格式
                                            if base_para.alignment is not None:
                                                new_para.alignment = base_para.alignment
                                            if base_para.paragraph_format.left_indent is not None:
                                                new_para.paragraph_format.left_indent = base_para.paragraph_format.left_indent
                                            if base_para.paragraph_format.first_line_indent is not None:
                                                new_para.paragraph_format.first_line_indent = base_para.paragraph_format.first_line_indent
                                        else:
                                            new_para = cell.add_paragraph()
                                        
                                        # 应用HTML中的格式信息（覆盖基础格式）
                                        FormatMapper.apply_format_to_docx_paragraph(new_para, format_info)
                                        
                                        # 处理占位符和文本（保留原始格式，不修改颜色）
                                        parts = re.split(r'(\{\{[a-zA-Z_][a-zA-Z0-9_]*\}\})', text)
                                        for part in parts:
                                            if not part:
                                                continue
                                            # 占位符和普通文本都直接添加，保留原始格式
                                            new_para.add_run(part)
                                col_idx += 1
                    row_idx += 1
                table_idx += 1
        
        # 保存更新后的DOCX
        docx_bytes_updated = io.BytesIO()
        doc.save(docx_bytes_updated)
        docx_bytes_updated.seek(0)
        
        return docx_bytes_updated.read()
    except Exception as e:
        logger.error(f"更新DOCX内容失败: {e}")
        raise


def html_to_docx(html_content: str) -> bytes:
    """
    将 HTML 内容转换为 DOCX 格式的字节数据
    
    ⚠️ 废弃警告：此函数已废弃，格式保留能力有限。
    应该使用 update_docx_content_from_html 来更新现有DOCX文件，以保留格式。
    此函数仅用于以下场景：
    1. 手动创建模板时（create_template）
    2. 导入模板失败时的回退方案（import_template）
    3. 导出模板时如果没有DOCX文件（export_template）
    
    注意：这是一个简化实现，主要用于保存模板内容。
    复杂的 HTML 格式可能无法完全保留。
    
    Args:
        html_content: HTML 格式的内容
        
    Returns:
        DOCX 文件的字节数据
        
    Raises:
        Exception: 如果转换失败
    """
    try:
        # 创建新的 Document
        doc = Document()
        
        # 简单的 HTML 解析（处理基本标签）
        # 移除占位符周围的 HTML 标签，保留占位符本身
        # 这是一个简化实现，实际项目中可能需要更复杂的 HTML 解析
        
        # 将 HTML 内容按段落分割
        paragraphs = html_content.split('\n')
        
        for para_text in paragraphs:
            if not para_text.strip():
                # 空行
                doc.add_paragraph()
                continue
            
            # 处理段落中的占位符和文本
            # 提取占位符和普通文本
            parts = re.split(r'(\{\{[a-zA-Z_][a-zA-Z0-9_]*\}\})', para_text)
            
            paragraph = doc.add_paragraph()
            
            for part in parts:
                if not part:
                    continue
                
                # 检查是否是占位符
                if re.match(r'\{\{[a-zA-Z_][a-zA-Z0-9_]*\}\}', part):
                    # 占位符：添加为普通文本（保持 {{field_name}} 格式，不修改颜色）
                    paragraph.add_run(part)
                else:
                    # 普通文本
                    # 简单的 HTML 标签处理
                    text = re.sub(r'<[^>]+>', '', part)  # 移除 HTML 标签
                    text = text.replace('&nbsp;', ' ')
                    text = text.replace('&lt;', '<')
                    text = text.replace('&gt;', '>')
                    text = text.replace('&amp;', '&')
                    
                    if text.strip():
                        paragraph.add_run(text)
        
        # 将 Document 转换为字节数据
        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        
        return docx_bytes.read()
    except Exception as e:
        logger.error(f"HTML 转 DOCX 失败: {e}")
        raise


def parse_placeholder_metadata(
    html_content: str,
    existing_metadata: Optional[Dict[str, PlaceholderMetadata]] = None
) -> Dict[str, PlaceholderMetadata]:
    """
    从模板内容中解析并构建占位符元数据
    
    如果占位符在 existing_metadata 中已存在，则保留现有配置。
    如果占位符不存在，则创建默认配置。
    
    Args:
        html_content: HTML 格式的模板内容
        existing_metadata: 现有的占位符元数据（可选）
        
    Returns:
        占位符元数据字典
    """
    if existing_metadata is None:
        existing_metadata = {}
    
    # 提取所有占位符
    placeholders = extract_placeholders(html_content)
    
    # 构建元数据字典
    metadata: Dict[str, PlaceholderMetadata] = {}
    
    for placeholder_name in placeholders:
        # 验证占位符名称
        try:
            validated_name = validate_placeholder_name(placeholder_name)
        except ValueError as e:
            logger.warning(f"跳过无效的占位符名称 '{placeholder_name}': {e}")
            continue
        
        # 如果已存在元数据，使用现有配置
        if validated_name in existing_metadata:
            metadata[validated_name] = existing_metadata[validated_name]
        else:
            # 创建默认元数据
            metadata[validated_name] = PlaceholderMetadata(
                type='text',  # 默认类型为文本
                label=validated_name.replace('_', ' ').title(),  # 从名称生成标签
                required=False,
                default_value=None
            )
    
    return metadata


def validate_template_content(html_content: str) -> tuple[bool, Optional[str]]:
    """
    验证模板内容是否有效
    
    Args:
        html_content: HTML 格式的模板内容
        
    Returns:
        (是否有效, 错误信息)
    """
    if not html_content or not html_content.strip():
        return False, "模板内容不能为空"
    
    # 提取所有可能的占位符（包括无效的）
    # 使用更宽松的正则表达式来匹配所有 {{...}} 格式
    pattern = r'\{\{([^}]+)\}\}'
    all_placeholders = re.findall(pattern, html_content)
    
    # 验证每个占位符名称
    for placeholder in all_placeholders:
        try:
            validate_placeholder_name(placeholder)
        except ValueError as e:
            return False, f"占位符 '{{{{ {placeholder} }}}}' 无效: {e}"
    
    return True, None


def _chinese_to_placeholder_name(chinese_text: str) -> str:
    """
    将中文标签转换为占位符名称（英文）
    
    使用简单的拼音映射和规则：
    - 移除标点符号
    - 将常见中文词汇映射为英文
    - 使用下划线连接多个词
    
    Args:
        chinese_text: 中文标签文本
        
    Returns:
        占位符名称（英文，符合命名规范）
    """
    # 常见中文到英文的映射
    common_mappings = {
        '姓名': 'name',
        '性别': 'gender',
        '年龄': 'age',
        '出生日期': 'birthday',
        '民族': 'nation',
        '住所地': 'address',
        '经常居住地': 'current_residence',
        '身份证号': 'id_card',
        '联系电话': 'contact_phone',
        '联系方式': 'contact',
        '邮箱': 'email',
        '地址': 'address',
        '邮政编码': 'postal_code',
        '名称': 'name',
        '统一社会信用代码': 'unified_social_credit_code',
        '法定代表人': 'legal_representative',
        '负责人': 'representative',
        '职务': 'position',
        '工作单位': 'employer',
        '证件类型': 'id_type',
        '证件号码': 'id_number',
        '注册地': 'registration_place',
        '登记地': 'registration_place',
    }
    
    # 清理文本：移除标点符号和空格
    cleaned = re.sub(r'[：:()（）【】\[\]（）]', '', chinese_text)
    cleaned = cleaned.strip()
    
    # 如果完全匹配，直接返回
    if cleaned in common_mappings:
        return common_mappings[cleaned]
    
    # 尝试部分匹配（如果包含常见词汇）
    for key, value in common_mappings.items():
        if key in cleaned:
            # 提取前缀（如果有）
            prefix = ''
            if '原告' in cleaned:
                prefix = 'plaintiff_'
            elif '被告' in cleaned:
                prefix = 'defendant_'
            elif '申请人' in cleaned:
                prefix = 'applicant_'
            elif '被申请人' in cleaned:
                prefix = 'respondent_'
            elif '第三人' in cleaned:
                prefix = 'third_party_'
            
            # 移除前缀部分
            remaining = cleaned.replace('原告', '').replace('被告', '').replace('申请人', '').replace('被申请人', '').replace('第三人', '')
            remaining = remaining.strip()
            
            if remaining in common_mappings:
                return prefix + common_mappings[remaining]
            elif remaining:
                # 如果还有剩余文本，尝试生成名称
                # 简单处理：使用拼音首字母或直接使用value
                return prefix + value
    
    # 如果没有匹配，生成一个简单的名称
    # 移除所有非字母数字字符，用下划线替换
    # 注意：这里不能保留中文字符，必须全部转换为英文
    name = re.sub(r'[^\w]', '_', cleaned)
    
    # 如果包含中文字符，需要移除（因为占位符名称必须是英文）
    # 使用正则表达式移除所有中文字符
    name = re.sub(r'[\u4e00-\u9fff]', '', name)
    
    # 如果名称为空或只包含下划线，使用默认名称
    if not name or name == '_' * len(name) or not name.strip('_'):
        name = 'field'
    
    # 确保以字母或下划线开头
    if not re.match(r'^[a-zA-Z_]', name):
        name = 'field_' + name
    
    # 转换为小写，移除多余下划线
    name = re.sub(r'_+', '_', name).lower().strip('_')
    
    # 最终验证：确保名称只包含字母、数字、下划线
    name = re.sub(r'[^a-zA-Z0-9_]', '', name)
    
    # 如果清理后名称为空，使用默认名称
    if not name:
        name = 'field'
    
    # 再次确保以字母或下划线开头
    if not re.match(r'^[a-zA-Z_]', name):
        name = 'field_' + name
    
    return name


def identify_placeholder_positions(html_content: str) -> List[Dict[str, Any]]:
    """
    识别HTML内容中可能的占位符位置
    
    规则：
    1. 查找包含冒号（：或:）的段落
    2. 冒号之前的部分作为标签（label）
    3. 冒号之后的部分作为内容，需要分析其类型
    
    Args:
        html_content: HTML格式的模板内容
        
    Returns:
        占位符位置列表，每个元素包含：
        - label: 标签文本
        - content: 内容文本
        - placeholder_name: 生成的占位符名称
        - html_match: 匹配的HTML片段（用于替换）
    """
    import re
    from html import unescape
    
    positions = []
    
    # 匹配包含冒号的段落
    # 支持中文冒号（：）和英文冒号（:）
    # 匹配格式：<p>标签：内容</p> 或 <p>标签:内容</p>
    pattern = r'<p[^>]*>([^<]*[：:][^<]*)</p>'
    
    matches = list(re.finditer(pattern, html_content))
    
    for match in matches:
        full_match = match.group(0)
        content_text = match.group(1)
        
        # 提取标签和内容
        # 支持中文冒号和英文冒号
        colon_match = re.search(r'([^：:]+)[：:](.+)', content_text)
        if not colon_match:
            continue
        
        label = colon_match.group(1).strip()
        content = colon_match.group(2).strip()
        
        # 清理HTML实体
        label = unescape(label)
        content = unescape(content)
        
        # 移除HTML标签（如果有）
        label = re.sub(r'<[^>]+>', '', label).strip()
        content = re.sub(r'<[^>]+>', '', content).strip()
        
        # 跳过空标签
        if not label:
            continue
        
        # 过滤掉不应该识别为占位符的内容
        # 1. 法律条文引用（包含"第X条"、"规定"等）
        if re.search(r'第[一二三四五六七八九十\d]+条|规定|法律|法规|条例', label):
            continue
        # 2. 说明性文字（过长或包含特殊字符）
        if len(label) > 50:  # 标签过长，可能是说明文字
            continue
        # 3. 纯标点符号或特殊字符
        if re.match(r'^[：:()（）【】\[\]★\s]+$', label):
            continue
        
        # 生成占位符名称
        placeholder_name = _chinese_to_placeholder_name(label)
        
        # 验证占位符名称是否有效（必须只包含字母、数字、下划线，且以字母或下划线开头）
        if not re.match(r'^[a-zA-Z_][a-zA-Z0-9_]*$', placeholder_name):
            # 如果生成的名称无效，使用默认名称加序号
            placeholder_name = f"field_{len(positions) + 1}"
        
        # 确保占位符名称唯一（如果重复，添加序号）
        existing_names = {pos['placeholder_name'] for pos in positions}
        base_name = placeholder_name
        counter = 1
        while placeholder_name in existing_names:
            placeholder_name = f"{base_name}_{counter}"
            counter += 1
        
        positions.append({
            'label': label,
            'content': content,
            'placeholder_name': placeholder_name,
            'html_match': full_match,
            'match_start': match.start(),
            'match_end': match.end(),
        })
    
    return positions


def generate_placeholders_from_positions(
    html_content: str,
    positions: List[Dict[str, Any]]
) -> str:
    """
    根据识别的位置生成占位符，替换HTML内容
    
    Args:
        html_content: 原始HTML内容
        positions: 占位符位置列表（从identify_placeholder_positions获取）
        
    Returns:
        包含占位符的HTML内容
    """
    # 从后往前替换，避免索引问题
    result = html_content
    for pos in sorted(positions, key=lambda x: x['match_start'], reverse=True):
        # 构建新的HTML：保留标签部分，用占位符替换内容部分
        placeholder = f"{{{{{pos['placeholder_name']}}}}}"
        
        # 提取原始HTML的标签属性
        original_match = pos['html_match']
        tag_match = re.search(r'<p([^>]*)>', original_match)
        tag_attrs = tag_match.group(1) if tag_match else ''
        
        # 构建新的HTML片段：<p>标签：{{placeholder}}</p>
        label = pos['label']
        new_content = f"<p{tag_attrs}>{label}：{placeholder}</p>"
        
        # 替换
        result = result[:pos['match_start']] + new_content + result[pos['match_end']:]
    
    return result

