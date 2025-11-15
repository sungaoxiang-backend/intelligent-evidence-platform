"""
模板解析和文档转换工具函数的单元测试
"""
import os
import tempfile

import pytest
from docx import Document

from app.lex_docx.utils import (
    docx_bytes_to_html,
    docx_to_html,
    extract_placeholders,
    html_to_docx,
    parse_placeholder_metadata,
    validate_template_content,
)


class TestExtractPlaceholders:
    """测试占位符提取功能"""
    
    def test_extract_simple_placeholders(self):
        """测试提取简单占位符"""
        html = "Hello {{name}}, your age is {{age}}."
        placeholders = extract_placeholders(html)
        assert placeholders == {"name", "age"}
    
    def test_extract_placeholders_with_underscores(self):
        """测试提取带下划线的占位符"""
        html = "User {{user_name}} has {{total_count}} items."
        placeholders = extract_placeholders(html)
        assert placeholders == {"user_name", "total_count"}
    
    def test_extract_no_placeholders(self):
        """测试没有占位符的情况"""
        html = "This is plain text without placeholders."
        placeholders = extract_placeholders(html)
        assert placeholders == set()
    
    def test_extract_duplicate_placeholders(self):
        """测试重复占位符只提取一次"""
        html = "{{name}} and {{name}} are the same."
        placeholders = extract_placeholders(html)
        assert placeholders == {"name"}
    
    def test_extract_placeholders_in_html_tags(self):
        """测试 HTML 标签中的占位符"""
        html = "<p>Hello {{name}}</p><div>{{age}}</div>"
        placeholders = extract_placeholders(html)
        assert placeholders == {"name", "age"}


class TestValidateTemplateContent:
    """测试模板内容验证功能"""
    
    def test_validate_valid_content(self):
        """测试验证有效内容"""
        html = "Hello {{user_name}}, your age is {{age}}."
        is_valid, error = validate_template_content(html)
        assert is_valid
        assert error is None
    
    def test_validate_empty_content(self):
        """测试验证空内容"""
        is_valid, error = validate_template_content("")
        assert not is_valid
        assert error == "模板内容不能为空"
    
    def test_validate_invalid_placeholder_starts_with_number(self):
        """测试验证以数字开头的占位符"""
        html = "Hello {{123invalid}}."
        is_valid, error = validate_template_content(html)
        assert not is_valid
        assert "无效" in error
    
    def test_validate_invalid_placeholder_special_chars(self):
        """测试验证包含特殊字符的占位符"""
        html = "Hello {{user-name}}."
        is_valid, error = validate_template_content(html)
        assert not is_valid


class TestParsePlaceholderMetadata:
    """测试占位符元数据解析功能"""
    
    def test_parse_metadata_without_existing(self):
        """测试解析元数据（无现有配置）"""
        html = "Hello {{user_name}}, your age is {{age}}."
        metadata = parse_placeholder_metadata(html)
        
        assert "user_name" in metadata
        assert "age" in metadata
        assert metadata["user_name"].type == "text"
        assert metadata["user_name"].label == "User Name"
        assert metadata["user_name"].required is False
    
    def test_parse_metadata_with_existing(self):
        """测试解析元数据（有现有配置）"""
        from app.lex_docx.schemas import PlaceholderMetadata
        
        html = "Hello {{user_name}}, your age is {{age}}."
        existing = {
            "user_name": PlaceholderMetadata(
                type="text",
                label="用户名",
                required=True
            )
        }
        metadata = parse_placeholder_metadata(html, existing)
        
        # 现有配置应该保留
        assert metadata["user_name"].label == "用户名"
        assert metadata["user_name"].required is True
        
        # 新占位符应该有默认配置
        assert metadata["age"].type == "text"
        assert metadata["age"].required is False


class TestDocxToHtml:
    """测试 DOCX 转 HTML 功能"""
    
    def test_docx_to_html_simple(self):
        """测试简单 DOCX 转 HTML"""
        # 创建临时 DOCX 文件
        doc = Document()
        doc.add_paragraph("Hello {{name}}, your age is {{age}}.")
        
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            doc.save(tmp.name)
            tmp_path = tmp.name
        
        try:
            html = docx_to_html(tmp_path)
            assert len(html) > 0
            assert "{{name}}" in html or "name" in html.lower()
        finally:
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)
    
    def test_docx_to_html_file_not_found(self):
        """测试文件不存在的情况"""
        with pytest.raises(FileNotFoundError):
            docx_to_html("/nonexistent/file.docx")
    
    def test_docx_bytes_to_html(self):
        """测试 DOCX 字节转 HTML"""
        # 创建 DOCX 文档
        doc = Document()
        doc.add_paragraph("Test content with {{placeholder}}.")
        
        # 保存到字节
        import io
        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        
        html = docx_bytes_to_html(docx_bytes.read())
        assert len(html) > 0


class TestHtmlToDocx:
    """测试 HTML 转 DOCX 功能"""
    
    def test_html_to_docx_simple(self):
        """测试简单 HTML 转 DOCX"""
        html = "<p>Hello {{user_name}}, your age is {{age}}.</p>"
        docx_bytes = html_to_docx(html)
        
        assert len(docx_bytes) > 0
        assert isinstance(docx_bytes, bytes)
    
    def test_html_to_docx_with_placeholders(self):
        """测试包含占位符的 HTML 转 DOCX"""
        html = "<p>Name: {{name}}</p><p>Age: {{age}}</p>"
        docx_bytes = html_to_docx(html)
        
        assert len(docx_bytes) > 0
        
        # 验证生成的 DOCX 可以读取
        import io

        from docx import Document
        doc = Document(io.BytesIO(docx_bytes))
        assert len(doc.paragraphs) > 0
    
    def test_html_to_docx_empty(self):
        """测试空 HTML"""
        html = ""
        docx_bytes = html_to_docx(html)
        assert len(docx_bytes) > 0  # 至少应该生成一个空的 DOCX

