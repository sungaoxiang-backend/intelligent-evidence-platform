"""
文档模板服务
用于加载、解析和管理文书模板

新架构：
- 源模板：Markdown（使用Jinja2模板语法）
- 展示/打印：HTML + CSS（从Markdown渲染）
- 严格版式：DOCX（从Markdown转换，保留作为导出载体）
- 数据对齐：统一占位符命名，与前端Schema/后端数据树一致
"""

import yaml
import uuid
from pathlib import Path
from typing import Dict, Any, List, Optional
from datetime import datetime
import logging

from docxtpl import DocxTemplate
from jinja2 import Environment, FileSystemLoader, Template
import markdown
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)


class DocumentTemplateService:
    """文档模板服务"""
    
    def __init__(self, template_dir: Optional[Path] = None):
        """
        初始化模板服务
        
        Args:
            template_dir: 模板目录路径，默认使用 app/documents_template
        """
        if template_dir is None:
            template_dir = Path(__file__).parent
        self.template_dir = Path(template_dir)
        self.config_file = self.template_dir / "template_schema.yaml"
        self.markdown_template_dir = self.template_dir / "templates"
        self.markdown_template_dir.mkdir(parents=True, exist_ok=True)
        
        # 初始化Jinja2环境
        self.jinja_env = Environment(
            loader=FileSystemLoader(str(self.markdown_template_dir)),
            autoescape=False,
            trim_blocks=True,
            lstrip_blocks=True
        )
        
        self.templates: Dict[str, Any] = {}
        self._load_templates()
    
    def _load_templates(self):
        """加载模板配置"""
        try:
            logger.info(f"尝试加载模板配置文件: {self.config_file}")
            logger.info(f"文件是否存在: {self.config_file.exists()}")
            if self.config_file.exists():
                with open(self.config_file, 'r', encoding='utf-8') as f:
                    config = yaml.safe_load(f)
                    logger.info(f"YAML解析结果类型: {type(config)}")
                    logger.info(f"YAML解析结果键: {list(config.keys()) if isinstance(config, dict) else 'Not a dict'}")
                    self.templates = config.get('templates', {})
                    logger.info(f"成功加载 {len(self.templates)} 个模板")
                    logger.info(f"模板键: {list(self.templates.keys())}")
                    for key, value in self.templates.items():
                        logger.info(f"模板 {key}: {type(value)}, 包含键: {list(value.keys()) if isinstance(value, dict) else 'Not a dict'}")
            else:
                logger.warning(f"模板配置文件不存在: {self.config_file}")
                logger.warning(f"配置文件绝对路径: {self.config_file.resolve()}")
                self.templates = {}
        except yaml.YAMLError as e:
            logger.error(f"YAML解析失败: {str(e)}")
            logger.error(f"错误位置: {e.problem_mark if hasattr(e, 'problem_mark') else '未知'}")
            self.templates = {}
        except Exception as e:
            logger.error(f"加载模板配置失败: {str(e)}", exc_info=True)
            self.templates = {}
    
    def get_template_list(self) -> List[Dict[str, Any]]:
        """
        获取模板列表
        
        Returns:
            模板列表，每个模板包含基本信息
        """
        template_list = []
        # YAML结构是: templates -> "模板名称" -> template_data
        for template_key, template_data in self.templates.items():
            # 如果template_data是字典且包含template_id，则是一个模板
            if isinstance(template_data, dict):
                template_list.append({
                    "template_id": template_data.get("template_id", template_key),
                    "name": template_data.get("name", template_key),
                    "type": template_data.get("type", ""),
                    "category": template_data.get("category", ""),
                    "description": template_data.get("description", ""),
                })
        return template_list
    
    def get_template_detail(self, template_id: str) -> Optional[Dict[str, Any]]:
        """
        获取模板详情（包含完整的表单结构）
        
        Args:
            template_id: 模板ID或模板名称
            
        Returns:
            模板详情，包含所有表单字段定义
        """
        # 先按 template_id 查找
        for key, template_data in self.templates.items():
            if isinstance(template_data, dict):
                # 匹配template_id或模板名称
                if template_data.get("template_id") == template_id or key == template_id:
                    return template_data
        
        # 如果找不到，返回 None
        return None
    
    def get_template_form_schema(self, template_id: str) -> Optional[Dict[str, Any]]:
        """
        获取模板的表单结构（用于前端渲染）
        
        Args:
            template_id: 模板ID
            
        Returns:
            表单结构，包含模板名称、说明、块列表等
        """
        template = self.get_template_detail(template_id)
        if not template:
            return None
        
        # 构建表单结构
        form_schema = {
            "template_id": template.get("template_id", template_id),
            "name": template.get("name", ""),
            "type": template.get("type", ""),
            "category": template.get("category", ""),
            "description": template.get("description", ""),
            "instructions": template.get("instructions", {}),
            "special_notice": template.get("special_notice", {}),
            "blocks": []
        }
        
        # 处理每个块
        for block in template.get("blocks", []):
            block_schema = {
                "block_id": block.get("block_id", ""),
                "title": block.get("title", ""),
                "description": block.get("description", ""),
                "rows": []
            }
            
            # 处理每个行
            for row in block.get("rows", []):
                row_schema = {
                    "row_id": row.get("row_id", ""),
                    "subtitle": row.get("subtitle", ""),
                    "subtitle_width": row.get("subtitle_width", 120),
                    "fields": []
                }
                
                # 处理每个字段
                for field in row.get("fields", []):
                    field_schema = {
                        "field_id": field.get("field_id", ""),
                        "label": field.get("label", ""),
                        "type": field.get("type", "text"),
                        "required": field.get("required", False),
                        "placeholder": field.get("placeholder", ""),
                        "default": field.get("default", ""),
                    }
                    
                    # 添加字段类型特定的属性
                    if field.get("type") == "select" or field.get("type") == "radio" or field.get("type") == "checkbox":
                        field_schema["options"] = field.get("options", [])
                    
                    if field.get("type") == "textarea":
                        field_schema["rows"] = field.get("rows", 3)
                    
                    if field.get("type") == "date" or field.get("type") == "datetime":
                        field_schema["format"] = field.get("format", "YYYY-MM-DD")
                    
                    if field.get("validation"):
                        field_schema["validation"] = field.get("validation")
                    
                    row_schema["fields"].append(field_schema)
                
                block_schema["rows"].append(row_schema)
            
            form_schema["blocks"].append(block_schema)
        
        return form_schema
    
    def validate_form_data(self, template_id: str, form_data: Dict[str, Any]) -> Dict[str, Any]:
        """
        验证表单数据
        
        Args:
            template_id: 模板ID
            form_data: 表单数据
            
        Returns:
            验证结果，包含是否通过和错误信息
        """
        template = self.get_template_detail(template_id)
        if not template:
            return {
                "valid": False,
                "errors": [f"模板 {template_id} 不存在"]
            }
        
        errors = []
        
        # 遍历所有字段进行验证
        for block in template.get("blocks", []):
            for row in block.get("rows", []):
                for field in row.get("fields", []):
                    field_id = field.get("field_id", "")
                    field_value = form_data.get(field_id, "")
                    
                    # 不检查必填字段，允许空值
                    # 移除必填验证，允许空字段生成文档
                    
                    # 字段验证
                    validation = field.get("validation", {})
                    if field_value and validation:
                        # 最小长度验证
                        min_length = validation.get("min_length")
                        if min_length and len(str(field_value)) < min_length:
                            errors.append(f"字段 '{field.get('label', field_id)}' 长度不能少于 {min_length} 个字符")
                        
                        # 最大长度验证
                        max_length = validation.get("max_length")
                        if max_length and len(str(field_value)) > max_length:
                            errors.append(f"字段 '{field.get('label', field_id)}' 长度不能超过 {max_length} 个字符")
                        
                        # 正则表达式验证
                        pattern = validation.get("pattern")
                        if pattern:
                            import re
                            if not re.match(pattern, str(field_value)):
                                errors.append(f"字段 '{field.get('label', field_id)}' 格式不正确")
        
        return {
            "valid": len(errors) == 0,
            "errors": errors
        }
    
    def generate_document_from_form(self, template_id: str, form_data: Dict[str, Any]) -> Dict[str, Any]:
        """
        根据表单数据生成文书
        
        Args:
            template_id: 模板ID
            form_data: 表单数据
            
        Returns:
            生成结果，包含文件路径
        """
        # 不进行必填验证，允许空字段生成文档
        template = self.get_template_detail(template_id)
        if not template:
            return {
                "success": False,
                "message": f"模板 {template_id} 不存在"
            }
        
        try:
            # 准备变量数据（将表单数据转换为模板变量）
            variables = self._prepare_template_variables(form_data, template)
            
            # 添加模板的instructions和special_notice到variables
            # 确保items是列表而不是方法
            if template.get("instructions"):
                instructions = template.get("instructions")
                # 如果是字典，创建副本并确保items是列表
                if isinstance(instructions, dict):
                    instructions = instructions.copy()
                    # 如果items存在且不是列表，转换为列表
                    if "items" in instructions:
                        items = instructions["items"]
                        if not isinstance(items, list):
                            # 如果不是列表，尝试转换
                            if hasattr(items, '__iter__') and not isinstance(items, str):
                                instructions["items"] = list(items)
                            else:
                                instructions["items"] = []
                    else:
                        instructions["items"] = []
                variables["instructions"] = instructions
            if template.get("special_notice"):
                variables["special_notice"] = template.get("special_notice")
            
            # 优先使用DOCX模板（保持原有格式结构，使用docxtpl填充占位符）
            template_file_path = template.get("file_path", "")
            if template_file_path:
                # 处理相对路径
                template_path = Path(template_file_path)
                if not template_path.is_absolute():
                    # 相对于项目根目录或template_dir
                    if (self.template_dir / template_file_path).exists():
                        template_path = self.template_dir / template_file_path
                    elif Path(template_file_path).exists():
                        template_path = Path(template_file_path)
                
                if template_path.exists():
                    logger.info(f"使用DOCX模板: {template_path}")
                    # 使用docxtpl加载模板
                    doc = DocxTemplate(str(template_path))
                    doc.render(variables)
                    
                    # 创建输出目录
                    output_dir = Path("static/documents")
                    output_dir.mkdir(parents=True, exist_ok=True)
                    
                    # 生成输出文件名
                    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                    output_filename = f"{template_id}_{timestamp}_{uuid.uuid4().hex[:8]}.docx"
                    output_path = output_dir / output_filename
                    
                    # 保存文档
                    doc.save(str(output_path))
                    
                    logger.info(f"文书生成成功: {output_path}")
                    
                    return {
                        "success": True,
                        "message": "文书生成成功",
                        "file_path": str(output_path),
                        "filename": output_filename,
                        "template_id": template_id
                    }
            
            return {
                "success": False,
                "message": "模板文件不存在（既无Markdown模板也无DOCX模板）"
            }
            
        except Exception as e:
            logger.error(f"生成文书失败: {str(e)}", exc_info=True)
            return {
                "success": False,
                "message": f"生成文书失败: {str(e)}"
            }
    
    def _prepare_template_variables(self, form_data: Dict[str, Any], template: Dict[str, Any]) -> Dict[str, Any]:
        """
        准备模板变量，将表单数据转换为DOCX模板可以使用的变量格式
        
        Args:
            form_data: 表单数据
            template: 模板配置
            
        Returns:
            模板变量字典
        """
        variables = {}
        
        # 直接将所有字段ID作为变量名（使用field_id）
        # 这样DOCX模板中可以使用 {{field_id}} 来引用
        for block in template.get("blocks", []):
            for row in block.get("rows", []):
                for field in row.get("fields", []):
                    field_id = field.get("field_id", "")
                    if field_id:
                        field_value = form_data.get(field_id, "")
                        
                        # 处理不同类型的字段值
                        if isinstance(field_value, list):
                            # 对于checkbox类型，将列表转换为字符串
                            if field.get("type") == "checkbox":
                                # 将选中的值用逗号连接
                                variables[field_id] = "，".join(field_value)
                            else:
                                variables[field_id] = "，".join(str(v) for v in field_value)
                        elif field_value is None:
                            variables[field_id] = ""
                        else:
                            variables[field_id] = str(field_value)
                        
                        # 处理嵌套选项（如"国有"下的"控股"/"参股"）
                        if field.get("type") == "checkbox":
                            for option in field.get("options", []):
                                if option.get("value") in field_value and option.get("sub_options"):
                                    # 获取嵌套选项的值
                                    nested_key = f"{field_id}_{option.get('value')}"
                                    nested_value = form_data.get(nested_key, "")
                                    if nested_value:
                                        variables[nested_key] = str(nested_value)
        
        # 添加一些通用变量
        variables["current_date"] = datetime.now().strftime("%Y年%m月%d日")
        variables["generated_at"] = datetime.now().isoformat()
        
        return variables
    
    def _generate_docx_from_markdown(self, markdown_path: Path, variables: Dict[str, Any], template_id: str) -> Path:
        """
        从Markdown模板生成DOCX文档（新架构）
        
        Args:
            markdown_path: Markdown模板文件路径
            variables: 模板变量
            template_id: 模板ID
            
        Returns:
            生成的DOCX文件路径
        """
        # 1. 使用Jinja2渲染Markdown模板
        template_name = markdown_path.name
        jinja_template = self.jinja_env.get_template(template_name)
        rendered_markdown = jinja_template.render(**variables)
        
        # 2. 将Markdown转换为HTML（用于预览）
        html_content = markdown.markdown(
            rendered_markdown,
            extensions=['tables', 'fenced_code', 'nl2br']
        )
        
        # 3. 创建DOCX文档
        doc = Document()
        
        # 设置默认样式
        style = doc.styles['Normal']
        font = style.font
        font.name = '宋体'
        font.size = Pt(12)
        
        # 设置中文字体
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        # 解析Markdown内容并转换为DOCX
        lines = rendered_markdown.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            
            if not line:
                # 空行，添加段落
                doc.add_paragraph()
                i += 1
                continue
            
            # 处理标题
            if line.startswith('# '):
                # 一级标题
                para = doc.add_heading(line[2:], level=1)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                i += 1
            elif line.startswith('## '):
                # 二级标题
                para = doc.add_heading(line[3:], level=2)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                i += 1
            elif line.startswith('### '):
                # 三级标题
                para = doc.add_heading(line[4:], level=3)
                i += 1
            elif line.startswith('**') and line.endswith('**'):
                # 加粗文本
                text = line[2:-2]
                para = doc.add_paragraph()
                run = para.add_run(text)
                run.bold = True
                i += 1
            elif line.startswith('---'):
                # 分隔线
                doc.add_paragraph('─' * 50)
                i += 1
            else:
                # 普通段落
                # 处理变量占位符（如果还有未替换的）
                para = doc.add_paragraph(line)
                i += 1
        
        # 4. 保存DOCX文档
        output_dir = Path("static/documents")
        output_dir.mkdir(parents=True, exist_ok=True)
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_filename = f"{template_id}_{timestamp}_{uuid.uuid4().hex[:8]}.docx"
        output_path = output_dir / output_filename
        
        doc.save(str(output_path))
        
        logger.info(f"从Markdown模板生成DOCX成功: {output_path}")
        
        return output_path
    
    def generate_html_from_markdown(self, template_id: str, form_data: Dict[str, Any]) -> Optional[str]:
        """
        从Markdown模板生成HTML（用于预览和打印）
        
        Args:
            template_id: 模板ID
            form_data: 表单数据
            
        Returns:
            HTML内容，如果失败返回None
        """
        try:
            template = self.get_template_detail(template_id)
            if not template:
                return None
            
            # 准备变量
            variables = self._prepare_template_variables(form_data, template)
            if template.get("instructions"):
                variables["instructions"] = template.get("instructions")
            if template.get("special_notice"):
                variables["special_notice"] = template.get("special_notice")
            
            # 查找Markdown模板
            markdown_template_path = self.markdown_template_dir / f"{template_id}.md"
            if not markdown_template_path.exists():
                return None
            
            # 使用Jinja2渲染Markdown模板
            template_name = markdown_template_path.name
            jinja_template = self.jinja_env.get_template(template_name)
            rendered_markdown = jinja_template.render(**variables)
            
            # 将Markdown转换为HTML
            html_content = markdown.markdown(
                rendered_markdown,
                extensions=['tables', 'fenced_code', 'nl2br']
            )
            
            # 包装在完整的HTML文档中（添加CSS样式）
            full_html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <style>
        body {{
            font-family: 'SimSun', '宋体', serif;
            line-height: 1.8;
            margin: 2cm;
            font-size: 12pt;
        }}
        h1 {{
            text-align: center;
            font-size: 18pt;
            font-weight: bold;
            margin: 20pt 0;
        }}
        h2 {{
            text-align: center;
            font-size: 16pt;
            font-weight: bold;
            margin: 15pt 0;
        }}
        h3 {{
            font-size: 14pt;
            font-weight: bold;
            margin: 12pt 0;
        }}
        p {{
            margin: 6pt 0;
            text-align: justify;
        }}
        @media print {{
            body {{
                margin: 0;
            }}
        }}
    </style>
</head>
<body>
{html_content}
</body>
</html>"""
            
            return full_html
            
        except Exception as e:
            logger.error(f"生成HTML失败: {str(e)}", exc_info=True)
            return None


# 全局模板服务实例
template_service = DocumentTemplateService()

