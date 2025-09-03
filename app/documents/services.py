"""
文书生成服务
提供Word文档生成、模板管理等功能
"""

import os
import uuid
import yaml
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, Optional, List
import logging

from docxtpl import DocxTemplate
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from sqlalchemy.ext.asyncio import AsyncSession

logger = logging.getLogger(__name__)


class DocumentGenerator:
    """文书生成器"""
    
    def __init__(self):
        self.template_dir = Path("templates/documents")
        self.output_dir = Path("static/documents")
        self.config_file = Path("app/documents/templates.yaml")
        self._ensure_directories()
        self._load_config()
    
    def _ensure_directories(self):
        """确保必要的目录存在"""
        self.template_dir.mkdir(parents=True, exist_ok=True)
        self.output_dir.mkdir(parents=True, exist_ok=True)
    
    def _load_config(self):
        """加载配置文件"""
        try:
            if self.config_file.exists():
                with open(self.config_file, 'r', encoding='utf-8') as f:
                    self.config = yaml.safe_load(f)
            else:
                self.config = self._create_default_config()
                self._save_config()
        except Exception as e:
            logger.error(f"加载配置文件失败: {str(e)}")
            self.config = self._create_default_config()
    
    def _save_config(self):
        """保存配置文件"""
        try:
            with open(self.config_file, 'w', encoding='utf-8') as f:
                yaml.dump(self.config, f, allow_unicode=True, default_flow_style=False)
        except Exception as e:
            logger.error(f"保存配置文件失败: {str(e)}")
    
    def _create_default_config(self) -> Dict[str, Any]:
        """创建默认配置"""
        return {
            "templates": {
                "complaint_template": {
                    "name": "借款纠纷起诉状模板",
                    "type": "complaint",
                    "description": "适用于个人借款纠纷案件的起诉状模板",
                    "file_path": "templates/documents/complaint_template.docx",
                    "variables": [
                        {"name": "creditor_name", "description": "债权人姓名", "required": True, "default": ""},
                        {"name": "debtor_name", "description": "债务人姓名", "required": True, "default": ""},
                        {"name": "case_type", "description": "案件类型", "required": True, "default": "借款纠纷"},
                        {"name": "loan_amount", "description": "借款金额", "required": False, "default": ""}
                    ]
                },
                "evidence_list_template": {
                    "name": "证据材料清单模板",
                    "type": "evidence_list",
                    "description": "适用于整理案件证据材料的清单模板",
                    "file_path": "templates/documents/evidence_list_template.docx",
                    "variables": [
                        {"name": "case_title", "description": "案件标题", "required": True, "default": ""},
                        {"name": "case_id", "description": "案件编号", "required": True, "default": ""}
                    ]
                }
            },
            "variable_mapping": {
                "auto_mapping": True,
                "generated_variables": [
                    "case_title",
                    "current_date"
                ]
            }
        }
    
    def get_available_templates(self) -> List[Dict[str, Any]]:
        """获取可用的模板列表"""
        try:
            templates = []
            for template_id, template_info in self.config.get("templates", {}).items():
                template_info["template_id"] = template_id
                templates.append(template_info)
            return templates
        except Exception as e:
            logger.error(f"获取模板列表失败: {str(e)}")
            return []
    
    def get_template(self, template_id: str) -> Optional[Dict[str, Any]]:
        """获取特定模板"""
        try:
            template = self.config.get("templates", {}).get(template_id)
            if template:
                template["template_id"] = template_id
            return template
        except Exception as e:
            logger.error(f"获取模板失败: {str(e)}")
            return None
    
    async def generate_document_by_case_id(
        self,
        db: AsyncSession,
        template_id: str,
        case_id: int,
        custom_variables: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """
        通过案件ID生成文书
        
        Args:
            db: 数据库会话
            template_id: 模板ID
            case_id: 案件ID
            custom_variables: 自定义变量
            
        Returns:
            生成结果
        """
        try:
            # 导入cases服务
            from app.cases.services import get_by_id as get_case_by_id
            
            # 获取案件信息
            case = await get_case_by_id(db, case_id)
            if not case:
                return {
                    "success": False,
                    "message": f"案件 {case_id} 不存在"
                }
            
            # 将案件模型转换为字典格式
            case_data = self._convert_case_model_to_dict(case)
            
            # 调用原有的生成方法
            return self.generate_document(template_id, case_data, custom_variables)
            
        except Exception as e:
            logger.error(f"通过案件ID生成文书失败: {str(e)}")
            return {
                "success": False,
                "message": f"通过案件ID生成文书失败: {str(e)}"
            }
    
    def _convert_case_model_to_dict(self, case) -> Dict[str, Any]:
        """将案件模型转换为字典格式，用于文书生成"""
        # 分离债权人和债务人信息
        creditor = None
        debtor = None
        
        # 添加调试日志
        logger.info(f"案件 {case.id} 的当事人数量: {len(case.case_parties) if case.case_parties else 0}")
        
        for party in case.case_parties:
            logger.info(f"当事人: {party.party_name}, 角色: {party.party_role}, 类型: {party.party_type}, 法定代表人: {party.owner_name}")
            if party.party_role == "creditor":
                creditor = party
            elif party.party_role == "debtor":
                debtor = party
        
        # 构建基础案件信息
        case_data = {
            "case_id": case.id,
            "case_type": case.case_type.value if hasattr(case.case_type, 'value') else str(case.case_type) if case.case_type else "DEBT",
            "case_status": case.case_status.value if hasattr(case.case_status, 'value') else str(case.case_status) if case.case_status else "DRAFT",
            "loan_amount": case.loan_amount or 0.0,
            "loan_date": case.loan_date,
            "court_name": case.court_name or "",
            "description": case.description or "",
            "created_at": case.created_at if hasattr(case, 'created_at') else datetime.now()
        }
        
        # 添加债权人信息
        if creditor:
            case_data.update({
                # 债权人基本信息
                "creditor_name": creditor.party_name or "",
                "creditor_type": creditor.party_type or "",
                "creditor_phone": creditor.phone or "",
                "creditor_legal_rep_name": creditor.name or "",  # 法人姓名
                
                # 债权人个人/公司信息
                "creditor_gender": creditor.gender or "",
                "creditor_birthday": creditor.birthday or "",
                "creditor_nation": creditor.nation or "",
                "creditor_address": creditor.address or "",
                "creditor_id_card": creditor.id_card or "",
                "creditor_company_name": creditor.company_name or "",
                "creditor_company_address": creditor.company_address or "",
                "creditor_company_code": creditor.company_code or "",
                
                # 债权人银行信息
                "creditor_bank_account": creditor.bank_account or "",
                "creditor_bank_address": creditor.bank_address or "",
                "creditor_bank_phone": creditor.bank_phone or "",
                "creditor_owner_name": creditor.owner_name or "",  # 银行持卡人姓名
            })
        else:
            # 如果没有债权人信息，使用空值
            case_data.update({
                "creditor_name": "",
                "creditor_type": "",
                "creditor_phone": "",
                "creditor_gender": "",
                "creditor_birthday": "",
                "creditor_nation": "",
                "creditor_address": "",
                "creditor_id_card": "",
                "creditor_company_name": "",
                "creditor_company_address": "",
                "creditor_company_code": "",
                "creditor_bank_account": "",
                "creditor_bank_address": "",
                "creditor_bank_phone": "",
                "creditor_owner_name": "",
            })
        
        # 添加债务人信息
        if debtor:
            logger.info(f"债务人信息: {debtor.party_name}, 法定代表人: {debtor.name}")
            case_data.update({
                # 债务人基本信息
                "debtor_name": debtor.party_name or "",
                "debtor_type": debtor.party_type or "",
                "debtor_phone": debtor.phone or "",
                "debtor_legal_rep_name": debtor.name or "",  # 法人姓名
                
                # 债务人个人/公司信息
                "debtor_gender": debtor.gender or "",
                "debtor_birthday": debtor.birthday or "",
                "debtor_nation": debtor.nation or "",
                "debtor_address": debtor.address or "",
                "debtor_id_card": debtor.id_card or "",
                "debtor_company_name": debtor.company_name or "",
                "debtor_company_address": debtor.company_address or "",
                "debtor_company_code": debtor.company_code or "",
                
                # 债务人银行信息
                "debtor_bank_account": debtor.bank_account or "",
                "debtor_bank_address": debtor.bank_address or "",
                "debtor_bank_phone": debtor.bank_phone or "",
                "debtor_owner_name": debtor.owner_name or "",  # 银行持卡人姓名
            })
        else:
            # 如果没有债务人信息，使用空值
            case_data.update({
                "debtor_name": "",
                "debtor_type": "",
                "debtor_phone": "",
                "debtor_gender": "",
                "debtor_birthday": "",
                "debtor_nation": "",
                "debtor_address": "",
                "debtor_id_card": "",
                "debtor_company_name": "",
                "debtor_company_address": "",
                "debtor_company_code": "",
                "debtor_bank_account": "",
                "debtor_bank_address": "",
                "debtor_bank_phone": "",
                "debtor_owner_name": "",
            })
        
        return case_data

    def generate_document(
        self, 
        template_id: str,
        case_data: Dict[str, Any],
        custom_variables: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """
        生成文书
        
        Args:
            template_id: 模板ID
            case_data: 案件数据
            custom_variables: 自定义变量
            
        Returns:
            生成结果
        """
        try:
            # 获取模板
            template = self.get_template(template_id)
            if not template:
                return {
                    "success": False,
                    "message": f"模板 {template_id} 不存在"
                }
            
            # 准备变量数据
            variables = self._prepare_variables(template, case_data, custom_variables)
            
            # 生成文档
            output_file_path = self._create_document(template, variables, case_data)
            
            # 生成记录信息
            record_id = str(uuid.uuid4())
            record_info = {
                "id": record_id,
                "template_id": template_id,
                "case_id": case_data.get("case_id"),
                "document_type": template.get("type"),
                "file_path": output_file_path,
                "filename": Path(output_file_path).name,
                "variables_used": variables,
                "generated_at": datetime.now()
            }
            
            return {
                "success": True,
                "message": "文书生成成功",
                "file_path": output_file_path,
                "filename": Path(output_file_path).name,
                "record_info": record_info
            }
            
        except Exception as e:
            logger.error(f"生成文书失败: {str(e)}")
            return {
                "success": False,
                "message": f"生成文书失败: {str(e)}"
            }
    
    def _prepare_variables(
        self, 
        template: Dict[str, Any], 
        case_data: Dict[str, Any],
        custom_variables: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """准备模板变量"""
        # 确保case_data不为空
        case_data = case_data or {}
        custom_variables = custom_variables or {}
        
        # 格式化loan_amount，去除尾随零
        loan_amount = case_data.get("loan_amount", 0.0)
        if isinstance(loan_amount, (int, float)):
            # 如果是整数，直接显示；如果是小数，去除尾随零
            if loan_amount == int(loan_amount):
                formatted_loan_amount = str(int(loan_amount))
            else:
                formatted_loan_amount = f"{loan_amount:.2f}".rstrip('0').rstrip('.')
        else:
            formatted_loan_amount = str(loan_amount)
        
        # 格式化created_at为中文日期格式
        created_at = case_data.get("created_at")
        if created_at:
            if isinstance(created_at, datetime):
                formatted_created_at = created_at.strftime("%Y年%m月%d日")
            elif isinstance(created_at, str):
                try:
                    # 尝试解析ISO格式日期
                    parsed_date = datetime.fromisoformat(created_at.replace('Z', '+00:00'))
                    formatted_created_at = parsed_date.strftime("%Y年%m月%d日")
                except:
                    formatted_created_at = created_at
            else:
                formatted_created_at = str(created_at)
        else:
            formatted_created_at = datetime.now().strftime("%Y年%m月%d日")
        
        # 基础案件信息
        variables = {
            "title": "民事起诉状",
            "case_type": case_data.get("case_type", "货款纠纷"),
            "case_status": case_data.get("case_status", "draft"),
            
            # 当事人信息（动态生成，不包含标签）
            "creditor_info": self._generate_party_info(case_data, "creditor"),
            "debtor_info": self._generate_party_info(case_data, "debtor"),
            "creditor_name": case_data.get("creditor_name", ""),
            "debtor_name": case_data.get("debtor_name", ""),
            
            # 保留原有字段以兼容其他模板
            "creditor_type": case_data.get("creditor_type", ""),
            "creditor_gender": case_data.get("creditor_gender", ""),
            "creditor_birthday": case_data.get("creditor_birthday", ""),
            "creditor_nation": case_data.get("creditor_nation", "汉族"),
            "creditor_address": case_data.get("creditor_address", ""),
            "creditor_id_card": case_data.get("creditor_id_card", ""),
            "creditor_phone": case_data.get("creditor_phone", ""),
            "creditor_company_name": case_data.get("creditor_company_name", ""),
            "creditor_company_address": case_data.get("creditor_company_address", ""),
            "creditor_company_code": case_data.get("creditor_company_code", ""),
            "creditor_bank_account": case_data.get("creditor_bank_account", ""),
            "creditor_bank_address": case_data.get("creditor_bank_address", ""),
            "creditor_bank_phone": case_data.get("creditor_bank_phone", ""),
            "creditor_owner_name": case_data.get("creditor_owner_name", ""),
            
            "debtor_type": case_data.get("debtor_type", ""),
            "debtor_gender": case_data.get("debtor_gender", ""),
            "debtor_birthday": case_data.get("debtor_birthday", ""),
            "debtor_nation": case_data.get("debtor_nation", "汉族"),
            "debtor_address": case_data.get("debtor_address", ""),
            "debtor_id_card": case_data.get("debtor_id_card", ""),
            "debtor_phone": case_data.get("debtor_phone", ""),
            "debtor_company_name": case_data.get("debtor_company_name", ""),
            "debtor_company_address": case_data.get("debtor_company_address", ""),
            "debtor_company_code": case_data.get("debtor_company_code", ""),
            "debtor_bank_account": case_data.get("debtor_bank_account", ""),
            "debtor_bank_address": case_data.get("debtor_bank_address", ""),
            "debtor_bank_phone": case_data.get("debtor_bank_phone", ""),
            "debtor_owner_name": case_data.get("debtor_owner_name", ""),
            
            # 案件详情
            "loan_amount": formatted_loan_amount,
            "loan_date": case_data.get("loan_date", ""),
            "case_description": case_data.get("description", ""),
            
            # 自动生成的变量
            "current_date": datetime.now().strftime("%Y年%m月%d日"),
            "case_title": self._generate_case_title(case_data),
            
            # 诉讼内容（智能生成）
            "claims": self._generate_claims(case_data),
            "reasons": self._generate_reasons(case_data),
            
            # 结尾信息
            "court_name": case_data.get("court_name", "某某人民法院"),
            "court_address": case_data.get("court_name", "某某人民法院"),
            "created_at": formatted_created_at
        }
        
        # 合并自定义变量，自定义变量优先级最高
        if custom_variables:
            variables.update(custom_variables)
        
        return variables
    
    def _generate_party_info(self, case_data: Dict[str, Any], party_role: str) -> str:
        """根据当事人类型动态生成当事人信息"""
        if party_role == "creditor":
            party_type = case_data.get("creditor_type", "")
            party_name = case_data.get("creditor_name", "")
            party_phone = case_data.get("creditor_phone", "")
            
            # 个人信息
            gender = case_data.get("creditor_gender", "")
            birthday = case_data.get("creditor_birthday", "")
            nation = case_data.get("creditor_nation", "")
            address = case_data.get("creditor_address", "")
            id_card = case_data.get("creditor_id_card", "")
            name = case_data.get("creditor_legal_rep_name", "")  # 法人姓名
            
            # 公司信息
            company_name = case_data.get("creditor_company_name", "")
            company_address = case_data.get("creditor_company_address", "")
            company_code = case_data.get("creditor_company_code", "")
            owner_name = case_data.get("creditor_owner_name", "")  # 银行持卡人姓名
        else:  # debtor
            party_type = case_data.get("debtor_type", "")
            party_name = case_data.get("debtor_name", "")
            party_phone = case_data.get("debtor_phone", "")
            
            # 个人信息
            gender = case_data.get("debtor_gender", "")
            birthday = case_data.get("debtor_birthday", "")
            nation = case_data.get("debtor_nation", "")
            address = case_data.get("debtor_address", "")
            id_card = case_data.get("debtor_id_card", "")
            name = case_data.get("debtor_legal_rep_name", "")  # 法人姓名
            
            # 公司信息
            company_name = case_data.get("debtor_company_name", "")
            company_address = case_data.get("debtor_company_address", "")
            company_code = case_data.get("debtor_company_code", "")
            owner_name = case_data.get("debtor_owner_name", "")  # 银行持卡人姓名
        
        if not party_name:
            return ""
        
        # 根据当事人类型生成不同的信息格式
        if party_type == "company":
            # 公司信息格式 - 按照法律文书标准格式
            # 第一行：公司基本信息
            basic_info_parts = []
            basic_info_parts.append(party_name)
            
            if company_address:
                basic_info_parts.append(f"住所地{company_address}")
            
            if company_code:
                basic_info_parts.append(f"统一社会信用代码{company_code}")
            
            basic_info = "，".join(basic_info_parts) + "。"
            
            # 第二行：法定代表人信息（换行并与"原告"/"被告"对齐）
            legal_rep_info = ""
            if name:
                # 计算缩进：与"原告"/"被告"对齐（约2个字符的缩进）
                indent = "　　"  # 使用2个全角空格缩进
                legal_rep_parts = [f"法定代表人：{name}"]
                
                # 添加职务信息
                if party_role == "creditor":
                    legal_rep_parts.append("职务：执行董事兼经理")
                elif party_role == "debtor":
                    legal_rep_parts.append("职务：执行董事兼总经理")
                
                # 添加联系电话
                if party_phone:
                    if party_role == "creditor":
                        legal_rep_parts.append(f"代理人电话：{party_phone}")
                    else:
                        legal_rep_parts.append(f"联系电话：{party_phone}")
                
                legal_rep_info = indent + "，".join(legal_rep_parts)
            # 组合两行信息
            if legal_rep_info:
                return basic_info + "\n" + legal_rep_info
            else:
                return basic_info
        
        elif party_type == "person":
            # 个人信息格式
            info_parts = [party_name]
            
            if gender:
                info_parts.append(gender)
            
            if birthday:
                info_parts.append(f"{birthday}出生")
            
            if nation:
                info_parts.append(f"{nation}族")
            
            if address:
                info_parts.append(f"住{address}")
            
            if id_card:
                info_parts.append(f"公民身份号码{id_card}")
            
            if party_phone:
                info_parts.append(f"联系电话{party_phone}")
            
            return "，".join(info_parts)
        
        else:
            # 默认格式（个体工商户或其他类型）
            info_parts = [party_name]
            
            if company_name and company_name != party_name:
                info_parts.append(f"（{company_name}）")
            
            if company_address:
                info_parts.append(f"住所地{company_address}")
            
            if company_code:
                info_parts.append(f"统一社会信用代码{company_code}")
            
            if owner_name:
                info_parts.append(f"经营者{owner_name}")
            
            if party_phone:
                info_parts.append(f"联系电话{party_phone}")
            
            return "，".join(info_parts) + "。"
    
    def _generate_case_title(self, case_data: Dict[str, Any]) -> str:
        """生成案件标题"""
        creditor_name = case_data.get("creditor_name", "")
        debtor_name = case_data.get("debtor_name", "")
        case_type = case_data.get("case_type", "货款纠纷")
        
        # 根据案件类型调整标题
        if case_type == "CONTRACT":
            case_type_display = "买卖合同纠纷"
        elif case_type == "DEBT":
            case_type_display = "民间借贷纠纷"
        else:
            case_type_display = "货款纠纷"
        
        if creditor_name and debtor_name:
            return f"{creditor_name}诉{debtor_name}{case_type_display}案"
        elif creditor_name:
            return f"{creditor_name}诉被告{case_type_display}案"
        elif debtor_name:
            return f"原告诉{debtor_name}{case_type_display}案"
        else:
            return f"民事{case_type_display}案"
    
    def _generate_claims(self, case_data: Dict[str, Any]) -> str:
        """智能生成诉讼请求"""
        case_type = case_data.get("case_type", "DEBT")
        loan_amount = case_data.get("loan_amount", 0)
        
        # 格式化loan_amount，去除尾随零
        if isinstance(loan_amount, (int, float)):
            if loan_amount == int(loan_amount):
                formatted_amount = str(int(loan_amount))
            else:
                formatted_amount = f"{loan_amount:.2f}".rstrip('0').rstrip('.')
        else:
            formatted_amount = str(loan_amount)
        
        # 确保案件类型是大写字符串
        case_type = str(case_type).upper()
        
        if case_type == "DEBT":
            # 民间借贷纠纷
            return f"一、请求判令被告偿还原告借款本金{formatted_amount}元及逾期利息（自起诉之日起，按全国银行间同业拆借中心公布的一年期贷款市场报价利率计算至借款实际清偿完毕之日止）。"
        elif case_type == "CONTRACT":
            # 买卖合同纠纷
            return f"一、请求判令被告向原告支付货款{formatted_amount}元及逾期付款损失（自起诉之日起，按全国银行间同业拆借中心公布的一年期贷款市场报价利率加计50%计算至货款实际清偿完毕之日止）。"
        else:
            # 默认使用借款纠纷
            return f"一、请求判令被告偿还原告借款本金{formatted_amount}元及逾期利息（自起诉之日起，按全国银行间同业拆借中心公布的一年期贷款市场报价利率计算至借款实际清偿完毕之日止）。"

    
    def _generate_reasons(self, case_data: Dict[str, Any]) -> str:
        """智能生成事实与理由"""
        case_type = case_data.get("case_type", "DEBT")
        loan_amount = case_data.get("loan_amount", 0)
        creditor_type = case_data.get("creditor_type", "")
        creditor_company_name = case_data.get("creditor_company_name", "")
        debtor_company_name = case_data.get("debtor_company_name", "")
        
        # 格式化loan_amount，去除尾随零
        if isinstance(loan_amount, (int, float)):
            if loan_amount == int(loan_amount):
                formatted_amount = str(int(loan_amount))
            else:
                formatted_amount = f"{loan_amount:.2f}".rstrip('0').rstrip('.')
        else:
            formatted_amount = str(loan_amount)
        
        # 确保案件类型是大写字符串
        case_type = str(case_type).upper()
        
        if case_type == "DEBT":
            # 民间借贷纠纷
            return f"原被告系朋友关系，xx年原告陆续出借被告借款{formatted_amount}元，截至起诉之日，被告余欠原告借款{formatted_amount}元。原告多次催讨未果，故双方纠纷成讼。"
        elif case_type == "CONTRACT":
            # 买卖合同纠纷 - 根据当事人类型生成不同的理由
            if creditor_company_name and debtor_company_name:
                return f"原告{creditor_company_name}与被告{debtor_company_name}之间素有交易往来。截至起诉之日，被告余欠原告货款{formatted_amount}元。经原告多次催讨，被告仍未履行付款义务，故双方纠纷成讼。"
            elif creditor_company_name:
                return f"原告{creditor_company_name}与被告之间素有交易往来。截至起诉之日，被告余欠原告货款{formatted_amount}元。经原告多次催讨，被告仍未履行付款义务，故双方纠纷成讼。"
            else:
                return f"原告与被告之间素有交易往来。截至起诉之日，被告余欠原告货款{formatted_amount}元。经原告多次催讨，被告仍未履行付款义务，故双方纠纷成讼。"
        else:
            # 默认使用借款纠纷
            return f"原被告系朋友关系，xx年原告陆续出借被告借款{formatted_amount}元，截至起诉之日，被告余欠原告借款{formatted_amount}元。原告多次催讨未果，故双方纠纷成讼。"
    
    def _create_document(
        self, 
        template: Dict[str, Any], 
        variables: Dict[str, Any],
        case_data: Dict[str, Any]
    ) -> str:
        """创建Word文档"""
        try:
            # 检查模板文件是否存在
            template_path = Path(template.get("file_path", ""))
            if not template_path.exists():
                # 如果模板文件不存在，创建默认模板
                template_path = self._create_default_template(template, variables)
            
            # 使用docxtpl生成文档
            doc = DocxTemplate(template_path)
            doc.render(variables)
            
            # 生成输出文件名
            output_filename = f"{template.get('type', 'document')}_{case_data.get('case_id', 'unknown')}_{uuid.uuid4().hex[:8]}.docx"
            output_path = self.output_dir / output_filename
            
            # 保存文档
            doc.save(output_path)
            
            return str(output_path)
            
        except Exception as e:
            logger.error(f"创建文档失败: {str(e)}")
            raise
    
    def _create_default_template(
        self, 
        template: Dict[str, Any], 
        variables: Dict[str, Any]
    ) -> Path:
        """创建默认模板文件"""
        try:
            # 使用python-docx创建简单的默认模板
            doc = Document()
            
            # 添加标题
            title = doc.add_heading(variables.get("case_title", "起诉状"), 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加基本信息
            doc.add_paragraph(f"原告：{variables.get('creditor_name', '')}")
            doc.add_paragraph(f"被告：{variables.get('debtor_name', '')}")
            doc.add_paragraph(f"案由：{variables.get('case_type', '')}")
            
            # 添加正文内容
            doc.add_heading("诉讼请求", level=1)
            doc.add_paragraph("1. 请求法院判令被告偿还借款本金及利息")
            if variables.get("loan_amount"):
                doc.add_paragraph(f"2. 借款金额：{variables.get('loan_amount')}元")
            
            doc.add_heading("事实与理由", level=1)
            doc.add_paragraph(variables.get("case_description", "根据案件具体情况填写..."))
            
            # 保存模板文件
            template_path = self.template_dir / f"{template.get('type', 'document')}_template.docx"
            doc.save(str(template_path))
            
            return template_path
            
        except Exception as e:
            logger.error(f"创建默认模板失败: {str(e)}")
            # 如果docxtpl创建失败，创建简单的文本文件作为占位符
            template_path = self.template_dir / f"{template.get('type', 'document')}_template.txt"
            with open(template_path, 'w', encoding='utf-8') as f:
                f.write(f"案件标题：{variables.get('case_title', '起诉状')}\n")
                f.write(f"原告：{variables.get('creditor_name', '')}\n")
                f.write(f"被告：{variables.get('debtor_name', '')}\n")
                f.write(f"案由：{variables.get('case_type', '')}\n")
            
            return template_path
    
    def create_sample_templates(self):
        """创建示例模板文件"""
        try:
            # 创建起诉状示例模板
            complaint_template = self.template_dir / "complaint_template.docx"
            if not complaint_template.exists():
                self._create_sample_complaint_template()
            
            # 创建证据清单示例模板
            evidence_template = self.template_dir / "evidence_list_template.docx"
            if not evidence_template.exists():
                self._create_sample_evidence_template()
                
            logger.info("示例模板文件创建完成")
            
        except Exception as e:
            logger.error(f"创建示例模板文件失败: {str(e)}")
    
    def _create_sample_complaint_template(self):
        """创建起诉状示例模板"""
        try:
            doc = Document()
            
            # 添加标题
            doc.add_heading("{{case_title}}", 0)
            
            # 添加基本信息
            doc.add_paragraph("原告：{{creditor_name}}")
            doc.add_paragraph("被告：{{debtor_name}}")
            doc.add_paragraph("案由：{{case_type}}")
            
            # 添加正文内容
            doc.add_heading("诉讼请求", level=1)
            doc.add_paragraph("1. 请求法院判令被告偿还借款本金及利息")
            doc.add_paragraph("2. 借款金额：{{loan_amount}}元")
            
            doc.add_heading("事实与理由", level=1)
            doc.add_paragraph("{{case_description}}")
            
            # 保存模板
            template_path = self.template_dir / "complaint_template.docx"
            doc.save(str(template_path))
            
        except Exception as e:
            logger.error(f"创建起诉状示例模板失败: {str(e)}")
    
    def _create_sample_evidence_template(self):
        """创建证据清单示例模板"""
        try:
            doc = Document()
            
            # 添加标题
            doc.add_heading("{{case_title}}", 0)
            doc.add_paragraph("案件编号：{{case_id}}")
            
            # 添加证据清单
            doc.add_heading("证据材料清单", level=1)
            doc.add_paragraph("证据数量：{{evidence_count}}")
            doc.add_paragraph("总页数：{{total_pages}}")
            
            # 保存模板
            template_path = self.template_dir / "evidence_list_template.docx"
            doc.save(str(template_path))
            
        except Exception as e:
            logger.error(f"创建证据清单示例模板失败: {str(e)}")
    
    def get_health_status(self) -> Dict[str, Any]:
        """获取健康状态"""
        try:
            template_dir_exists = self.template_dir.exists()
            output_dir_exists = self.output_dir.exists()
            
            return {
                "status": "healthy",
                "template_dir_exists": template_dir_exists,
                "output_dir_exists": output_dir_exists,
                "template_dir": str(self.template_dir),
                "output_dir": str(self.output_dir)
            }
        except Exception as e:
            return {
                "status": "unhealthy",
                "error": str(e)
            }

    def _build_case_data_from_template(self, template: Dict[str, Any], custom_variables: Dict[str, Any]) -> Dict[str, Any]:
        """从模板配置构建案件数据"""
        case_data = {}
        
        # 从模板配置中获取变量定义
        template_variables = template.get("variables", [])
        
        for var_def in template_variables:
            var_name = var_def.get("name", "")
            default_value = var_def.get("default", "")
            
            # 优先使用自定义变量，其次使用模板默认值
            if var_name in custom_variables:
                case_data[var_name] = custom_variables[var_name]
            else:
                case_data[var_name] = default_value
        
        # 设置一些基础默认值
        case_data.setdefault("case_type", "货款纠纷")
        case_data.setdefault("creditor_nation", "汉族")
        case_data.setdefault("debtor_nation", "汉族")
        case_data.setdefault("court_address", "某某人民法院")
        
        return case_data
