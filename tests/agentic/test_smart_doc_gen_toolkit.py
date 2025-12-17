import pytest
import docx
import os
import json
from app.agentic.tools.smart_doc_gen_toolkit import SmartDocGenToolkit

@pytest.fixture
def temp_docx(tmp_path):
    doc = docx.Document()
    doc.add_paragraph("Title Paragraph")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Label"
    table.cell(0, 1).text = "Value"
    
    path = tmp_path / "test.docx"
    doc.save(path)
    return str(path)

def test_analyze_template(temp_docx):
    toolkit = SmartDocGenToolkit()
    result_json = toolkit.analyze_template(temp_docx)
    result = json.loads(result_json)
    
    assert "type" in result
    assert "stats" in result
    # "Title Paragraph" (15 chars) vs "LabelValue" (9 chars) -> likely narrative or unknown depending on thresholds,
    # but let's just check structure.
    # Actually logic: if t_chars > p_chars: element.
    # p_chars = 15. t_chars = 9. -> narrative.
    assert result["type"] == "narrative"

def test_extract_structure(temp_docx):
    toolkit = SmartDocGenToolkit()
    result_json = toolkit.extract_structure(temp_docx)
    blocks = json.loads(result_json)
    
    assert len(blocks) >= 2
    # Check paragraph
    p = next(b for b in blocks if b["type"] == "paragraph")
    assert p["text"] == "Title Paragraph"
    assert "p_" in p["id"]
    
    # Check table
    t = next(b for b in blocks if b["type"] == "table")
    assert "t_" in t["id"]
    assert len(t["rows"]) == 1
    assert t["rows"][0][0]["text"] == "Label"

def test_fill_template(temp_docx, tmp_path):
    toolkit = SmartDocGenToolkit()
    
    # Get IDs first
    structure = json.loads(toolkit.extract_structure(temp_docx))
    p_id = next(b for b in structure if b["type"] == "paragraph")["id"]
    t_id_parts = next(b for b in structure if b["type"] == "table")["id"]
    # We need specific cell ID. 
    # extract_structure output for table: 
    # "rows": [ [ {"id": "...", "text": "..."}, ... ] ]
    # t_0_r0_c1 is likely the Value cell.
    
    # Let's derive it or reuse known pattern since we just created it.
    # It should be t_0_r0_c1 if it's the 0th table.
    
    fillings = {
        p_id: "New Title",
        "t_0_r0_c1": "Filled Value"
    }
    
    output_path = tmp_path / "filled.docx"
    result = toolkit.fill_template(temp_docx, str(output_path), json.dumps(fillings))
    assert "success" in result
    
    # Verify content
    doc = docx.Document(output_path)
    assert doc.paragraphs[0].text == "New Title"
    assert doc.tables[0].cell(0, 1).text == "Filled Value"

def test_delete_paragraph(temp_docx, tmp_path):
    toolkit = SmartDocGenToolkit()
    structure = json.loads(toolkit.extract_structure(temp_docx))
    p_id = next(b for b in structure if b["type"] == "paragraph")["id"]
    
    fillings = {
        p_id: "__DELETE__"
    }
    
    output_path = tmp_path / "deleted.docx"
    toolkit.fill_template(temp_docx, str(output_path), json.dumps(fillings))
    
    doc = docx.Document(output_path)
    # The paragraph should be gone.
    # Note: docx.Document() creates an empty paragraph by default sometimes.
    # But our original doc had 1 paragraph. If deleted, it might have 0 or just empty depending on impl.
    # Let's check text is not "Title Paragraph"
    for p in doc.paragraphs:
        assert "Title Paragraph" not in p.text
