import pytest
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import RGBColor

from app.template_editor.mappers import DocxToProseMirrorMapper


def _find_text_style_mark(node):
    content = node.get("content", [])
    if not content:
        return None
    marks = content[0].get("marks", [])
    for mark in marks:
        if mark.get("type") == "textStyle":
            return mark
    return None


def test_paragraph_list_and_line_height_attrs():
    doc = Document()
    ordered_para = doc.add_paragraph("Ordered item")
    ordered_para.style = "List Number"
    ordered_para.paragraph_format.line_spacing = 1.5
    run = ordered_para.runs[0]
    run.font.color.rgb = RGBColor(0x11, 0x22, 0x33)
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    run.font.small_caps = True

    bullet_para = doc.add_paragraph("Bullet item")
    bullet_para.style = "List Bullet"

    mapper = DocxToProseMirrorMapper()
    result = mapper.map_document(doc)

    ordered_node = result["content"][0]
    bullet_node = result["content"][1]

    ordered_list = ordered_node["attrs"]["list"]
    assert ordered_list["type"] == "ordered"
    assert ordered_list["level"] == 0
    assert pytest.approx(1.5) == ordered_node["attrs"]["lineHeight"]

    bullet_list = bullet_node["attrs"]["list"]
    assert bullet_list["type"] == "unordered"

    style_mark = _find_text_style_mark(ordered_node)
    assert style_mark is not None
    style_attrs = style_mark["attrs"]
    assert style_attrs["color"] == "#112233"
    assert style_attrs["backgroundColor"] == "#ffff00"
    assert style_attrs["smallCaps"] is True


def test_docx_to_prosemirror_handles_newlines():
    doc = Document()
    para = doc.add_paragraph()
    run = para.add_run("第一行\n第二行")

    mapper = DocxToProseMirrorMapper()
    result = mapper.map_document(doc)

    paragraph_node = result["content"][0]
    nodes = paragraph_node["content"]
    assert nodes[0]["text"] == "第一行"
    assert nodes[1]["type"] == "hardBreak"
    assert nodes[2]["text"] == "第二行"

