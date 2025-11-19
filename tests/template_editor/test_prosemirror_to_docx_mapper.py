from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

from app.template_editor.mappers import (
    DocxToProseMirrorMapper,
    ProseMirrorToDocxMapper,
)


def test_prosemirror_to_docx_applies_paragraph_and_run_styles():
    pm_json = {
        "type": "doc",
        "content": [
            {
                "type": "paragraph",
                "attrs": {
                    "textAlign": "center",
                    "lineHeight": 18,
                    "list": {"type": "ordered", "level": 1},
                },
                "content": [
                    {
                        "type": "text",
                        "text": "格式测试",
                        "marks": [
                            {"type": "bold"},
                            {
                                "type": "textStyle",
                                "attrs": {
                                    "color": "#123456",
                                    "fontSize": "14pt",
                                    "fontFamily": "Arial",
                                    "smallCaps": True,
                                },
                            },
                        ],
                    }
                ],
            }
        ],
    }

    exporter = ProseMirrorToDocxMapper()
    importer = DocxToProseMirrorMapper()

    doc = exporter.map_document(pm_json)
    para = doc.paragraphs[0]

    assert para.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert abs(para.paragraph_format.line_spacing.pt - 18) < 0.1
    assert para.style.name in {"List Number", "Ordered List"}

    run = para.runs[0]
    assert run.bold
    assert str(run.font.color.rgb) == "123456"
    assert abs(run.font.size.pt - 14) < 0.1
    assert run.font.name == "Arial"
    assert run.font.small_caps

    roundtrip = importer.map_document(doc)
    para_node = roundtrip["content"][0]
    assert para_node["attrs"]["list"]["type"] == "ordered"
    text_marks = para_node["content"][0]["marks"]
    text_style = next(
        (mark for mark in text_marks if mark.get("type") == "textStyle"), {}
    )
    assert text_style.get("attrs", {}).get("color") == "#123456"
    assert exporter.export_warnings == []


def test_line_height_multiple_mapping():
    pm_json = {
        "type": "doc",
        "content": [
            {
                "type": "paragraph",
                "attrs": {
                    "lineHeight": 0.9,
                },
                "content": [{"type": "text", "text": "多行距"}],
            }
        ],
    }

    exporter = ProseMirrorToDocxMapper()
    doc = exporter.map_document(pm_json)
    para = doc.paragraphs[0]

    assert para.paragraph_format.line_spacing_rule == WD_LINE_SPACING.MULTIPLE
    assert abs(para.paragraph_format.line_spacing - 0.9) < 0.01


def test_null_attrs_do_not_break_export():
    pm_json = {
        "type": "doc",
        "content": [
            {
                "type": "paragraph",
                "attrs": {
                    "textAlign": "left",
                    "indent": None,
                    "firstLineIndent": None,
                    "lineHeight": None,
                    "spacing": {"before": None, "after": 6},
                },
                "content": [{"type": "text", "text": "空值处理"}],
            }
        ],
    }

    exporter = ProseMirrorToDocxMapper()
    doc = exporter.map_document(pm_json)
    para = doc.paragraphs[0]

    assert para.paragraph_format.left_indent is None
    assert para.paragraph_format.first_line_indent is None
    assert para.paragraph_format.space_before is None
    assert abs(para.paragraph_format.space_after.pt - 6) < 0.01
    assert exporter.export_warnings == []


def test_prosemirror_table_attrs_applied_on_export():
    pm_json = {
        "type": "doc",
        "content": [
            {
                "type": "table",
                "attrs": {
                    "style": "Table Grid",
                    "colWidths": [2000, 3000],
                    "tableWidth": {"width": 5000, "type": "dxa"},
                    "tableLayout": "fixed",
                },
                "content": [
                    {
                        "type": "tableRow",
                        "content": [
                            {
                                "type": "tableCell",
                                "attrs": {
                                    "colspan": 1,
                                    "rowspan": 1,
                                    "cellWidth": {"width": 2000, "type": "dxa"},
                                    "backgroundColor": "#ABCDEF",
                                },
                                "content": [
                                    {
                                        "type": "paragraph",
                                        "attrs": {},
                                        "content": [{"type": "text", "text": "A"}],
                                    }
                                ],
                            },
                            {
                                "type": "tableCell",
                                "attrs": {
                                    "colspan": 1,
                                    "rowspan": 1,
                                    "cellWidth": {"width": 3000, "type": "dxa"},
                                },
                                "content": [
                                    {
                                        "type": "paragraph",
                                        "attrs": {},
                                        "content": [{"type": "text", "text": "B"}],
                                    }
                                ],
                            },
                        ],
                    }
                ],
            }
        ],
    }

    exporter = ProseMirrorToDocxMapper()
    parser = DocxToProseMirrorMapper()

    doc = exporter.map_document(pm_json)
    assert exporter.export_warnings == []
    table = doc.tables[0]

    assert table.style.name == "Table Grid"

    grid_cols = table._tbl.tblGrid.gridCol_lst
    assert [int(col.get(qn("w:w"))) for col in grid_cols] == [2000, 3000]

    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    tbl_w = table._tbl.tblPr.find("./w:tblW", ns)
    assert tbl_w.get(qn("w:w")) == "5000"
    assert tbl_w.get(qn("w:type")) == "dxa"

    tbl_layout = table._tbl.tblPr.find("./w:tblLayout", ns)
    assert tbl_layout.get(qn("w:type")) == "fixed"

    first_cell = table.cell(0, 0)
    shading = first_cell._tc.tcPr.find("./w:shd", ns)
    assert shading is not None
    assert shading.get(qn("w:fill")) == "ABCDEF"

    tcw = first_cell._tc.tcPr.tcW
    assert tcw is not None
    assert tcw.get(qn("w:w")) == "2000"

    roundtrip = parser.map_document(doc)
    table_node = roundtrip["content"][0]
    assert table_node["attrs"]["colWidths"] == [2000, 3000]
    first_cell_attrs = table_node["content"][0]["content"][0]["attrs"]
    assert first_cell_attrs["backgroundColor"] == "#ABCDEF"

