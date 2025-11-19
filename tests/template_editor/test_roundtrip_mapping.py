from docx import Document
from docx.shared import RGBColor

from app.template_editor.mappers import (
    DocxToProseMirrorMapper,
    ProseMirrorToDocxMapper,
)


def _get_first_text_style_marks(node):
    for child in node.get("content", []):
        if child.get("type") == "text":
            return child.get("marks", [])
        marks = _get_first_text_style_marks(child)
        if marks:
            return marks
    return []


def test_docx_to_pm_roundtrip_preserves_basic_styles():
    doc = Document()
    doc._body.clear_content()

    heading = doc.add_heading("章节标题", level=2)
    heading.alignment = 1  # center

    para = doc.add_paragraph()
    run = para.add_run("彩色文本")
    run.bold = True
    run.font.color.rgb = RGBColor(0x12, 0x34, 0x56)

    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"

    parser = DocxToProseMirrorMapper()
    exporter = ProseMirrorToDocxMapper()

    pm_json = parser.map_document(doc)
    roundtrip_doc = exporter.map_document(pm_json)
    pm_roundtrip = parser.map_document(roundtrip_doc)

    heading_node = pm_roundtrip["content"][0]
    assert heading_node["type"] == "heading"
    assert heading_node["attrs"]["level"] == 2

    paragraph_node = pm_roundtrip["content"][1]
    style_marks = _get_first_text_style_marks(paragraph_node)
    text_style = next(
        (mark for mark in style_marks if mark.get("type") == "textStyle"), {}
    )
    assert any(mark.get("type") == "bold" for mark in style_marks)
    assert text_style.get("attrs", {}).get("color") == "#123456"

    table_node = pm_roundtrip["content"][2]
    original_table = pm_json["content"][2]
    assert table_node["type"] == "table"
    assert table_node["attrs"].get("colWidths") == original_table["attrs"].get(
        "colWidths"
    )

