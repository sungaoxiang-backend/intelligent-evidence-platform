from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from app.template_editor.mappers import DocxToProseMirrorMapper


def test_table_attrs_include_colwidths_and_style():
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"

    for grid_col, width in zip(
        table._tbl.tblGrid.findall(".//w:gridCol", {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}),
        (2000, 3000),
    ):
        grid_col.set(qn("w:w"), str(width))

    mapper = DocxToProseMirrorMapper()
    table_node = mapper.map_document(doc)["content"][0]

    assert table_node["attrs"]["style"] == "Table Grid"
    assert table_node["attrs"]["colWidths"] == [2000, 3000]


def test_table_cell_attrs_include_background_and_rowspan():
    doc = Document()
    table = doc.add_table(rows=2, cols=2)

    merged = table.cell(0, 0)
    merged.merge(table.cell(1, 0))
    tc = merged._tc
    tc_pr = tc.tcPr
    if tc_pr is None:
        tc_pr = OxmlElement("w:tcPr")
        tc.insert(0, tc_pr)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "ABCDEF")
    tc_pr.append(shd)

    mapper = DocxToProseMirrorMapper()
    table_node = mapper.map_document(doc)["content"][0]
    first_cell_attrs = table_node["content"][0]["content"][0]["attrs"]

    assert first_cell_attrs["rowspan"] == 2
    assert first_cell_attrs["backgroundColor"] == "#ABCDEF"

