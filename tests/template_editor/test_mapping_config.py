import pytest
from docx import Document
from docx.shared import RGBColor

from app.template_editor.mapping_config import hex_to_rgb, rgb_to_hex


def test_rgb_to_hex_from_color_format():
    doc = Document()
    run = doc.add_paragraph().add_run("color")
    run.font.color.rgb = RGBColor(0x12, 0x34, 0xAB)

    assert rgb_to_hex(run.font.color) == "#1234AB"


def test_rgb_to_hex_from_rgbcolor():
    color = RGBColor(0xDE, 0xAD, 0xBE)
    assert rgb_to_hex(color) == "#DEADBE"


def test_rgb_to_hex_invalid_input_returns_none():
    class Dummy:
        rgb = "XYZ"

    assert rgb_to_hex(None) is None
    assert rgb_to_hex(Dummy()) is None


@pytest.mark.parametrize(
    "hex_color,expected",
    [
        ("#123456", "123456"),
        ("123456", "123456"),
        ("#abcdef", "ABCDEF"),
    ],
)
def test_hex_to_rgb_valid(hex_color, expected):
    rgb = hex_to_rgb(hex_color)
    assert rgb is not None
    assert str(rgb) == expected


@pytest.mark.parametrize(
    "invalid_color",
    ["", "#12", "#XXXXXX", None],
)
def test_hex_to_rgb_invalid(invalid_color):
    assert hex_to_rgb(invalid_color) is None

